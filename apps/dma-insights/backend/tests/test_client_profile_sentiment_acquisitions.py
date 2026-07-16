"""D5 sentiment + acquisition history mined from the Client Profile report.

The Client Research/Profile report carries a "Sentiment Overview" table
(metric | value) and an "Acquisition History" table (Date | Target | …) that
the parser previously ignored. We mine them as faithful fallbacks for
`firmographics.sentiment` and the D5 acquisitions list (timeline_events
kind='acquisition'), with strict heading-anchoring + shape validation so a
wrong table under a drifted heading yields nothing rather than mismapped data.
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest

from app.services.parsers.client_profile import parse_client_profile_path

REAL = Path(__file__).parent / "fixtures" / "dma_packages_real_samples"
WSFS = REAL / "WSFS_Bank__DMA/04_reports/WSFS_Client_Profile_Research_Report.docx"
ALMA = REAL / "Alma_Bank__DMA/04_reports/AlmaBank_ClientProfile_Research_Report.docx"
NICOLA = REAL / "Nicola_Wealth__DMA/04_reports/NicolaWealth_ClientProfile_Research_Report.docx"
BATCHES = Path(__file__).parent / "fixtures" / "dma_packages_batches"
REGIONS_CP = (
    BATCHES / "batch_11/Regions Bank - DMA/04_reports"
    / "DMA_Client_Profile_RegionsBank_20260518.docx"
)


# ── Real-fixture assertions (WSFS is the report shape that qualifies) ──

def test_wsfs_sentiment_extracted() -> None:
    if not WSFS.exists():
        pytest.skip("WSFS fixture missing")
    sent = parse_client_profile_path(WSFS).sentiment
    sources = sent.get("sources", [])
    assert len(sources) >= 10
    assert all("source" in s and "rating" in s for s in sources)
    assert any("indeed" in s["source"].lower() for s in sources)


def test_wsfs_acquisitions_extracted() -> None:
    if not WSFS.exists():
        pytest.skip("WSFS fixture missing")
    acqs = parse_client_profile_path(WSFS).acquisition_events
    assert len(acqs) >= 2
    for a in acqs:
        assert a.kind == "acquisition"
        assert isinstance(a.event_date, date)
        assert a.title and a.title.strip()
    assert any("bryn mawr" in a.title.lower() for a in acqs)


@pytest.mark.parametrize("path", [ALMA, NICOLA])
def test_non_conforming_reports_yield_no_false_sentiment_or_acquisitions(path: Path) -> None:
    """Reports without a rating-shaped sentiment table / dated acquisition
    table must extract nothing — no mismapped entity-profile data."""
    if not path.exists():
        pytest.skip(f"fixture missing: {path}")
    r = parse_client_profile_path(path)
    assert r.sentiment.get("sources", []) == []
    assert r.acquisition_events == []


# ── Shape-validation unit tests (synthetic in-memory docx) ──

def _doc_with(heading: str, header: list[str] | None, rows: list[list[str]]):
    from docx import Document

    d = Document()
    d.add_heading(heading, level=2)
    cols = len(header) if header else len(rows[0])
    t = d.add_table(rows=0, cols=cols)
    if header:
        hc = t.add_row().cells
        for i, h in enumerate(header):
            hc[i].text = h
    for row in rows:
        rc = t.add_row().cells
        for i, v in enumerate(row):
            rc[i].text = v
    # terminate the section with a following heading
    d.add_heading("5. Next Section", level=2)
    return d


def test_sentiment_rating_table_accepted() -> None:
    from app.services.parsers.client_profile import parse_client_profile_doc

    d = _doc_with(
        "3.4 Sentiment Overview",
        None,
        [["Indeed Overall Rating", "3.6/5.0"], ["Culture & Values", "3.5/5.0"]],
    )
    sent = parse_client_profile_doc(d).sentiment
    assert {s["source"] for s in sent["sources"]} == {"Indeed Overall Rating", "Culture & Values"}


def test_sentiment_entity_kv_table_rejected() -> None:
    """A 'Field | Value' entity table under a sentiment heading is NOT
    rating-shaped → rejected (no fabricated sentiment)."""
    from app.services.parsers.client_profile import parse_client_profile_doc

    d = _doc_with(
        "3.4 Sentiment Overview",
        None,
        [["Entity Name", "Acme Bank"], ["Sub-Vertical", "Regional Banks"]],
    )
    assert parse_client_profile_doc(d).sentiment == {}


def test_acquisitions_dated_table_accepted() -> None:
    from app.services.parsers.client_profile import parse_client_profile_doc

    d = _doc_with(
        "4.4 Acquisition History",
        ["Date", "Target / Event", "Strategic Rationale", "Digital Implication"],
        [["Jan 2022", "Acme Corp", "Scale", "Data merge"],
         ["bad-date", "Skip Me", "x", "y"]],
    )
    acqs = parse_client_profile_doc(d).acquisition_events
    assert len(acqs) == 1  # undated row skipped
    assert acqs[0].title == "Acme Corp"
    assert acqs[0].event_date == date(2022, 1, 1)
    assert acqs[0].kind == "acquisition"


def test_acquisitions_undated_narrative_table_rejected() -> None:
    """An M&A strategy table with no date column is rejected entirely."""
    from app.services.parsers.client_profile import parse_client_profile_doc

    d = _doc_with(
        "4.4 M&A Activity",
        ["Aspect", "Detail"],
        [["M&A as Buyer", "IPO readiness"], ["M&A as Target", "OTCQX listing"]],
    )
    assert parse_client_profile_doc(d).acquisition_events == []


def test_acquisitions_entity_column_synonym_accepted() -> None:
    """Regions §4.4 labels the acquired-party column "Entity" (not
    "Target") — the synonym must match so the table isn't dropped whole."""
    from app.services.parsers.client_profile import parse_client_profile_doc

    d = _doc_with(
        "4.4 Acquisition History",
        ["Year", "Entity", "Type", "Integration Status"],
        [["2021", "EnerBank USA", "POS lending", "Integrated"],
         ["2022", "Sabal Capital Partners", "CRE lending", "Integrated"]],
    )
    acqs = parse_client_profile_doc(d).acquisition_events
    assert {a.title for a in acqs} == {"EnerBank USA", "Sabal Capital Partners"}
    assert {a.event_date.year for a in acqs} == {2021, 2022}


def test_acquisitions_dedup_on_target_year_not_datetime() -> None:
    """The same deal recorded at differing date grains ("2020" vs
    "Aug 2020") must collapse to ONE row (the Wintrust duplicate symptom),
    while a distinct same-year deal is preserved."""
    from app.services.parsers.client_profile import parse_client_profile_doc

    d = _doc_with(
        "4.4 Acquisition History",
        ["Date", "Entity", "Rationale"],
        [["2020", "Macatawa Bank", "MI expansion"],
         ["Aug 2020", "Macatawa Bank", "MI expansion (detail)"],
         ["2020", "Rush-Oak Corporation", "Chicago add"]],
    )
    acqs = parse_client_profile_doc(d).acquisition_events
    titles = [a.title for a in acqs]
    assert titles.count("Macatawa Bank") == 1  # differing-grain dup collapsed
    assert "Rush-Oak Corporation" in titles     # distinct same-year deal kept
    assert len(acqs) == 2


def test_regions_real_fixture_acquisitions_extracted() -> None:
    """Regions batch_11 §4.4 ("Entity"-headed table) → EnerBank 2021 +
    Sabal 2022 (was `acquisitions: 0` before the synonym fix)."""
    if not REGIONS_CP.exists():
        pytest.skip("Regions batch_11 fixture missing")
    acqs = parse_client_profile_path(REGIONS_CP).acquisition_events
    by_year = {a.event_date.year: a.title.lower() for a in acqs}
    assert 2021 in by_year and "enerbank" in by_year[2021]
    assert 2022 in by_year and "sabal" in by_year[2022]
    assert all(a.kind == "acquisition" for a in acqs)
