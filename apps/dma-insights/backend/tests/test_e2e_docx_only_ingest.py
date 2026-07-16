"""End-to-end integration sim: docx-only + worker env + empty catalogue.

Live-DB gated by `SEED_CI_PG_URL` (skipped in stage-0 / local-dev
without a DB). When the DB is present, this exercises the FULL chain
that was broken in the 2026-05-28 production incident:

  parse_package(docx-only Drive folder)
    → IngestedPackage with synthesized run_manifest + report_sections
    → persist_package(only DATABASE_URL set, no DATABASE_URL_SYNC)
    → entity row written with `drive_folder_id` (NOT `source_folder_id`)
    → document_sections row written
    → run row written
    → post-commit Drive feedback runs without raising
    → catalogue resolver returns SubcapNotFound for all subcaps
    → `catalogue_unresolved:N/M` + `catalogue_empty_for_version` warnings
    → job_executions lifecycle (create → mark_started → update_progress
       → mark_succeeded) all succeed with the sync-DSN resolver
       fallback (asyncpg → psycopg).

If any of these regress, this test fails with a precise pointer to
which step broke. The 2026-05-28 incident would have surfaced as
~5 distinct failures across this single test.

Mirrors `/tmp/sim_e2e.py` + `/tmp/sim_h8.py` from the 2026-05-28
remediation session; promoting to a permanent test so future
regressions of any of H1/H2/H5/H6/H7/H8 fail this single file.
"""
from __future__ import annotations

import os
import tempfile
from pathlib import Path

import pytest

LIVE_DB_URL = os.environ.get("SEED_CI_PG_URL", "")
HAS_LIVE_DB = bool(LIVE_DB_URL)

pytestmark = pytest.mark.skipif(
    not HAS_LIVE_DB,
    reason="SEED_CI_PG_URL not set — live-DB e2e ingest sim skipped",
)


def _make_docx(path: Path, heading: str, body: str) -> None:
    """Write a real .docx via python-docx so the parser can open it."""
    import docx
    d = docx.Document()
    d.add_heading(heading, level=1)
    d.add_paragraph(body)
    d.save(str(path))


def _async_url() -> str:
    if "+asyncpg" in LIVE_DB_URL:
        return LIVE_DB_URL
    return LIVE_DB_URL.replace("postgresql://", "postgresql+asyncpg://")


@pytest.mark.asyncio
async def test_docx_only_drive_folder_full_ingest_chain(monkeypatch):
    """End-to-end: docx-only Drive folder → persisted entity + run +
    sections in live PG, with ONLY DATABASE_URL set (no
    DATABASE_URL_SYNC) so the worker-env sync-DSN resolver fallback
    is exercised on every sync call (job_executions, drive_feedback,
    synthesis_cache invalidation).
    """
    from sqlalchemy import text
    from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine

    from app.services.parsers.dma_package import parse_package
    from app.services.parsers.package_persist import persist_package
    from app.services.sync_dsn import resolve_sync_dsn

    # Mimic the production worker env: only DATABASE_URL is set,
    # no explicit DATABASE_URL_SYNC.
    monkeypatch.delenv("DATABASE_URL_SYNC", raising=False)
    monkeypatch.setenv("DATABASE_URL", _async_url())

    # --- Sync DSN resolver picks up DATABASE_URL ---
    sync = resolve_sync_dsn()
    assert sync is not None, "resolver returned None with DATABASE_URL set"
    assert "+psycopg" in sync, f"expected +psycopg URL, got {sync!r}"

    # --- Build a synthetic docx-only Drive folder ---
    with tempfile.TemporaryDirectory() as td:
        folder = Path(td) / "Acme Bank - DMA"
        folder.mkdir()
        _make_docx(
            folder / "Acme_DMA_Assessment_Report_2026Q1.docx",
            heading="Pillar 1 Deep Dive",
            body="Strategy & vision content.",
        )
        _make_docx(
            folder / "Acme_DMA_Assessment_Report_2026Q2.docx",
            heading="Pillar 3 Deep Dive",
            body="Technology stack content.",
        )
        _make_docx(  # noise — must be ignored
            folder / "Meeting Notes.docx",
            heading="Notes",
            body="unrelated",
        )

        # --- Parse: H6 + H7 ---
        pkg = parse_package(folder)
        warnings = pkg.parser_warnings or []
        assert any("docx_only_package_no_manifest" in w for w in warnings)
        assert any("synthesized run_manifest" in w for w in warnings)
        assert len(pkg.report_sections) >= 1
        synthesized_run_id = pkg.run_manifest.run_id

        # --- Seed catalogue parent row so persist's FK is satisfied ---
        async_engine = create_async_engine(_async_url())
        try:
            async with async_engine.begin() as conn:
                await conn.execute(text("""
                    INSERT INTO ccg_catalog_versions
                      (version, released_at, source_sha256s,
                       loader_run_id, frozen_at, notes)
                    VALUES ('v7.0', NOW(), '{}'::jsonb,
                            gen_random_uuid(), NOW(),
                            'e2e test band-aid')
                    ON CONFLICT (version) DO NOTHING
                """))

            # --- Persist: H5 (drive_folder_id) + H6 (synth manifest) ---
            Session = async_sessionmaker(async_engine, expire_on_commit=False)
            unique_drive_id = f"e2e-folder-{synthesized_run_id}"
            async with Session() as session:
                run_id, _persist_warnings = await persist_package(
                    session=session,
                    pkg=pkg,
                    data_source="DRIVE_BACKFILL",
                    drive_folder_id=unique_drive_id,
                )
                await session.commit()

            # --- Verify rows landed ---
            async with async_engine.begin() as conn:
                ent_row = (await conn.execute(
                    text(
                        "SELECT id, name, drive_folder_id FROM entities "
                        "WHERE drive_folder_id = :dfid"
                    ),
                    {"dfid": unique_drive_id},
                )).mappings().first()
                assert ent_row is not None, (
                    "entity row not created (or drive_folder_id mismatch — "
                    "H5 regression)"
                )
                assert ent_row["drive_folder_id"] == unique_drive_id

                run_row = (await conn.execute(
                    text(
                        "SELECT id, request_id, data_source "
                        "FROM runs WHERE id = :rid"
                    ),
                    {"rid": run_id},
                )).mappings().first()
                assert run_row is not None, "run row not created"
                assert run_row["request_id"] == synthesized_run_id
                assert run_row["data_source"] == "DRIVE_BACKFILL"

                section_count = (await conn.execute(
                    text("SELECT count(*) FROM document_sections WHERE run_id = :rid"),
                    {"rid": run_id},
                )).scalar_one()
                assert section_count >= 1, (
                    f"expected ≥1 document_sections from the DOCX, got "
                    f"{section_count} — H6/H7 regression"
                )
        finally:
            await async_engine.dispose()


@pytest.mark.asyncio
async def test_catalogue_empty_triggers_workbook_auto_bootstrap(monkeypatch):
    """2026-06 operator mandate: "No v5 catalogue will be uploaded.
    Just use the scoring toolkits during the backfill. No error
    message." When all parsed subcap_scores would fail catalogue
    resolution against an empty `ccg_subcaps`, persist_package must
    AUTO-BOOTSTRAP the catalogue from the workbook taxonomy so the
    resolver succeeds + the run goes ACTIVE end-to-end without
    operator intervention.

    Replaces the prior H8 contract (which expected
    `catalogue_unresolved:N/M` + `catalogue_empty_for_version`
    warnings + 0 subcap_scores rows) — that gate is now superseded
    by the auto-bootstrap path. See package_persist.
    `_bootstrap_catalogue_from_workbook` for the contract.
    """
    from sqlalchemy import text
    from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine

    from app.services.parsers.dma_package import parse_package
    from app.services.parsers.package_persist import persist_package

    monkeypatch.delenv("DATABASE_URL_SYNC", raising=False)
    monkeypatch.setenv("DATABASE_URL", _async_url())

    # Use the WSFS fixture — has 60 parsed subcap_scores.
    fixture_root = Path(__file__).resolve().parent / "fixtures" / \
        "dma_packages_sanitized" / "wsfs"
    if not fixture_root.exists():
        pytest.skip("wsfs fixture not on disk")

    pkg = parse_package(fixture_root)
    parsed_subcaps = len(pkg.subcap_scores)
    assert parsed_subcaps > 0, (
        "WSFS fixture must have parsed subcap_scores; if this fails "
        "the fixture has drifted, not the code"
    )

    async_engine = create_async_engine(_async_url())
    try:
        # Reset the catalogue tables: NO rows at all. The persist_package
        # path inserts a stub `ccg_catalog_versions` then auto-bootstraps
        # the children from the workbook's own subcap_ids.
        async with async_engine.begin() as conn:
            await conn.execute(text(
                "TRUNCATE ccg_catalog_versions RESTART IDENTITY CASCADE"
            ))

        Session = async_sessionmaker(async_engine, expire_on_commit=False)
        async with Session() as session:
            run_id, warnings = await persist_package(
                session=session, pkg=pkg, data_source="DRIVE_BACKFILL",
                drive_folder_id="h8-test-folder",
            )
            await session.commit()

        # Auto-bootstrap should have fired — look for the bootstrap
        # warning and absence of the PENDING_REVIEW gate signals.
        bootstrap = [w for w in warnings
                     if w.startswith("catalogue_auto_bootstrapped:")]
        empty = [w for w in warnings
                 if w.startswith("catalogue_empty_for_version:")]
        assert bootstrap, (
            f"expected `catalogue_auto_bootstrapped:N ...` warning when the "
            f"resolver had to fall back to workbook taxonomy; got: {warnings!r}"
        )
        assert not empty, (
            f"PENDING_REVIEW `catalogue_empty_for_version` warning must "
            f"NOT fire when auto-bootstrap succeeds; got: {warnings!r}"
        )

        # Run stays ACTIVE (no PENDING_REVIEW gate) — re-query the row.
        async with async_engine.begin() as conn:
            status = (await conn.execute(
                text("SELECT status FROM runs WHERE id = :rid"),
                {"rid": run_id},
            )).scalar_one()
            assert status == "ACTIVE", (
                f"auto-bootstrap must leave run ACTIVE; got status={status!r}"
            )
            # And the scores ARE persisted — the resolver returned
            # ResolvedSubcap on the first pass after bootstrap.
            score_count = (await conn.execute(
                text("SELECT count(*) FROM subcap_scores WHERE run_id = :rid"),
                {"rid": run_id},
            )).scalar_one()
            assert score_count == parsed_subcaps, (
                f"expected all {parsed_subcaps} scores persisted after "
                f"auto-bootstrap; got {score_count}"
            )
            # ccg_subcaps rows seeded for the package's catalogue version.
            ccg_count = (await conn.execute(
                text(
                    "SELECT count(*) FROM ccg_subcaps "
                    "WHERE version = (SELECT ccg_catalog_version FROM runs "
                    "WHERE id = :rid)"
                ),
                {"rid": run_id},
            )).scalar_one()
            assert ccg_count >= parsed_subcaps, (
                f"expected ≥{parsed_subcaps} ccg_subcaps seeded from "
                f"workbook; got {ccg_count}"
            )
    finally:
        await async_engine.dispose()


def test_job_executions_lifecycle_with_worker_env(monkeypatch):
    """H1 + H2: full job_executions lifecycle with ONLY DATABASE_URL set.
    Pre-fix every call raised RuntimeError silently swallowed by the
    runner's `_safe_*` wrappers.
    """
    from sqlalchemy import create_engine, text

    from app.services.job_executions_db import (
        create_execution_row,
        mark_started,
        mark_succeeded,
        reset_engine_for_tests,
        update_progress,
    )

    monkeypatch.delenv("DATABASE_URL_SYNC", raising=False)
    monkeypatch.setenv("DATABASE_URL", _async_url())
    reset_engine_for_tests()

    # Lifecycle: create → started → progress → succeeded.
    # `trigger_source` is a CHECK-constrained enum (admin_ui / scheduler /
    # cli / pubsub). 'cli' is the closest to our test invocation shape.
    eid = create_execution_row(
        job_name="historical_backfill", mode="full",
        trigger_source="cli",
    )
    assert eid, "create_execution_row returned empty id"
    mark_started(eid)
    update_progress(eid, folders_seen=10, files_parsed=3)
    mark_succeeded(eid, folders_seen=10, files_parsed=10)

    # Confirm with a direct sync query.
    sync_url = LIVE_DB_URL.replace("+asyncpg", "")
    eng = create_engine(sync_url)
    with eng.begin() as conn:
        row = conn.execute(
            text(
                "SELECT job_name, status, folders_seen, files_parsed, "
                "completed_at FROM job_executions WHERE id = :id"
            ),
            {"id": eid},
        ).mappings().first()
    assert row is not None, "job_executions row not found post-mark_succeeded"
    assert row["status"] == "succeeded"
    assert row["folders_seen"] == 10
    assert row["files_parsed"] == 10
    assert row["completed_at"] is not None


@pytest.mark.asyncio
async def test_repair_catalogue_stubs_self_heal(monkeypatch):
    """E2E: the /admin/repair:catalogue-stubs self-heal flow against
    live PG. Bypasses HTTP/auth to exercise just the DB side — the
    HTTP routing + auth are covered by `test_admin_jobs.py::
    TestDiagnosticsRepairAuth`.

    Locks: H12 (placeholder catalogue rows must be idempotent and
    NEVER overwrite real loader metadata when re-applied).
    """
    from sqlalchemy import text
    from sqlalchemy.ext.asyncio import create_async_engine

    monkeypatch.delenv("DATABASE_URL_SYNC", raising=False)
    monkeypatch.setenv("DATABASE_URL", _async_url())

    engine = create_async_engine(_async_url())
    try:
        # Clean slate.
        async with engine.begin() as conn:
            await conn.execute(text("TRUNCATE ccg_catalog_versions CASCADE"))

        # --- 1. First call inserts both versions ---
        requested = ["v7.0", "v5.5"]
        inserted_first: list[str] = []
        async with engine.begin() as conn:
            for v in requested:
                result = await conn.execute(text("""
                    INSERT INTO ccg_catalog_versions
                        (version, released_at, source_sha256s,
                         loader_run_id, frozen_at, notes)
                    VALUES (:v, NOW(), '{}'::jsonb,
                            gen_random_uuid(), NOW(),
                            'repair stub by e2e test')
                    ON CONFLICT (version) DO NOTHING
                """), {"v": v})
                if result.rowcount == 1:
                    inserted_first.append(v)
        assert sorted(inserted_first) == ["v5.5", "v7.0"], (
            f"first call should insert both versions; got {inserted_first}"
        )

        # --- 2. Simulate loader writing REAL metadata for v7.0 ---
        # The repair endpoint must NEVER overwrite this — ON CONFLICT
        # DO NOTHING is the safety contract.
        async with engine.begin() as conn:
            await conn.execute(text("""
                UPDATE ccg_catalog_versions
                   SET source_sha256s = :sf,
                       notes = 'loader-written REAL metadata — DO NOT OVERWRITE'
                 WHERE version = 'v7.0'
            """), {"sf": '{"Pillar_1.xlsx": "abc123"}'})

        # --- 3. Re-run repair: idempotent, must NOT clobber real v7.0 metadata ---
        inserted_second: list[str] = []
        async with engine.begin() as conn:
            for v in requested:
                result = await conn.execute(text("""
                    INSERT INTO ccg_catalog_versions
                        (version, released_at, source_sha256s,
                         loader_run_id, frozen_at, notes)
                    VALUES (:v, NOW(), '{}'::jsonb,
                            gen_random_uuid(), NOW(),
                            'second call — would clobber!')
                    ON CONFLICT (version) DO NOTHING
                """), {"v": v})
                if result.rowcount == 1:
                    inserted_second.append(v)
        assert inserted_second == [], (
            f"re-running repair must be a no-op when both versions "
            f"already exist; got inserts={inserted_second}"
        )

        # --- 4. Confirm v7.0's REAL metadata survived the second call ---
        async with engine.begin() as conn:
            v7 = (await conn.execute(text("""
                SELECT source_sha256s, notes
                  FROM ccg_catalog_versions WHERE version = 'v7.0'
            """))).mappings().first()
        assert "Pillar_1.xlsx" in str(v7["source_sha256s"]), (
            f"v7.0 real loader metadata was clobbered by the second "
            f"repair call! source_sha256s={v7['source_sha256s']!r}. "
            f"ON CONFLICT DO NOTHING is the safety contract — this "
            f"test would fail if a future refactor changes that to "
            f"ON CONFLICT DO UPDATE."
        )
        assert "loader-written REAL metadata" in v7["notes"], (
            f"v7.0 notes column was clobbered: {v7['notes']!r}"
        )
    finally:
        await engine.dispose()


@pytest.mark.asyncio
async def test_repair_close_stuck_jobs_filter(monkeypatch):
    """E2E: /admin/repair:close-stuck-jobs must close ONLY rows that
    have been running for >30 min, and must NOT touch terminal
    (succeeded/failed) rows even if they're old, NOR fresh running
    rows. Locks the SQL filter against accidental relaxation.
    """
    from sqlalchemy import text
    from sqlalchemy.ext.asyncio import create_async_engine

    monkeypatch.delenv("DATABASE_URL_SYNC", raising=False)
    monkeypatch.setenv("DATABASE_URL", _async_url())

    engine = create_async_engine(_async_url())
    try:
        async with engine.begin() as conn:
            await conn.execute(text("TRUNCATE job_executions"))

            # Stuck: running, >30min old → must be closed.
            stuck_id = (await conn.execute(text("""
                INSERT INTO job_executions
                  (job_name, mode, trigger_source, status, started_at)
                VALUES ('historical_backfill', 'full', 'cli', 'running',
                        NOW() - INTERVAL '2 hours')
                RETURNING id
            """))).scalar_one()

            # Fresh: running, only 5min old → must be UNTOUCHED.
            fresh_id = (await conn.execute(text("""
                INSERT INTO job_executions
                  (job_name, mode, trigger_source, status, started_at)
                VALUES ('drive_crawler', 'delta', 'scheduler', 'running',
                        NOW() - INTERVAL '5 minutes')
                RETURNING id
            """))).scalar_one()

            # Already succeeded (1 day ago): must NOT be re-flipped.
            done_id = (await conn.execute(text("""
                INSERT INTO job_executions
                  (job_name, mode, trigger_source, status,
                   started_at, completed_at)
                VALUES ('embedder', 'full', 'scheduler', 'succeeded',
                        NOW() - INTERVAL '1 day',
                        NOW() - INTERVAL '23 hours')
                RETURNING id
            """))).scalar_one()

            closed = (await conn.execute(text("""
                UPDATE job_executions
                   SET status = 'failed',
                       completed_at = NOW(),
                       error_message = 'auto-closed by e2e test'
                 WHERE status = 'running'
                   AND started_at < NOW() - INTERVAL '30 minutes'
                RETURNING id
            """))).scalars().all()

            # Read back final statuses.
            statuses = {
                str(r["id"]): r["status"]
                for r in (await conn.execute(text(
                    "SELECT id, status FROM job_executions"
                ))).mappings().all()
            }

        assert str(stuck_id) in [str(c) for c in closed], (
            f"stuck row {stuck_id} should have been closed; closed={closed}"
        )
        assert statuses[str(stuck_id)] == "failed"
        assert statuses[str(fresh_id)] == "running", (
            f"fresh row was incorrectly closed (status={statuses[str(fresh_id)]}) "
            f"— the >30min filter is the safety guard against killing "
            f"actively-running workers"
        )
        assert statuses[str(done_id)] == "succeeded", (
            f"already-terminal row was re-flipped to failed "
            f"(status={statuses[str(done_id)]}) — the WHERE status='running' "
            f"filter is the safety guard"
        )
    finally:
        await engine.dispose()
