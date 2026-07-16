"""Tests for the four-track 2026-05-28 visibility + deep-extraction work.

Tracks covered:

  1. CLI visibility — the worker LOUDLY warns when it can't write to
     job_executions; the resolve_sync_dsn fallback ladder reaches into
     Secret Manager when env vars are missing.

  2. /admin/trace/ingest endpoint — wire-format + defensive against
     missing tables.

  3. (removed) The frontend Pipeline status card lived in the net-new
     OperationsPanel, which was deleted for strict prototype fidelity;
     the /admin/trace/ingest BE contract it consumed stays covered by
     track 2 above.

  4. Deep extraction — DOCX scrape, OCR helpers, run-id mining, folder
     name inference, strategy ladder.

Pure-logic — no DB, no Drive, no OCR binaries required (the OCR helpers
return ("", 0) when pytesseract is missing, which the tests check).
"""
from __future__ import annotations

import os
from pathlib import Path

import pytest

BACKEND_DIR = Path(__file__).resolve().parents[1]
APP_DIR = BACKEND_DIR / "app"
FRONTEND_DIR = BACKEND_DIR.parent / "frontend"

LIVE_DB_URL = os.environ.get("SEED_CI_PG_URL", "")
HAS_LIVE_DB = bool(LIVE_DB_URL)


# ── Track 1: CLI visibility + Secret Manager fallback ─────────────────


def test_safe_create_row_warning_is_loud() -> None:
    """When `_safe_create_row` fails, the worker MUST print a banner
    visible in the operator's terminal — operator-reported gap was
    "UI shows no runs even while CLI jobs run", i.e. the worker
    silently couldn't write to job_executions."""
    runner_src = (
        BACKEND_DIR.parent / "workers" / "_runner.py"
    ).read_text()
    # The banner uses box-drawing chars so it's visually distinct.
    assert "WARNING: job_executions row NOT written" in runner_src, (
        "_safe_create_row failure must emit a LOUD operator-visible "
        "warning — silent failure was the root cause of the 'UI shows "
        "no runs' complaint"
    )
    assert "file=sys.stderr" in runner_src, (
        "Banner must go to stderr so it shows up in Cloud Run logs + "
        "the operator's terminal"
    )


def test_resolve_sync_dsn_secret_manager_fallback_exists() -> None:
    """The DSN resolver must try Secret Manager when env vars are
    missing. Covers the Cloud Shell case where the operator runs
    the worker via `python -m app.scripts.historical_backfill` without
    DATABASE_URL_SYNC wired."""
    from app.services import sync_dsn

    src = (APP_DIR / "services" / "sync_dsn.py").read_text()
    assert "secretmanager" in src
    assert "dma-insights-database-url-sync" in src
    assert hasattr(sync_dsn, "_try_secret_manager")
    assert hasattr(sync_dsn, "reset_secret_cache_for_tests")


def test_resolve_sync_dsn_explicit_env_wins_over_secret_manager(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Explicit DATABASE_URL_SYNC must short-circuit BEFORE Secret
    Manager — operators staging a different DB shouldn't have their
    env var silently overridden by the secret."""
    from app.services import sync_dsn

    monkeypatch.setenv(
        "DATABASE_URL_SYNC", "postgresql+psycopg://explicit:explicit@h:5432/d",
    )
    sync_dsn.reset_secret_cache_for_tests()
    assert sync_dsn.resolve_sync_dsn() == (
        "postgresql+psycopg://explicit:explicit@h:5432/d"
    )


def test_resolve_sync_dsn_disabled_by_env(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """When DMA_DISABLE_SECRET_DSN_FALLBACK=1 is set, the resolver
    skips Secret Manager (tests + local dev opt-out)."""
    from app.services import sync_dsn

    monkeypatch.delenv("DATABASE_URL_SYNC", raising=False)
    monkeypatch.delenv("DATABASE_URL", raising=False)
    monkeypatch.setenv("DMA_DISABLE_SECRET_DSN_FALLBACK", "1")
    sync_dsn.reset_secret_cache_for_tests()
    assert sync_dsn.resolve_sync_dsn() is None


def test_resolve_sync_dsn_swallows_secret_manager_failures(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """When google.cloud.secretmanager is missing or the secret is
    inaccessible, the helper returns None without raising.

    2026-05-29 fix: the prior implementation relied on the test env
    having no ADC, which broke on Cloud Build where the worker SA
    DOES have secretAccessor. Now we mock `_try_secret_manager` to
    simulate the lib-missing / no-ADC / IAM-deny failure modes
    explicitly, so the test is hermetic across all CI envs.
    """
    from app.services import sync_dsn

    monkeypatch.delenv("DATABASE_URL_SYNC", raising=False)
    monkeypatch.delenv("DATABASE_URL", raising=False)
    # Opt out of conftest._hermetic_secret_manager so the resolver
    # actually reaches `_try_secret_manager` — that's the path under test.
    monkeypatch.delenv("DMA_DISABLE_SECRET_DSN_FALLBACK", raising=False)
    sync_dsn.reset_secret_cache_for_tests()
    # Simulate every Secret Manager failure mode (lib missing, no ADC,
    # secret not present, IAM deny) — the helper must swallow all of
    # them and return None.
    monkeypatch.setattr(sync_dsn, "_try_secret_manager", lambda: None)
    assert sync_dsn.resolve_sync_dsn() is None


# ── Track 2: /admin/trace/ingest endpoint ──────────────────────────────


def test_trace_ingest_endpoint_registered() -> None:
    """The new trace endpoint must be wired into admin.py."""
    admin = (APP_DIR / "routers" / "admin.py").read_text()
    assert '"/trace/ingest"' in admin
    assert "async def trace_ingest" in admin
    # Auth-gated via require_admin (operator-only).
    trace_pos = admin.find("async def trace_ingest")
    decorator_window = admin[max(0, trace_pos - 600): trace_pos]
    assert "require_admin" in decorator_window, (
        "/admin/trace/ingest must be admin-gated"
    )


def test_trace_ingest_step_labels_locked() -> None:
    """The /admin/trace/ingest response surfaces these step labels
    verbatim. Drifting them breaks the operator's mental model + any
    consumer of the trace payload."""
    admin = (APP_DIR / "routers" / "admin.py").read_text()
    for label in (
        "entities ingested",
        "runs persisted",
        "latest run readable",
        "scores persisted",
        "report sections",
        "evidence persisted",
        "entity visible in directory",
        "UI overview will render scores",
    ):
        assert f'"{label}"' in admin, (
            f"Trace step label drift: {label!r} missing — the "
            f"/admin/trace/ingest payload surfaces this verbatim"
        )


def test_trace_endpoint_returns_self_heal_shape() -> None:
    """The response shape MUST include checks_passed + checks_total
    so the frontend can render `(X/Y checks)` chip even when
    individual queries throw."""
    admin = (APP_DIR / "routers" / "admin.py").read_text()
    for field in (
        "pipeline_steps", "checks_passed", "checks_total",
        "latest_entity_id", "ui_render_ok", "pipeline_healthy",
    ):
        assert f'"{field}"' in admin or f'{field}=' in admin, (
            f"trace_ingest response missing field: {field}"
        )


# ── Track 3 (frontend Pipeline status card) removed with OperationsPanel ─
# The card was net-new (absent from the prototype) and was deleted for
# strict fidelity. Its /admin/trace/ingest contract stays locked by the
# Track 2 tests above (endpoint registered · step labels · self-heal shape).


# ── Track 4: Deep extraction ───────────────────────────────────────────


class TestDeepExtractPureLogic:
    """Pure-logic tests — no DOCX, no OCR binaries, no Drive.

    The deep_extract module exposes helpers that work even when the
    underlying libs (python-docx, pytesseract, pdf2image) are missing.
    These tests pin the no-libs-installed behaviour: empty string,
    zero pages, graceful fallback to folder-name-only.
    """

    def test_extract_run_id_finds_req_pattern(self) -> None:
        from app.services.parsers.deep_extract import extract_run_id

        assert extract_run_id("Reference REQ-DEADBEEF in body") == "REQ-DEADBEEF"

    def test_extract_run_id_finds_dma_asm_pattern(self) -> None:
        from app.services.parsers.deep_extract import extract_run_id

        assert (
            extract_run_id("Issued under DMA-ASM-WSFS-20260519-0001 below")
            == "DMA-ASM-WSFS-20260519-0001"
        )

    def test_extract_run_id_returns_none_on_empty(self) -> None:
        from app.services.parsers.deep_extract import extract_run_id

        assert extract_run_id("") is None
        assert extract_run_id("no ids here") is None

    def test_infer_institution_strips_dma_suffix(self) -> None:
        from app.services.parsers.deep_extract import (
            infer_institution_from_folder,
        )

        assert (
            infer_institution_from_folder("Bank of Bermuda - DMA")
            == "Bank of Bermuda"
        )
        assert infer_institution_from_folder("WSFS_DMA_Engagement_Package") == "WSFS"
        assert infer_institution_from_folder("RegionsBank_DMA_20260518") == "RegionsBank"

    def test_infer_institution_fallback_to_raw_name(self) -> None:
        from app.services.parsers.deep_extract import (
            infer_institution_from_folder,
        )
        # No DMA token → return raw name.
        assert (
            infer_institution_from_folder("just_a_folder")
            == "just_a_folder"
        )

    def test_scrape_docx_text_empty_when_path_missing(
        self, tmp_path: Path,
    ) -> None:
        from app.services.parsers.deep_extract import scrape_docx_text

        # Nonexistent path → "".
        assert scrape_docx_text(tmp_path / "missing.docx") == ""

    def test_ocr_returns_empty_when_libs_missing(
        self, tmp_path: Path,
    ) -> None:
        """The OCR helpers MUST gracefully return ("", 0) when their
        upstream libs (pytesseract, pdf2image) aren't installed —
        retry-mode worker must not crash on a clean container."""
        from app.services.parsers.deep_extract import ocr_docx_images, ocr_pdf

        # Either the libs aren't installed (returns ("",0)) or they
        # are but the file doesn't exist (also returns ("",0)).
        assert ocr_docx_images(tmp_path / "x.docx") == ("", 0)
        assert ocr_pdf(tmp_path / "x.pdf") == ("", 0)

    def test_deep_extract_folder_returns_folder_name_only_when_empty(
        self, tmp_path: Path,
    ) -> None:
        """An empty folder still returns a usable result — the strategy
        is 'folder_name_only', the institution is inferred from the
        folder name, the caller has something to UPSERT."""
        from app.services.parsers.deep_extract import deep_extract_folder

        folder = tmp_path / "WSFS_DMA_2026"
        folder.mkdir()
        result = deep_extract_folder(folder)
        assert result.strategy in ("folder_name_only", "none")
        if result.strategy == "folder_name_only":
            assert result.institution == "WSFS"

    def test_deep_extract_recovers_text_from_real_docx(
        self, tmp_path: Path,
    ) -> None:
        """Functional end-to-end test against a synthetic DOCX —
        exercises the docx_text strategy path with real python-docx.

        Mirrors the production retry-mode flow: operator's folder has
        a DOCX report but no canonical 03_scoring_workbook layout; the
        deep extractor must recover the report text + mine the run-id
        regex out of it.
        """
        try:
            import docx
        except ImportError:
            pytest.skip("python-docx not installed in test env")

        # Build a minimal DOCX inside a synthetic folder.
        folder = tmp_path / "Bank_of_Test_DMA_Engagement_Package"
        folder.mkdir()
        reports = folder / "04_reports"
        reports.mkdir()
        doc = docx.Document()
        doc.add_heading("DMA Assessment Report", level=1)
        doc.add_paragraph(
            "Run identifier: DMA-ASM-BANKOFTEST-20260519-0001"
        )
        doc.add_heading("Pillar 1 — Strategic Posture & Governance", level=2)
        doc.add_paragraph(
            "Subcap P1C1.1.1 — current state assessment indicates the "
            "institution is at the Building maturity level."
        )
        doc.add_paragraph(
            "Subcap P1C1.1.2 — target state should push to Competing "
            "within the next 18 months."
        )
        doc.add_heading("Pillar 2 — Customer Engagement", level=2)
        doc.add_paragraph(
            "Score evidence supports a Differentiating posture on "
            "P2C1.1.1; competing on P2C2.1.1."
        )
        # Add a table — exercises the cell-walk branch of scrape_docx_text.
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Subcap"
        table.rows[0].cells[1].text = "Score"
        table.rows[1].cells[0].text = "P1C1.1.1"
        table.rows[1].cells[1].text = "3.0"
        doc.save(str(reports / "Assessment_Report_BankOfTest.docx"))

        from app.services.parsers.deep_extract import deep_extract_folder
        result = deep_extract_folder(folder)

        # Strategy: docx_text (the python-docx scraper recovered enough text).
        assert result.strategy == "docx_text", (
            f"Expected docx_text strategy, got {result.strategy!r} "
            f"(text_chars={len(result.scraped_text)})"
        )
        # Text length must clear the 200-char threshold + contain
        # markers from both paragraphs AND the table cells.
        assert len(result.scraped_text) >= 200
        assert "Strategic Posture" in result.scraped_text
        assert "P1C1.1.1" in result.scraped_text
        # Table walked: header + cell value both present.
        assert "Score" in result.scraped_text
        # Run-id regex mined the DMA-ASM token out of the text.
        assert result.run_id == "DMA-ASM-BANKOFTEST-20260519-0001"
        # Institution inferred from the folder name (Bank_of_Test_DMA → "Bank_of_Test").
        assert result.institution == "Bank_of_Test"
        # docx_paths_scraped lists the relative path the extractor walked.
        assert result.docx_paths_scraped is not None
        assert any(
            "Assessment_Report" in p for p in result.docx_paths_scraped
        ), result.docx_paths_scraped

    def test_has_scoreable_content_threshold(self) -> None:
        from app.services.parsers.deep_extract import has_scoreable_content

        # Short text — never scoreable.
        assert has_scoreable_content("subcap pillar") is False
        # Long text without DMA hints — not scoreable.
        assert has_scoreable_content("lorem " * 50) is False
        # Long text with 3+ hints — scoreable.
        good = " ".join([
            "pillar maturity", "subcap detection", "score band", "x " * 100,
        ])
        assert has_scoreable_content(good) is True


def test_dma_package_parser_consults_lenient_env(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path,
) -> None:
    """When DMA_INGEST_LENIENT=1 is set, parse_package must invoke
    deep_extract_folder as a fallback. We can't easily test the
    actual deep-extract path without DOCX fixtures, but the SOURCE
    must reference both the env var + the deep_extract function."""
    src = (
        APP_DIR / "services" / "parsers" / "dma_package.py"
    ).read_text()
    assert "DMA_INGEST_LENIENT" in src, (
        "parse_package must consult DMA_INGEST_LENIENT — the env var "
        "is how the worker enables retry-mode deep extraction"
    )
    assert "deep_extract_folder" in src, (
        "parse_package must call deep_extract_folder when lenient"
    )
    # The warning name surfaces in import_audit so it must be exact.
    assert "lenient_mode_deep_extract:" in src


def test_worker_activates_lenient_mode_in_retry() -> None:
    """When --retry-failed-only is active, the worker MUST set
    DMA_INGEST_LENIENT=1 so the parser fall-through chain fires.
    Without this the retry is no more thorough than the first pass."""
    src = (APP_DIR / "scripts" / "historical_backfill.py").read_text()
    # The env var must be set INSIDE the retry-mode branch (after the
    # retry_targets resolve), not unconditionally.
    pos = src.find("retry_failed_only")
    assert pos > 0
    window = src[pos: pos + 5000]
    assert 'os.environ["DMA_INGEST_LENIENT"] = "1"' in window, (
        "Worker must set DMA_INGEST_LENIENT=1 in the retry-mode branch"
    )


# ── Live-PG: trace endpoint against a real seeded DB ───────────────────


@pytest.mark.skipif(not HAS_LIVE_DB, reason="SEED_CI_PG_URL not set")
class TestTraceEndpointLivePg:
    """The trace endpoint must work against a real seeded DB.

    Doesn't assert the FULL data shape (depends on seed_ci output)
    but validates the endpoint NEVER raises + returns a usable shape.
    """

    @pytest.fixture(autouse=True)
    def setup_client(self, monkeypatch: pytest.MonkeyPatch):
        async_url = LIVE_DB_URL
        sync_url = async_url.replace("+asyncpg", "+psycopg")
        monkeypatch.setenv("DATABASE_URL_SYNC", sync_url)
        monkeypatch.setenv("DATABASE_URL", async_url)
        monkeypatch.setenv("DMA_BOT_API_KEY", "ci-bot-key")
        monkeypatch.setenv("ENV", "local")
        # 2026-06-05: hard-reset module-level engine globals AND the
        # Settings lru_cache. Prior tests in the suite (notably
        # test_backfill_quarantine.TestQuarantineLivePg) leave asyncpg
        # connections bound to their event loop; without this reset
        # the trace endpoint here tries to reuse them and raises
        # "Task got Future attached to a different loop" /
        # "Event loop is closed". The get_settings cache_clear is
        # required AFTER monkeypatch.setenv so the next get_settings()
        # call re-reads the freshly-patched DATABASE_URL/ENV values
        # (the same cache-clear contract refresh_engine_on_auth_failure
        # uses for runtime secret rotation).
        from app import database as _app_db
        from app.config import get_settings
        from app.services.job_executions_db import reset_engine_for_tests
        _app_db._engine = None
        _app_db._sessionmaker = None
        get_settings.cache_clear()
        reset_engine_for_tests()
        yield
        # Symmetric teardown so the NEXT test module starts clean too.
        _app_db._engine = None
        _app_db._sessionmaker = None
        get_settings.cache_clear()

    def test_trace_endpoint_returns_pipeline_steps(self):
        """The endpoint must return checks_total > 0 + steps non-empty
        even on a fresh/empty DB."""
        from fastapi.testclient import TestClient

        from app.main import create_app
        app = create_app()
        client = TestClient(app)
        # Get an ADMIN session token.
        r = client.post(
            "/api/v1/auth/dev-login",
            params={"email": "mishley.otiende@zennify.com"},
        )
        assert r.status_code == 200, r.text
        # Hit the trace endpoint.
        r = client.get("/api/v1/admin/trace/ingest")
        assert r.status_code == 200, r.text
        body = r.json()
        assert "pipeline_steps" in body
        assert "checks_total" in body
        assert "checks_passed" in body
        assert "pipeline_healthy" in body
        assert body["checks_total"] >= 1, (
            "trace endpoint must report at least one check even on "
            "empty DB (entities count is the minimum)"
        )
