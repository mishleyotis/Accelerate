"""F2 tests — parser handles every package shape we've actually seen.

These tests synthesize *minimal* package skeletons matching the layouts
of the 5 operator-uploaded packages (RegionsBank, Amalgamated, ANB,
WSFS, AmeriCU). Full real fixtures are committed in
`backend/tests/fixtures/dma_packages/` once sanitization completes
(plan F1); these tests give the parser a green signal in the
meantime and cover the structural quirks each shape exercises:

  - RegionsBank: flat layout (no top-level wrapper dir), no
    MANIFEST.json, no run_manifest.json — must synthesize from the
    governance_qa_verdict file (priority-4 glob).
  - Amalgamated: nested wrapper, canonical 07_governance/run_manifest.json.
  - ANB: nested wrapper, 08_appendices/run_manifest.json, layer1_/layer2_
    governance pattern.
  - WSFS: flat layout, 08_appendices/run_manifest.json with l1_run_id
    field (the WSFS schema variant).
  - AmeriCU: nested wrapper, 03_scoring_workbook/run_manifest.json
    (the non-canonical location the previous parser couldn't find).

State branches covered:
  manifest_in_root              | manifest_in_07_governance
  manifest_in_08_appendices     | manifest_in_03_scoring_workbook   ← AmeriCU fix
  qa_verdict_as_run_manifest    | wrapper_dir_single_child
  flat_layout_no_wrapper        | both_layouts_threshold_2
"""
from __future__ import annotations

import json
from pathlib import Path

import pytest

from app.services.parsers.dma_package import _find_root, parse_package


def _seed_minimal_package(
    root: Path,
    *,
    manifest_path: str = "07_governance/run_manifest.json",
    manifest_schema: str = "run_manifest_v2",
    run_id_field: str = "run_id",
    run_id_value: str = "DMA-ASM-TESTBANK-20260101-0001",
    extra_dirs: tuple[str, ...] = (
        "01_evidence", "03_scoring_workbook",
        "06_peers", "07_governance",
    ),
) -> None:
    """Drop a minimum-viable DMA package skeleton under `root`."""
    for d in extra_dirs:
        (root / d).mkdir(parents=True, exist_ok=True)
    manifest = root / manifest_path
    manifest.parent.mkdir(parents=True, exist_ok=True)
    manifest.write_text(json.dumps({
        "$schema": manifest_schema,
        run_id_field: run_id_value,
        "entity": "Test Bank",
        "overall_score": 2.0,
    }))


@pytest.fixture
def tmpdir(tmp_path: Path) -> Path:
    return tmp_path


def test_regionsbank_shape_flat_no_manifest_qa_verdict_as_manifest(tmpdir):
    """RegionsBank: flat layout, no MANIFEST.json, no run_manifest.json,
    governance_qa_verdict_*.json is the only manifest-ish file. Parser
    priority-4 glob `*qa_verdict*.json` should pick it up."""
    root = tmpdir / "regions"
    root.mkdir()
    for d in ("01_evidence", "03_scoring_workbook", "06_peers", "07_governance"):
        (root / d).mkdir()
    (root / "07_governance" / "governance_qa_verdict_RegionsBank_20260518.json").write_text(
        json.dumps({
            "$schema": "run_manifest_v2",
            "run_id": "DMA-ASM-REGIONS-20260518-0001",
            "verdict": "PASS_WITH_NOTES",
        }),
    )
    detected = _find_root(root)
    assert detected == root
    pkg = parse_package(root)
    assert pkg.run_manifest.run_id == "DMA-ASM-REGIONS-20260518-0001"


def test_amalgamated_shape_nested_wrapper_canonical_07_governance_manifest(tmpdir):
    """Amalgamated_Bank_DMA_2026/ wrapper, canonical
    07_governance/run_manifest.json. Parser priority-1 path."""
    root = tmpdir / "wrapper_outer"
    root.mkdir()
    inner = root / "Amalgamated_Bank_DMA_2026"
    inner.mkdir()
    _seed_minimal_package(
        inner,
        manifest_path="07_governance/run_manifest.json",
        run_id_value="DMA-ASM-AMALGAMATED-20260429-0001",
    )
    detected = _find_root(root)
    assert detected == inner
    pkg = parse_package(root)
    assert pkg.run_manifest.run_id == "DMA-ASM-AMALGAMATED-20260429-0001"


def test_anb_shape_nested_wrapper_08_appendices_manifest(tmpdir):
    """ANB_DMA_Complete_Bundle/ wrapper, 08_appendices/run_manifest.json.
    Parser priority-1 path."""
    root = tmpdir / "wrapper_outer"
    root.mkdir()
    inner = root / "ANB_DMA_Complete_Bundle"
    inner.mkdir()
    _seed_minimal_package(
        inner,
        manifest_path="08_appendices/run_manifest.json",
        run_id_value="DMA-ASM-ANB-20260421-0001",
        extra_dirs=("01_evidence", "03_scoring_workbook",
                    "06_peers", "07_governance", "08_appendices"),
    )
    detected = _find_root(root)
    assert detected == inner
    pkg = parse_package(root)
    assert pkg.run_manifest.run_id == "DMA-ASM-ANB-20260421-0001"


def test_wsfs_shape_flat_08_appendices_manifest_l1_run_id(tmpdir):
    """WSFS: flat layout, 08_appendices/run_manifest.json with the
    WSFS schema variant (l1_run_id instead of run_id)."""
    root = tmpdir / "wsfs"
    root.mkdir()
    for d in ("01_evidence", "03_scoring_workbook", "06_peers",
              "07_governance", "08_appendices"):
        (root / d).mkdir()
    (root / "08_appendices" / "run_manifest.json").write_text(json.dumps({
        "$schema": "run_manifest_v2",
        "assessment_id": "DMA-RES-WSFS-20260519-0001",
        "l1_run_id": "DMA-ASM-WSFS-20260519-0001",
        "entity": "WSFS Financial Corporation",
        "overall_score": 2.07,
    }))
    detected = _find_root(root)
    assert detected == root
    pkg = parse_package(root)
    assert pkg.run_manifest.run_id == "DMA-ASM-WSFS-20260519-0001"


def test_americu_shape_manifest_in_03_scoring_workbook(tmpdir):
    """AmeriCU bug fix — run_manifest.json lives in 03_scoring_workbook/.
    Previously parser priority-1/2 only checked
    `.`, `07_governance`, `08_appendices` — AmeriCU's manifest was
    invisible and the parser fell through to the qa_verdict fallback,
    losing scoring fields. F2 added `03_scoring_workbook` (and
    `02_research_workbook`) to the canonical lookup."""
    root = tmpdir / "wrapper_outer"
    root.mkdir()
    inner = root / "AmeriCU_DMA_Deliverable_2026-04-29"
    inner.mkdir()
    for d in ("01_evidence", "03_scoring_workbook", "06_peers", "07_governance"):
        (inner / d).mkdir()
    # Non-canonical schema: AmeriCU uses `schema` (no $) and
    # `assessment_id` not `run_id`. parse_run_manifest already accepts
    # both per its existing implementation.
    (inner / "03_scoring_workbook" / "run_manifest.json").write_text(json.dumps({
        "schema": "dma_assessment_run_manifest_v2",
        "assessment_id": "DMA-RES-AMERICU-20260427-0001",
        "entity": "AmeriCU Credit Union",
        "overall_score": 2.44,
    }))
    detected = _find_root(root)
    assert detected == inner
    pkg = parse_package(root)
    assert pkg.run_manifest.run_id == "DMA-RES-AMERICU-20260427-0001"
    # No warning about "used variant manifest file" because
    # 03_scoring_workbook/run_manifest.json IS now in the canonical
    # priority-1 list (after F2).
    assert not any("variant manifest file" in w for w in pkg.parser_warnings or [])


def test_threshold_relaxed_to_2_canonical_subfolders(tmpdir):
    """The _find_root threshold was relaxed from ≥3 to ≥2 of the
    canonical subfolders, AND requires at least one manifest-bearing
    kind. Verifies a sparse 2-folder package is accepted."""
    root = tmpdir / "sparse"
    root.mkdir()
    (root / "01_evidence").mkdir()
    (root / "07_governance").mkdir()    # manifest-bearing
    (root / "07_governance" / "run_manifest.json").write_text(json.dumps({
        "$schema": "run_manifest_v2",
        "run_id": "DMA-ASM-SPARSE-20260601-0001",
        "entity": "Sparse Bank",
    }))
    detected = _find_root(root)
    assert detected == root


def test_threshold_rejects_unmanifested_two_folder_layout(tmpdir):
    """Even with 2 numbered subfolders, _find_root rejects if neither
    is manifest-bearing (07_governance / 08_appendices / 03_scoring).
    Guards against accidentally parsing arbitrary trees."""
    root = tmpdir / "bogus"
    root.mkdir()
    (root / "01_evidence").mkdir()
    (root / "05_narrative_deck").mkdir()
    with pytest.raises(FileNotFoundError):
        _find_root(root)


# ── 2026-05-28 H6/H7 hotfix regressions ───────────────────────────────


def _make_minimal_docx(path: Path, heading: str = "Executive Summary",
                       body: str = "This is the body.") -> None:
    """Write a real .docx so `parse_assessment_report` can open it.

    We use python-docx (already a runtime dep — see
    `app/services/parsers/assessment_report.py::extract_paragraphs_from_docx`).
    """
    import docx
    d = docx.Document()
    d.add_heading(heading, level=1)
    d.add_paragraph(body)
    d.save(str(path))


def test_docx_only_folder_accepted_no_manifest_no_subfolders(tmpdir):
    """H6 — a Drive folder with ONLY an Assessment Report DOCX (no
    MANIFEST.json, no canonical 01_..08_ subfolders) must be accepted
    by `_find_root`. Pre-fix this raised FileNotFoundError, which
    caused 21 of 115 Drive folders to be miscounted as parse failures
    in the 2026-05-28 backfill.
    """
    root = tmpdir / "docx_only"
    root.mkdir()
    _make_minimal_docx(
        root / "DMA_Assessment_Report_Acme_Bank_2026.docx",
        heading="Pillar 1 Deep Dive",
        body="Strategy overview content.",
    )
    detected = _find_root(root)
    assert detected == root, (
        f"expected _find_root to accept docx-only folder; got {detected!r}"
    )


def test_docx_only_folder_emits_docx_only_warning(tmpdir):
    """H6 — `parse_package` must surface `docx_only_package_no_manifest`
    in parser_warnings when the folder was accepted via the docx-only
    branch. This lets import audit + admin UI distinguish docx-only
    ingest from canonical full-package ingest.
    """
    root = tmpdir / "docx_only_with_warn"
    root.mkdir()
    _make_minimal_docx(
        root / "Some_DMA_Assessment_Report.docx",
        heading="Executive Summary",
        body="Body for assessment.",
    )
    pkg = parse_package(root)
    warnings = pkg.parser_warnings or []
    assert any("docx_only_package_no_manifest" in w for w in warnings), (
        f"expected docx_only_package_no_manifest warning, got: {warnings!r}"
    )
    # And the report_sections should be populated since the DOCX
    # parsed successfully.
    assert len(pkg.report_sections) >= 1, (
        f"expected at least one report section from the DOCX; got "
        f"{len(pkg.report_sections)}"
    )


def test_docx_only_nested_in_reports_subdir(tmpdir):
    """H6 + H7 — DOCXs in a `Reports/` (non-canonical) subdirectory
    must be discovered via the recursive find_assessment_reports
    branch. Several Drive folders in production nest reports this way.
    """
    root = tmpdir / "nested"
    root.mkdir()
    (root / "Reports").mkdir()
    _make_minimal_docx(
        root / "Reports" / "DMA_Assessment_Report_Nested.docx",
        heading="Pillar 2 Deep Dive",
        body="Engagement story.",
    )
    detected = _find_root(root)
    assert detected == root
    pkg = parse_package(root)
    assert len(pkg.report_sections) >= 1, (
        f"expected report sections from nested DOCX; got {len(pkg.report_sections)}"
    )


def test_multiple_dma_docx_files_all_parsed(tmpdir):
    """H7 — when a folder has multiple Assessment Report DOCXs (e.g.
    initial + refresh), all of them must be parsed and their sections
    persisted (deduped by kind+heading). Pre-fix only the first was
    read.
    """
    root = tmpdir / "multi_docx"
    root.mkdir()
    _make_minimal_docx(
        root / "DMA_Assessment_Report_2025Q4.docx",
        heading="Pillar 1 Deep Dive",
        body="Q4 strategy content.",
    )
    _make_minimal_docx(
        root / "DMA_Assessment_Report_2026Q1.docx",
        heading="Pillar 3 Deep Dive",
        body="Q1 technology content.",
    )
    pkg = parse_package(root)
    # Both should contribute distinct sections (different kinds).
    kinds = {s.kind for s in pkg.report_sections}
    # We expect at least 2 distinct kinds since the two DOCXs cover
    # different pillars. The classifier may map "Pillar 1 Deep Dive"
    # to p1_deep_dive and "Pillar 3 Deep Dive" to p3_deep_dive — or
    # both to a generic kind; either way, having content from BOTH
    # files is what we're proving.
    assert len(pkg.report_sections) >= 2, (
        f"expected ≥2 sections from two DOCXs; got {len(pkg.report_sections)}. "
        f"kinds={kinds!r}"
    )


def test_unrelated_docx_folder_still_rejected(tmpdir):
    """H6 — DOCX-only acceptance must NOT cause unrelated Drive
    folders to be miscategorised as DMA packages. Only DOCXs whose
    filename matches the DMA report tokens count.
    """
    root = tmpdir / "unrelated"
    root.mkdir()
    # Plain "Meeting Notes.docx" must NOT trigger DOCX acceptance.
    _make_minimal_docx(
        root / "Meeting Notes 2026-05.docx",
        heading="Notes",
        body="Random.",
    )
    with pytest.raises(FileNotFoundError):
        _find_root(root)


def test_progress_flush_called_during_backfill(monkeypatch):
    """H2 — historical_backfill must pass `flush=True` to the
    execution tracker so counters land in the DB during the run, not
    just at completion. The pre-fix code skipped the DB write and
    showed `running` with all-NULL counters in the admin UI.
    """
    import inspect

    from app.scripts import historical_backfill as hb

    src = inspect.getsource(hb)
    # Every `_ex.update(` call should have `flush=True` somewhere
    # within its parenthesised arg block. Crude but catches the
    # exact regression we just fixed.
    # Strategy: find each `_ex.update(` and walk forward until the
    # matching close paren, checking for the keyword.
    i = 0
    bad: list[str] = []
    while True:
        i = src.find("_ex.update(", i)
        if i == -1:
            break
        depth = 0
        j = i
        while j < len(src):
            ch = src[j]
            if ch == "(":
                depth += 1
            elif ch == ")":
                depth -= 1
                if depth == 0:
                    break
            j += 1
        call_text = src[i:j + 1]
        if "flush=True" not in call_text:
            bad.append(call_text[:120].replace("\n", " "))
        i = j + 1
    assert not bad, (
        "found _ex.update(...) calls without flush=True — admin UI "
        "progress counters will not persist mid-run:\n  "
        + "\n  ".join(bad)
    )
