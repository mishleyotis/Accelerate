"""Deep-extraction helpers — used by retry-mode backfill to recover
data from packages the first-pass parser couldn't process.

When `--retry-failed-only` is active, the historical_backfill worker
sets env `DMA_INGEST_LENIENT=1` and parse_package consults these
helpers as a last-resort fallback chain BEFORE giving up on a folder:

    1. Scrape every `*.docx` in the folder tree (depth ≤3) for text
       — even when the canonical 03_scoring_workbook layout is missing.
    2. If text extraction returns < 200 chars, attempt OCR via
       pytesseract on rendered PDF pages (graceful fallback when the
       binary isn't installed).
    3. Synthesize a minimal RunManifest from scraped text + folder name
       so persist_package has something to anchor to.

The goal: "Retries should focus on deep information retrieval from
obtained reports and more robust parsing, not just trying drive access.
Even the Retries can try visual scans if stuff fails." (operator).

Self-heal contract:
  - Every helper returns a best-effort result; failures degrade quietly.
  - Missing dependencies (pytesseract / poppler / docx2txt) are
    detected at import time and the helper returns None.
  - All helpers are pure (no DB writes, no Drive calls) so they can be
    unit-tested without infra.
"""
from __future__ import annotations

import re
from collections.abc import Iterable
from dataclasses import dataclass
from pathlib import Path

# Conservative DMA-token regex — matches "DMA", "DMA-ASM", "REQ-XXXXXXXX",
# "DMA-ASM-{ENTITY}-{YYYYMMDD}-{NNNN}" in scraped DOCX text. Operators
# put these into folder names + filenames; we mine them as a last resort.
_RUN_ID_RE = re.compile(
    r"(REQ-[0-9A-F]{8}|DMA-ASM-[A-Z0-9_-]+-\d{8}-\d{4})",
    re.IGNORECASE,
)

# Heading hints that signal a deep DOCX contains scoreable content
# even when the canonical workbook is missing. Used to gate the
# "extract scores via heading-walk" path.
_SCORE_HEADING_HINTS = (
    "subcap", "capability", "pillar", "maturity", "score",
    "current state", "current-state", "currentstate",
    "target state", "target-state", "competing", "differentiating",
)


@dataclass
class DeepExtractResult:
    """Output of a deep-extraction attempt over one folder.

    Fields are deliberately all optional — the persistence layer
    checks each one independently so a partial recovery (e.g. only
    text + run_id, no scores) is still a useful enrichment over the
    `failed_parse` baseline.
    """
    scraped_text: str = ""
    run_id: str | None = None
    institution: str | None = None
    ocr_pages: int = 0
    docx_paths_scraped: list[str] | None = None
    strategy: str = "none"           # one of: docx_text / docx_ocr /
                                     # pdf_ocr / folder_name_only / none
    warnings: list[str] | None = None

    def __post_init__(self) -> None:
        if self.docx_paths_scraped is None:
            self.docx_paths_scraped = []
        if self.warnings is None:
            self.warnings = []


def scrape_docx_text(path: Path) -> str:
    """Extract every paragraph + table cell from a DOCX as plain text.

    Returns empty string on any failure (file unreadable, python-docx
    missing, etc.). Safe to call on non-DOCX paths — returns "".

    The default `python-docx` extraction misses content inside images;
    OCR fallback lives in `ocr_docx_images` below.
    """
    try:
        import docx  # type: ignore[import-untyped]
    except ImportError:
        return ""
    try:
        document = docx.Document(str(path))
    except Exception:
        return ""
    parts: list[str] = []
    for p in document.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def ocr_docx_images(path: Path) -> tuple[str, int]:
    """OCR every image embedded in the DOCX.

    Returns (extracted_text, page_count). Returns ("", 0) on any
    failure — pytesseract missing, tesseract binary not on PATH,
    no images in the file, etc.
    """
    try:
        from io import BytesIO

        import docx  # type: ignore[import-untyped]
        import pytesseract  # type: ignore[import-untyped]
        from PIL import Image  # type: ignore[import-untyped]
    except ImportError:
        return ("", 0)
    try:
        document = docx.Document(str(path))
    except Exception:
        return ("", 0)
    pieces: list[str] = []
    count = 0
    # python-docx exposes embedded images via the related parts.
    try:
        for rel in document.part.rels.values():
            if "image" not in rel.target_ref:
                continue
            try:
                blob = rel.target_part.blob
                img = Image.open(BytesIO(blob))
                txt = pytesseract.image_to_string(img) or ""
                if txt.strip():
                    pieces.append(txt.strip())
                count += 1
            except Exception:
                continue
    except Exception:
        return ("", count)
    return ("\n".join(pieces), count)


def ocr_pdf(path: Path) -> tuple[str, int]:
    """OCR a PDF by rasterizing each page + running tesseract.

    Returns (extracted_text, page_count). Failures (pdf2image missing,
    poppler not installed, pytesseract missing, file corrupt) all
    return ("", 0).
    """
    try:
        import pytesseract  # type: ignore[import-untyped]
        from pdf2image import convert_from_path  # type: ignore[import-untyped]
    except ImportError:
        return ("", 0)
    try:
        pages = convert_from_path(str(path), dpi=200)
    except Exception:
        return ("", 0)
    pieces: list[str] = []
    count = 0
    for page in pages:
        try:
            txt = pytesseract.image_to_string(page) or ""
            if txt.strip():
                pieces.append(txt.strip())
            count += 1
        except Exception:
            continue
    return ("\n".join(pieces), count)


def _walk_files(root: Path, suffix: str, max_depth: int = 3) -> Iterable[Path]:
    """Yield every file under `root` (depth ≤ max_depth) ending in
    `suffix` (case-insensitive). Resilient to permission errors."""
    root = root.resolve()
    suffix = suffix.lower()
    try:
        # rglob doesn't accept depth, but we can cap via parts comparison.
        for p in root.rglob(f"*{suffix}"):
            try:
                depth = len(p.relative_to(root).parts)
            except ValueError:
                continue
            if depth <= max_depth and p.is_file():
                yield p
    except Exception:
        return


def extract_run_id(text: str) -> str | None:
    """Pull the first DMA-ASM or REQ run-id out of `text`. None if no
    match. Pure regex extraction — feeds the synthesized RunManifest.
    """
    if not text:
        return None
    m = _RUN_ID_RE.search(text)
    return m.group(1).upper() if m else None


def infer_institution_from_folder(folder_name: str) -> str:
    """Best-effort institution name from a Drive folder name.

    Examples:
      'Bank of Bermuda - DMA'                  → 'Bank of Bermuda'
      'WSFS_DMA_Engagement_Package'             → 'WSFS'
      'RegionsBank_DMA_20260518'                → 'RegionsBank'
      'AmeriCU_DMA_Deliverable_2026-04-29'      → 'AmeriCU'

    Falls back to the raw folder name if no DMA token is present.
    """
    if not folder_name:
        return ""
    # Strip everything from the first DMA-like token onward. Word
    # boundary doesn't work here because `_` is a word char in Python
    # regex (so `\bDMA\b` matches inside `WSFS_DMA_…` only because the
    # leading `_` IS a word boundary, but `DMA_` has no boundary after
    # the A). Explicit separator class handles the common cases:
    # space/hyphen/underscore separator both before AND after the DMA
    # token, or end of string.
    parts = re.split(
        r"[\s_\-]+(?i:dma)(?:[\s_\-]+|$)", folder_name, maxsplit=1
    )
    return parts[0].strip(" _-") if parts else folder_name


def deep_extract_folder(folder: Path) -> DeepExtractResult:
    """Run the FULL deep-extraction chain over `folder`.

    Strategy ladder (returns at first material success):
      1. Walk every .docx (depth ≤3) + extract text via python-docx.
      2. If total scraped text < 200 chars, OCR every embedded image
         in those DOCX files.
      3. If still < 200 chars, walk every .pdf + OCR each page.
      4. If still < 200 chars, derive institution from folder name
         and return a folder_name_only result (caller has at least
         a hint to UPSERT).

    The result's `strategy` field tells the operator + the quarantine
    row which path actually succeeded — invaluable for diagnosing
    "why did this folder fail" without hand-running the parser.
    """
    result = DeepExtractResult()

    # Step 1: scrape DOCX text.
    docx_paths = list(_walk_files(folder, ".docx", max_depth=3))
    scraped: list[str] = []
    for p in docx_paths:
        txt = scrape_docx_text(p)
        if txt:
            scraped.append(txt)
            result.docx_paths_scraped.append(str(p.relative_to(folder)))
    text = "\n".join(scraped).strip()
    if len(text) >= 200:
        result.scraped_text = text
        result.strategy = "docx_text"
        result.run_id = extract_run_id(text)
        result.institution = infer_institution_from_folder(folder.name)
        return result

    # Step 2: OCR embedded images in those DOCX files.
    ocr_text_parts = []
    pages = 0
    for p in docx_paths:
        ocr_t, n = ocr_docx_images(p)
        if ocr_t:
            ocr_text_parts.append(ocr_t)
        pages += n
    if pages:
        result.ocr_pages = pages
    combined = (text + "\n" + "\n".join(ocr_text_parts)).strip()
    if len(combined) >= 200:
        result.scraped_text = combined
        result.strategy = "docx_ocr"
        result.run_id = extract_run_id(combined)
        result.institution = infer_institution_from_folder(folder.name)
        result.warnings.append(
            f"deep_extract_used_docx_ocr: scraped {pages} embedded images"
        )
        return result

    # Step 3: OCR PDF pages.
    pdf_paths = list(_walk_files(folder, ".pdf", max_depth=3))
    pdf_text_parts = []
    pdf_pages = 0
    for p in pdf_paths:
        pdf_t, n = ocr_pdf(p)
        if pdf_t:
            pdf_text_parts.append(pdf_t)
        pdf_pages += n
    if pdf_pages:
        result.ocr_pages += pdf_pages
    combined = (combined + "\n" + "\n".join(pdf_text_parts)).strip()
    if len(combined) >= 200:
        result.scraped_text = combined
        result.strategy = "pdf_ocr"
        result.run_id = extract_run_id(combined)
        result.institution = infer_institution_from_folder(folder.name)
        result.warnings.append(
            f"deep_extract_used_pdf_ocr: scraped {pdf_pages} pages"
        )
        return result

    # Step 4: folder-name-only fallback. Better than nothing — the
    # caller can UPSERT an `entities` row keyed on drive_folder_id +
    # surface "needs manual re-ingest" on the admin import audit.
    inst = infer_institution_from_folder(folder.name)
    if inst:
        result.strategy = "folder_name_only"
        result.scraped_text = ""
        result.institution = inst
        result.warnings.append(
            "deep_extract_only_folder_name: no text extracted from any file"
        )
        return result

    # Total failure — caller treats this as failed_other.
    result.strategy = "none"
    result.warnings.append(
        "deep_extract_yielded_nothing: 0 DOCX + 0 PDF readable in folder"
    )
    return result


def has_scoreable_content(text: str) -> bool:
    """True iff the scraped text looks like it might contain DMA score
    content (used to gate the LLM extractor in retry mode).

    Heuristic only — false positives are cheap (we just attempt LLM
    extraction + the validator rejects bad output); false negatives
    are operator-visible as a `failed_parse` quarantine row.
    """
    if not text or len(text) < 200:
        return False
    lower = text.lower()
    return sum(1 for h in _SCORE_HEADING_HINTS if h in lower) >= 3
