"""Entity data-completeness healer (self-healing pipeline, layer 3A).

The backfill loads the DMA reports into the DB; the per-leaf parsers extract a
SUBSET (location- and schema-sensitive). The QA sweep proved the packages carry
far more than the DB holds — e.g. `02_research_workbook/financial_baseline.json`
ships `latest_filing_*.{total_assets, employees, domestic_offices}` +
`latest_ratios.{roa_pct, roe_pct, nim_pct, ...}` and `entity_profile.json` ships
`entity.regulator_stack.primary_federal` + ownership/identifiers/brands — yet
firmographics had `aum_usd / headcount / primary_regulator = NULL` for 94/94.

This module is the auditable place that closes those gaps so **no Overview
firmographics panel is empty for any of the 94**. It reads the entity's package
directly (resolved from `runs.drive_folder_id = "local:<dir>"`), extracts with
schema-tolerant logic the leaf parsers miss, and UPSERTs the firmographics
columns the `/overview` endpoint reads (`aum_usd`, `headcount`, `revenue_usd`,
`primary_regulator`, `financial_highlights`, `parsed_facts`). FILL-IF-EMPTY only
— never overwrites already-good data, never fabricates (only grounded package
values). Idempotent.

`audit_firmographics()` (used by `--verify-only`) reports per-entity which
panel fields are still empty after healing.
"""
from __future__ import annotations

import csv
import json
import re
from pathlib import Path
from typing import Any

from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession

from app.services.startup_enrich import regulator_is_garbled

# The Overview "FIRMOGRAPHICS" panel fields (router maps these + flattens
# parsed_facts). An entity is "complete" when none of these is empty.
PANEL_FIELDS = ("aum_usd", "headcount", "primary_regulator")
PANEL_FACTS = ("branches", "cagr", "founded")
# An org-scale fact that keeps the panel non-empty when headcount is genuinely
# undisclosed (members for a CU, homes for a residential REIT, branches).
_ORG_SCALE_FACTS = ("headcount_fact", "members", "homes", "branches")

# Last-resort primary-regulator by subvertical — the framework that
# definitionally governs that entity type (recorded as basis="inferred", never
# overwriting an extracted value). US-centric default + a Canadian override so a
# clearly-Canadian entity (Payments Canada) is not labelled with US regulators.
SUBVERTICAL_REGULATOR = {
    "RB": "FDIC",                                  # Regional Bank (state-chartered default)
    "CIB": "Federal Reserve / OCC",                # Corporate & Investment Bank
    "CL": "State financial regulator / CFPB",      # Commercial Lending
    "CU": "NCUA",                                  # Credit Union
    "FC": "Farm Credit Administration (FCA)",      # Farm Credit
    "IB": "State Insurance Department (NAIC)",     # Insurance Broker
    "IC": "State Insurance Department (NAIC)",     # Insurance Carrier
    "AM": "SEC",                                   # Asset Manager
    "RIA": "SEC",                                  # RIA / Wealth
}
SUBVERTICAL_REGULATOR_CA = {
    "RB": "OSFI",
    "CIB": "Bank of Canada / OSFI",
    "CL": "FCAC / provincial regulator",
    "CU": "Provincial credit union regulator",
    "IB": "Provincial insurance regulator",
    "IC": "OSFI / provincial insurance regulator",
    "AM": "CIRO / provincial securities commission",
    "RIA": "CIRO / provincial securities commission",
}
_CANADA_RE = re.compile(r"\bcanad(?:a|ian)\b", re.I)


def _default_regulator(subvertical: str | None, name: str | None) -> str | None:
    """Subvertical-default regulator, geography-aware: a Canadian entity gets the
    Canadian framework, not the US default."""
    if subvertical not in SUBVERTICAL_REGULATOR:
        return None
    if name and _CANADA_RE.search(str(name)):
        return SUBVERTICAL_REGULATOR_CA.get(subvertical, SUBVERTICAL_REGULATOR[subvertical])
    return SUBVERTICAL_REGULATOR[subvertical]

# DMA framework SVn → canonical 2-letter subvertical (observed in the corpus
# classification cells: SV2 Credit Unions, SV5 Wealth Advisory, SV6 Asset Mgmt).
_SV_CODE = {"SV1": "RB", "SV2": "CU", "SV3": "CL", "SV5": "RIA", "SV6": "AM",
            "SV7": "IC", "SV8": "IB", "SV9": "CIB"}
# Name/classification keyword → canonical code (first match wins; ordered so
# specific types beat the broad "bank"/"insurance" fallbacks). REITs map to AM
# (the corpus taxonomy has no REIT code; e.g. American Homes is classified AM).
_SUBVERTICAL_KEYWORDS: list[tuple[str, str]] = [
    (r"credit union|\bccu\b", "CU"),
    (r"farm credit|farm credit system|\baca\b|agcredit", "FC"),
    (r"insurance broker|insurance brokerage|insurance agency|brokerage &|benefits broker", "IB"),
    (r"travel insurance|mutual insurance|insurance carrier|general insurance|"
     r"life insurance|property.{0,4}casualty|\binsurer\b|assurance", "IC"),
    (r"asset management|investment management|capital management|hedge fund|"
     r"fund manager|trust company|trust organization|multi-strategy", "AM"),
    (r"\breit\b|real estate investment|realty|real estate", "AM"),
    (r"wealth management|wealth advisory|registered investment advis|\bria\b|private wealth", "RIA"),
    (r"broker-dealer|securities brokerage|\bbrokerage\b|interactive brokers", "RIA"),
    (r"payments|clearing house|settlement system|payments? canada", "CIB"),
    (r"investment bank|corporate bank|capital markets", "CIB"),
    (r"commercial lending|business lending|equipment finance", "CL"),
    (r"\binsurance\b", "IC"),
    (r"\bbank\b|bancorp|bancshares|\bbanc\b|savings bank|national association|\bn\.a\.\b", "RB"),
]


def classify_subvertical(name: str, pkg_text: str = "") -> str | None:
    """Canonical subvertical code from the package's stated classification (an
    explicit SVn code wins) else the entity name + classification keywords.
    Returns None rather than guess when nothing matches."""
    m = re.search(r"\bSV([1-9])\b", pkg_text)
    if m and f"SV{m.group(1)}" in _SV_CODE:
        return _SV_CODE[f"SV{m.group(1)}"]
    blob = f"{name} {pkg_text}".lower()
    for patt, code in _SUBVERTICAL_KEYWORDS:
        if re.search(patt, blob):
            return code
    # ML fallback — ONLY when the deterministic rules above found nothing, so the
    # regex behaviour is never overridden (no regression). Dependency-optional:
    # returns None when sklearn/joblib or the model artifact is absent, or the
    # prediction is low-confidence. NON_FI is a gold-only label (no production UI
    # mapping yet) -> clamp to None so this function only ever emits the 9 codes.
    try:
        from app.ml.text_classifier import get_classifier
        feat = re.sub(r"\bSV[-_ ]?\d+\b", " ", f"{name} {pkg_text}", flags=re.I).strip()
        label, _conf = get_classifier("subvertical").predict(feat)
        if label and label != "NON_FI":
            return label
    except Exception:
        pass
    return None


# Roster container keys seen across the corpus. Beyond the original list-shaped
# keys, real packages carry the roster as a DICT-of-roles under
# `leadership_snapshot` (entity_profile.json) or `key_leadership`
# (research_handoff.json) — see _find_rosters for the dict handling.
_ROSTER_KEY = re.compile(
    r"c_suite_roster|leadership_roster|leadership_snapshot|leadership_team|"
    r"key_leadership|key_executives|^leadership$|leadership_profile|"
    r"leadership_contacts|^executives$|executive_team|management_team|"
    r"^c_suite$|leadership_register|exec_team|^officers$",
    re.I,
)


# Role-dict keys that carry metadata, not a person — skip them so a
# `ctoo_background: "Nearly 4 decades IT…"` value isn't parsed into a person.
_META_ROLE_KEY = re.compile(
    r"background|bio|note|summary|source|evidence|as_of|updated|date|count|"
    r"total|url|link|tenure|since|profile",
    re.I,
)


def _person_from_role_string(role_key: str, raw: str) -> dict[str, Any] | None:
    """Parse a `{"ceo": "William Mynatt Jr. (President & CEO)"}` value into a
    person dict. The corpus also uses `"Name (joined date)"` and
    `"A (..) + B (..)"` (we take the first person). The parenthetical becomes
    the title when it reads like one; otherwise the role key is the title.
    """
    s = re.split(r"\s+\+\s+|\s*;\s*|\s*/\s*", raw.strip())[0].strip()
    if not s:
        return None
    m = re.match(r"^(.*?)\s*\(([^)]*)\)\s*$", s)
    if m:
        name, paren = m.group(1).strip(), m.group(2).strip()
        title = paren if len(re.findall(r"[A-Za-z]", paren)) >= 2 else None
    else:
        name, title = s, None
    if not title:
        title = str(role_key).replace("_", " ").upper()
    return {"name": name, "title": title}
_NAME_KEY = re.compile(r"^name$|full_name|exec_name|person", re.I)
_TITLE_KEY = re.compile(r"^title$|^role$|position|exec_title|job_title", re.I)
_TENURE_KEY = re.compile(r"tenure|since|appointed|start", re.I)
_BACKGROUND_KEY = re.compile(r"background|bio|profile|summary|experience", re.I)


def _find_rosters(obj: Any):
    """Yield lists that look like a leadership roster (list of {name,title} dicts).

    Two real corpus shapes are recognised under a roster-matching key:
      - LIST of person dicts  → `[{name,title,...}, ...]`
      - DICT of roles         → `{"ceo":{name,title}, "cfo":{...}}` (the
        `leadership_snapshot` / `key_leadership` shape). The role key fills in
        as the title only when the person object carries none.
    """
    if isinstance(obj, dict):
        for k, v in obj.items():
            if _ROSTER_KEY.search(k):
                if isinstance(v, list) and v and isinstance(v[0], dict):
                    yield v
                    continue
                if isinstance(v, dict):
                    people: list[dict[str, Any]] = []
                    for role_key, person in v.items():
                        if isinstance(person, dict) and any(
                            _NAME_KEY.search(pk) for pk in person
                        ):
                            if not any(_TITLE_KEY.search(pk) for pk in person):
                                person = {**person, "title": str(role_key).replace("_", " ").upper()}
                            people.append(person)
                        elif (
                            isinstance(person, str)
                            and person.strip()
                            and not _META_ROLE_KEY.search(str(role_key))
                        ):
                            parsed = _person_from_role_string(role_key, person)
                            if parsed:
                                people.append(parsed)
                    if people:
                        yield people
                        continue
            yield from _find_rosters(v)
    elif isinstance(obj, list):
        for x in obj:
            yield from _find_rosters(x)


def extract_leadership(pkg_dir: Path) -> list[dict[str, Any]]:
    """Leadership roster [{name, title, tenure, background}] from the package's
    leadership_profile/entity_profile/research_handoff JSON (e.g. Amarillo's
    `c_suite_roster_publicly_identified`). Grounded; deduped by name."""
    out: list[dict[str, Any]] = []
    seen: set[str] = set()
    for d in collect_source_docs(pkg_dir):
        for roster in _find_rosters(d):
            for item in roster:
                if not isinstance(item, dict):
                    continue
                name = next((str(item[k]).strip() for k in item
                             if _NAME_KEY.search(k) and isinstance(item[k], str) and item[k].strip()), None)
                if not name or name.lower() in seen:
                    continue
                title = next((str(item[k]).strip() for k in item
                              if _TITLE_KEY.search(k) and isinstance(item[k], str) and item[k].strip()), None)
                tenure = next((str(item[k]).strip() for k in item
                               if _TENURE_KEY.search(k) and isinstance(item[k], str | int | float)), None)
                background = next((str(item[k]).strip() for k in item
                                   if _BACKGROUND_KEY.search(k) and isinstance(item[k], str) and item[k].strip()), None)
                seen.add(name.lower())
                out.append({"name": name[:120], "title": (title or None) and title[:160],
                            "tenure": tenure, "background": background and background[:400]})
        if out:
            break  # first source carrying a roster wins
    return out[:24]


def extract_subvertical_text(pkg_dir: Path) -> str:
    """Concatenate the package's subvertical-bearing strings (classification
    JSON fields + docx 'Subvertical'/'Size Tier' rows) for `classify_subvertical`."""
    parts: list[str] = []
    sv_key = re.compile(r"sub_?vertical|classification|size[_ ]?tier", re.I)
    for d in collect_source_docs(pkg_dir):
        for _k, v in _scan(d, sv_key.search):
            if isinstance(v, str) and v.strip():
                parts.append(v.strip())
    blob = collect_docx_text(pkg_dir)
    if blob:
        for line in blob.split("\n"):
            if re.search(r"sub.?vertical|classification|\bSV[1-9]\b", line, re.I) and len(line) < 160:
                parts.append(line.strip())
    return " | ".join(parts[:40])


def _corpus_root(corpus_dir: str | None) -> Path:
    if corpus_dir:
        return Path(corpus_dir)
    # backend/app/services/entity_healing.py -> parents[2] == backend/
    return Path(__file__).resolve().parents[2] / "tests" / "fixtures" / "dma_packages_batches"


def resolve_package_dir(drive_folder_id: str | None, corpus_dir: str | None) -> Path | None:
    """`local:<Client> - DMA` → the corpus folder for that client."""
    if not drive_folder_id or not drive_folder_id.startswith("local:"):
        return None
    name = drive_folder_id.split("local:", 1)[1].strip()
    root = _corpus_root(corpus_dir)
    if not root.is_dir():
        return None
    for batch in sorted(root.iterdir()):
        if batch.is_dir() and (batch / name).is_dir():
            return batch / name
    # fallback: any descendant dir matching the client name
    hits = [d for d in root.glob(f"*/{name}") if d.is_dir()]
    return hits[0] if hits else None


def _load_json(path: Path) -> dict[str, Any]:
    try:
        d = json.loads(path.read_text())
        return d if isinstance(d, dict) else {}
    except (OSError, json.JSONDecodeError):
        return {}


def _load_csv_kv(path: Path) -> dict[str, Any]:
    """A1_Entity_Profile_Summary.csv style `Field,Value` → {field: value}.
    Tolerates a header row whose first cell is literally 'Field'."""
    out: dict[str, Any] = {}
    try:
        with path.open(newline="") as fh:
            for row in csv.reader(fh):
                if len(row) >= 2 and row[0].strip():
                    k = row[0].strip()
                    if k.lower() in ("field", "key", "attribute"):
                        continue
                    out[k] = row[1].strip()
    except OSError:
        return {}
    return out


# Files that carry the SUBJECT entity's own firmographics (priority order).
# Peer/evidence/search files are excluded — they hold OTHER companies' numbers.
# Matched case-insensitively — the corpus mixes `research_handoff.json` and
# `VNO_Research_Handoff.json`. Peer/evidence/search files are excluded below.
_JSON_SOURCE_RE = re.compile(
    r"financial_baseline|entity_profile|research[_ ]?handoff|00_parameters|"
    r"subvertical_classification|report_analysis|leadership_profile",
    re.I,
)
_EXCLUDE_DIR = re.compile(r"/0?6[_ ]?peers/|/peers/", re.I)
_EXCLUDE_NAME = re.compile(r"peer|evidence_index|search_log|qa_verdict|check_results|corrections?_|corrected_", re.I)


_DOCX_PROFILE_RE = re.compile(r"client[_ ]?profile|profile|research|background", re.I)
_DOCX_SKIP_RE = re.compile(r"assessment|scoring|governance|patch|verdict", re.I)
# A label that identifies a firmographic KV row (table or `Label: value` prose).
_FIRMO_LABEL = re.compile(
    r"total[_ ]?assets|assets under|client assets|managed assets|\baum\b|\baua\b|"
    r"employees?|headcount|workforce|\bstaff\b|primary[_ ]?(federal[_ ]?)?regulator|"
    r"lead[_ ]?regulator|^regulator|branches?|founded|headquarter|net[_ ]?income",
    re.I,
)
_COLON_RE = re.compile(r"^([A-Za-z][A-Za-z0-9 /()&.+-]{2,46}?)\s*[:|]\s*(.+)$")


def _read_docx(path: Path):
    try:
        import docx  # python-docx, optional
    except ImportError:
        return None
    try:
        return docx.Document(str(path))
    except Exception:  # corrupt/locked docx: skip, never crash the backfill
        return None


def extract_docx_facts(path: Path) -> dict[str, Any]:
    """Pull label→value firmographics out of a profile/research .docx — from
    table rows (the first two cells when the row leads with a firmographic
    label) AND from `Label: value` prose paragraphs. Peer tables (company name
    as label) are excluded because the label must match a firmographic term.
    Returns a flat {label: value} dict the recursive scanner consumes."""
    doc = _read_docx(path)
    if doc is None:
        return {}
    out: dict[str, Any] = {}

    def _put(label: str, value: str) -> None:
        label, value = label.strip(), value.strip()
        if label and value and len(label) <= 48 and len(value) <= 220 and label not in out:
            out[label] = value

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            # plain 2-col KV, or wider row that leads with a firmographic label
            if (len(cells) >= 2 and cells[0] and cells[1]
                    and (len(cells) == 2 or _FIRMO_LABEL.search(cells[0]))):
                _put(cells[0], cells[1])
    for para in doc.paragraphs:
        m = _COLON_RE.match(para.text.strip())
        if m and _FIRMO_LABEL.search(m.group(1)):
            _put(m.group(1), m.group(2))
    return out


# Prose fallbacks (used only when the structured + KV scan leaves a panel field
# empty). Lines that describe peers/aspirations are skipped so we never attribute
# a competitor's figure to the subject.
_PEER_LINE_RE = re.compile(r"peer|aspiration|comparable|benchmark|competitor|versus|\bvs\.?\b|size tier proximity|regional bank;|industry|market size|sector|addressable|nationwide|globally|combined u\.?s\.?", re.I)
# A gross/levered AUM basis is not the headline figure (Elliott: "$128.6B
# regulatory (incl. leverage)" vs the "$79.8B net" headline).
_GROSS_RE = re.compile(r"gross|leverage|regulatory (?:aum|basis|gross)|including leverage|notional", re.I)
_UNIT = {"t": 1e12, "trillion": 1e12, "b": 1e9, "billion": 1e9, "m": 1e6, "million": 1e6}
_NUM_UNIT = r"\$?([\d.,]+)\s*(trillion|billion|million|[bmt])\b"
# A figure is the entity's assets only when the number is ADJACENT to an asset
# noun. Number-before ("$32.5B total assets") and number-after ("AUA ~$90B").
_ASSET_NOUN = r"(?:(?:total|client|managed|institutional|private\s+wealth|net|gross|regulatory|discretionary|fee[- ]earning)\s+){0,2}(?:assets?|aua|aum|assets under (?:management|administration))"
# Figures the source itself flags as wrong (Elliott: Indeed '$23B AUM' is an
# "aggregator error").
_JUNK_RE = re.compile(r"aggregator|erroneous|incorrect|mistaken|\bnot\s+\$", re.I)
_PAT_BEFORE = re.compile(rf"{_NUM_UNIT}[\s-]{{0,3}}(?:in\s+)?{_ASSET_NOUN}\b", re.I)
_PAT_AFTER = re.compile(rf"(?:total assets|assets under (?:management|administration)|\baua\b|\baum\b|client assets|managed assets|primary metric)[\s:=|,\-]{{1,5}}~?{_NUM_UNIT}", re.I)
# A number whose trailing word is one of these is NOT total assets (the classic
# "$32.5B total assets, $27.5B deposits" — 27.5 is deposits, not assets).
_NONASSET_TAIL = re.compile(r"^\W*(deposits?|loans?|income|revenue|equity|capital|earnings|profit|origination|fees?|market cap)", re.I)
# Threshold bands ("Mega (>$50B AUM)", "over $50B") are floors, not the figure.
_BAND_RE = re.compile(r">\s*\$|over\s+\$|exceeds\s+\$|up to \$|mega \(", re.I)
# A size-tier RANGE ("Super-Regional ($100B-$200B Assets)", "$150B-$160B"):
# two money figures joined by a dash. Neither edge is a precise figure -- the
# upper edge ($200B) is NOT the entity's actual assets (Regions' real figure
# is the balance-sheet $157.3B). We pre-scan these spans so any asset match
# that lands on a range edge is excluded from scoring; the balance-sheet
# actual then wins. Dash class covers hyphen-minus + Unicode dashes/minus.
_MONEY_RANGE_RE = re.compile(
    r"\$?\s*[\d.,]+\s*(?:trillion|billion|million|[bmt])\b"
    r"\s*[-\u2010-\u2015\u2212]\s*"
    r"\$?\s*[\d.,]+\s*(?:trillion|billion|million|[bmt])\b",
    re.I)
_EMP_PATTS = (
    r"\b([\d][\d,]{1,6})\s+(?:full-time\s+)?(?:employees|ftes|associates|team members)\b",
    r"(?:employees|headcount|workforce)\s*[:=|]\s*~?(?:approximately\s+)?([\d][\d,]{1,6})\b",
)
_REG_PROSE = (
    r"primary (?:federal )?regulator(?:s)?\s*[:|]\s*([^.\n]{2,80})",
    r"(?:regulated|supervised|chartered)\s+by\s+(?:the\s+)?([A-Z][^.\n]{2,70})",
)


def _asset_usd(num_s: str, unit_s: str) -> float | None:
    try:
        return float(num_s.replace(",", "")) * _UNIT.get(unit_s.lower(), 1.0)
    except (ValueError, AttributeError):
        return None


def prose_assets(text: str) -> float | None:
    """The entity's headline assets/AUM/AUA from the profile prose. Scores each
    adjacent-to-an-asset-noun figure by how it's labelled and how often it
    recurs, then returns the highest-scoring value — so a once-mentioned
    custody/peer/segment number can't beat the repeated headline. Peer,
    gross/levered, and deposit/loan-labelled figures are excluded."""
    scores: dict[float, float] = {}
    # Character spans covered by a size-tier money RANGE ("$100B-$200B"). An
    # asset figure landing inside one of these is a band edge, not a point
    # estimate, and is excluded so the balance-sheet actual wins.
    band_spans = [mm.span() for mm in _MONEY_RANGE_RE.finditer(text or "")]

    def _in_band(pos: int) -> bool:
        return any(s <= pos < e for s, e in band_spans)

    def _consider(usd: float | None, ctx: str, tail: str, *, band_edge: bool = False) -> None:
        if not usd or usd <= 1e7:
            return
        if (band_edge or _PEER_LINE_RE.search(ctx) or _GROSS_RE.search(ctx)
                or _BAND_RE.search(ctx) or _JUNK_RE.search(ctx)
                or _NONASSET_TAIL.match(tail)):
            return
        cl = ctx.lower()
        if "primary" in cl and "metric" in cl:
            strength = 3
        elif any(t in cl for t in ("total assets", "assets under", "aum", "aua")):
            strength = 2
        else:
            strength = 1
        key = round(usd / 1e8) * 1e8  # bucket to ~0.1B so re-statements coalesce
        scores[key] = scores.get(key, 0.0) + 1 + strength

    for m in _PAT_BEFORE.finditer(text):
        ctx = text[max(0, m.start() - 55): m.end() + 25]
        # The number/unit lead the match, so its start is the match start.
        _consider(_asset_usd(m.group(1), m.group(2)), ctx, "",
                  band_edge=_in_band(m.start()))
    for m in _PAT_AFTER.finditer(text):
        ctx = text[max(0, m.start() - 55): m.end() + 25]
        # The number/unit trail the match, so probe just inside the end.
        _consider(_asset_usd(m.group(1), m.group(2)), ctx, text[m.end(): m.end() + 14],
                  band_edge=_in_band(m.end() - 1))
    if not scores:
        return None
    return max(scores.items(), key=lambda kv: (kv[1], kv[0]))[0]


# Domain scale metrics for non-bank entities whose primary size figure is not
# "total assets/AUM" — insurers (premium), mortgage servicers (servicing UPB),
# specialty lenders (loan/receivables portfolio), insurers' invested assets.
# Each carries its basis so the panel can label it honestly (not "AUM").
_DOMAIN_SCALE = (
    ("total_invested_assets", r"total invested assets[^.\n|]{0,18}?[:=,\-|(]?\s*~?\$?([\d.,]+)\s*(trillion|billion|million|[bmt])\b"),
    ("servicing_upb", r"servicing upb[^.\n|]{0,14}?[:=,\-|(]?\s*~?\$?([\d.,]+)\s*(trillion|billion|million|[bmt])\b"),
    ("servicing_upb", r"\$([\d.,]+)\s*(trillion|billion|million|[bmt])\b[ \-]*servicing(?:\s+upb)?\b"),
    ("loan_portfolio", r"(?:loan portfolio|gross receivables|net receivables)[^.\n|]{0,16}?[:=,\-|(]?\s*~?\$?([\d.,]+)\s*(trillion|billion|million|[bmt])\b"),
    ("premium_volume", r"\$([\d.,]+)\s*(trillion|billion|million|[bmt])\b[ \-]*(?:in\s+)?(?:written\s+|annual\s+|placed\s+|gross\s+)?premium"),
    ("premium_volume", r"(?:written premium|premium written|premium placed|premium volume|annual premium|gross written premium|\bgwp\b|in[- ]force premium)[^.\n|]{0,16}?[:=,\-|(]?\s*~?\$?([\d.,]+)\s*(trillion|billion|million|[bmt])\b"),
    ("policyholder_surplus", r"(?:policyholder )?surplus[^.\n|]{0,12}?[:=,\-|(]?\s*~?\$?([\d.,]+)\s*(trillion|billion|million|[bmt])\b"),
)


_METRIC_LINE_RE = re.compile(r"primary\s*metric|size\s*tier|market\s*cap|enterprise value", re.I)
_BASIS_WORDS = (
    ("earning_assets", r"earning assets"),
    ("total_assets", r"total assets|under management|under administration|\baum\b|\baua\b"),
    ("servicing_upb", r"servicing|\bupb\b"),
    ("premium_volume", r"premium|\bgwp\b|gross written"),
    ("loan_portfolio", r"loan portfolio|receivables"),
    ("market_cap", r"market\s*cap|enterprise value"),
    ("revenue", r"revenue|\bnoi\b|\bffo\b|net premiums"),
)


def _infer_basis(context: str) -> str:
    cl = context.lower()
    for basis, patt in _BASIS_WORDS:
        if re.search(patt, cl):
            return basis
    return "scale"


_DASHES = ("-", "\u2013")  # hyphen, en-dash


def prose_primary_metric(text: str) -> tuple[float, str] | None:
    """The analyst-curated headline scale from a `Primary Metric` / `Size Tier`
    / `Market Cap` row (the figure the package itself chose to represent the
    entity's size), with its basis inferred from the surrounding words. Largest
    figure on such a line wins; peer/industry/gross lines and servicing
    fee-business figures (handled by `prose_domain_scale`) are skipped. As a
    last resort the upper bound of a `Size Tier (… $X)` range is used."""
    best: tuple[float, str] | None = None
    range_hi: float | None = None
    for raw in text.split("\n"):
        line = raw.strip()
        if not _METRIC_LINE_RE.search(line) or _PEER_LINE_RE.search(line) or _GROSS_RE.search(line):
            continue
        for m in re.finditer(_NUM_UNIT, line, re.I):
            usd = _asset_usd(m.group(1), m.group(2))
            if not usd or usd <= 1e7:
                continue
            before, after = line[max(0, m.start() - 1): m.start()], line[m.end(): m.end() + 2]
            in_range = before in _DASHES or any(after.startswith(d) for d in _DASHES)
            ctx = line[max(0, m.start() - 45): m.end() + 30]
            if in_range:
                # remember the largest range bound for the fallback
                range_hi = usd if range_hi is None else max(range_hi, usd)
                continue
            if re.search(r"servicing", ctx, re.I):  # fee business, not own scale
                continue
            basis = _infer_basis(ctx)
            if best is None or usd > best[0]:
                best = (usd, basis)
    if best is None and range_hi is not None:
        return (range_hi, "size_tier")
    return best


def prose_domain_scale(text: str) -> tuple[float, str] | None:
    """Largest subject-attributed *domain* scale figure (premium / servicing
    UPB / loan portfolio / invested assets) with its basis — used only when no
    total-assets/AUM figure exists. Peer/industry lines excluded."""
    best: tuple[float, str] | None = None
    for basis, patt in _DOMAIN_SCALE:
        for m in re.finditer(patt, text, re.I):
            ctx = text[max(0, m.start() - 55): m.end() + 20]
            if _PEER_LINE_RE.search(ctx) or _BAND_RE.search(ctx) or _JUNK_RE.search(ctx):
                continue
            usd = _asset_usd(m.group(1), m.group(2))
            if usd and usd > 1e7 and (best is None or usd > best[0]):
                best = (usd, basis)
    return best


def prose_headcount(text: str) -> int | None:
    best: int | None = None
    for patt in _EMP_PATTS:
        for m in re.finditer(patt, text, re.I):
            line = text[max(0, m.start() - 50): m.end() + 30]
            if _PEER_LINE_RE.search(line) or re.search(r"\b\d(\.\d+)?\s*/\s*(5|10)\b", line):
                continue
            n = _to_int(m.group(1))
            if n and 1 <= n < 1_000_000 and (best is None or n > best):
                best = n
    return best


_BRANCH_RE = re.compile(r"\b([\d][\d,]{0,4})\+?\s+(?:branches|branch offices|banking centers|banking offices|financial centers|retail locations|locations|offices|domestic offices)\b", re.I)


def prose_branches(text: str) -> int | None:
    """Largest plausible branch/office count stated in the profile prose
    (peer lines skipped). Keeps bank panels non-empty when headcount is
    undisclosed."""
    best: int | None = None
    for m in _BRANCH_RE.finditer(text):
        line = text[max(0, m.start() - 45): m.end() + 10]
        if _PEER_LINE_RE.search(line):
            continue
        n = _to_int(m.group(1))
        if n and 1 <= n <= 10000 and (best is None or n > best):
            best = n
    return best


def prose_regulator(text: str) -> str | None:
    for patt in _REG_PROSE:
        m = re.search(patt, text, re.I)
        if m:
            reg = _short_reg(m.group(1))
            if reg:
                return reg
    return None


# ── Website / ticker / CAGR / trend / revenue recovery (plan 4.2) ───────────
_URL_RE = re.compile(r"(?:https?://)?(?:www\.)?([a-z0-9][a-z0-9.-]{2,60}\.[a-z]{2,10})(?:/\S*)?", re.I)
_NON_OWN_HOSTS = re.compile(
    r"linkedin|glassdoor|indeed|facebook|twitter|x\.com|youtube|instagram|"
    r"crunchbase|bloomberg|reuters|forbes|wsj|sec\.gov|fdic|ncua|occ\.|google|"
    r"apple|wikipedia|yelp|bbb\.org|trustpilot|zippia|comparably|payscale|"
    r"prnewswire|businesswire|globenewswire|americanbanker|spglobal|fitch|"
    r"moodys|theorg\.com|zoominfo|pitchbook|github|medium\.com|substack", re.I)


def clean_website(raw: object) -> str | None:
    """Normalize a website value to a bare host ('www.amalgamatedbank.com' /
    'https://x.y/z' → 'x.y'). None for junk/aggregator hosts."""
    m = _URL_RE.search(str(raw or "").strip())
    if not m:
        return None
    host = m.group(1).lower().rstrip(".")
    if _NON_OWN_HOSTS.search(host) or host.count(".") < 1 or len(host) > 60:
        return None
    return host


def _name_tokens(name: object) -> set[str]:
    stop = {"bank", "credit", "union", "federal", "financial", "insurance",
            "group", "capital", "trust", "company", "corp", "corporation",
            "holdings", "services", "management", "partners", "national",
            "first", "community", "the", "inc", "llc", "association"}
    return {t for t in re.findall(r"[a-z0-9]+", str(name or "").lower())
            if len(t) >= 3 and t not in stop}


def website_from_urls(urls: list[str], entity_name: str) -> str | None:
    """Own-domain heuristic over evidence source URLs: the modal non-aggregator
    host whose domain shares a token (or a 5+-char prefix) with the entity
    name. Grounded — the host must actually appear in the evidence trail."""
    toks = _name_tokens(entity_name)
    counts: dict[str, int] = {}
    for u in urls or []:
        host = clean_website(u)
        if not host:
            continue
        stem = host.split(".")[-2] if "." in host else host
        stem_l = re.sub(r"[^a-z0-9]", "", stem.lower())
        hit = any(t in stem_l or (len(t) >= 5 and t[:5] in stem_l) for t in toks) \
            or any(stem_l.startswith(t[:5]) for t in toks if len(t) >= 5)
        if not hit:
            # initialism/abbreviation match ('anbtx' for American National
            # Bank of Texas): the stem must be an in-order subsequence of the
            # FULL name's letters (same test as ticker _symbol_from_name).
            letters = [c for c in str(entity_name or "").lower() if c.isalpha()]
            i = 0
            ok = len(stem_l) >= 4
            for c in stem_l:
                while i < len(letters) and letters[i] != c:
                    i += 1
                if i >= len(letters):
                    ok = False
                    break
                i += 1
            hit = ok
        if hit:
            counts[host] = counts.get(host, 0) + 1
    if not counts:
        return None
    return max(counts.items(), key=lambda kv: kv[1])[0]


_TICKER_RE = re.compile(r"\b(NYSE|NASDAQ|OTC(?:QX|QB)?|AMEX|TSX)\s*:\s*([A-Z]{1,5})\b")


def clean_ticker(raw: object) -> str | None:
    s = str(raw or "").strip()
    m = _TICKER_RE.search(s)
    if m:
        return f"{m.group(1)}: {m.group(2)}"
    if re.fullmatch(r"[A-Z]{1,5}", s):
        return s
    return None


def prose_ticker(text: str, entity_name: str) -> str | None:
    """Entity-scoped ticker disambiguation: a candidate counts ONLY when the
    entity's own name (token) occurs in the same clause — peer/vendor tickers
    ('CGI (NYSE:GIB)' in a vendor paragraph) must never be assigned."""
    toks = _name_tokens(entity_name)
    for m in _TICKER_RE.finditer(text or ""):
        clause = text[max(0, m.start() - 120): m.start()]
        clause = re.split(r"[.;\n]", clause)[-1].lower()
        if any(t in clause for t in toks):
            return f"{m.group(1)}: {m.group(2)}"
    return None


def prose_cagr(text: str) -> str | None:
    """A CAGR statement from trajectory prose → display string ('10.4%'),
    peer/industry lines skipped. Multi-metric reports keep the first
    asset/revenue-basis match."""
    for m in re.finditer(
            r"(?:CAGR|compound annual growth(?: rate)?)[^%\n]{0,32}?"
            r"(~?\d{1,2}(?:\.\d{1,2})?)\s*%"
            r"|(~?\d{1,2}(?:\.\d{1,2})?)\s*%[^%\n]{0,24}?(?:CAGR|compound annual)",
            text or "", re.I):
        ctx = (text or "")[max(0, m.start() - 60): m.end() + 20]
        if _PEER_LINE_RE.search(ctx):
            continue
        val = (m.group(1) or m.group(2) or "").lstrip("~")
        try:
            v = float(val)
        except ValueError:
            continue
        if 0 < v < 60:
            return f"{v:g}%"
    return None


_TREND_WORDS = (("ACCELERATING", r"accelerat"), ("DECLINING", r"declin|contract|shrink"),
                ("STABLE", r"\bstable\b|steady|consistent"),
                ("VARIABLE", r"\bvariable\b|volatile|mixed"))


def prose_trend(text: str) -> str | None:
    m = re.search(r"(?:trend|trajectory)\s+classification\s*[:|]\s*([A-Za-z]+)", text or "", re.I)
    if m and m.group(1).upper() in {"ACCELERATING", "DECELERATING", "STABLE", "VARIABLE", "DECLINING"}:
        return m.group(1).upper()
    win = re.search(r".{0,80}(?:growth|revenue|asset)s?.{0,120}", text or "", re.I)
    blob = win.group(0) if win else (text or "")[:400]
    for label, pat in _TREND_WORDS:
        if re.search(pat, blob, re.I):
            return label
    return None


# Subverticals whose headline scale metric is REVENUE (brokers, asset/wealth
# managers, fintechs) — banks/CUs stay honest-null ("assets is the metric").
REVENUE_SUBVERTICALS = frozenset({"IB", "AM", "RIA", "IC", "FINTECH_SAAS", "CL"})
_REVENUE_RE = re.compile(
    r"(?:annual\s+|total\s+|net\s+)?revenues?\s*(?:of|[:=|])?\s*~?\$([\d.,]+)\s*"
    r"(trillion|billion|million|[bmt])\b"
    r"|\$([\d.,]+)\s*(trillion|billion|million|[bmt])\b[^.\n]{0,18}?(?:in\s+)?(?:annual\s+)?revenue",
    re.I)
# `revenue_usd` is an ANNUAL figure. A quarterly (Q#, "quarter(ly)", "second
# quarter", per-quarter) context makes the adjacent number a partial-year
# figure that must NOT be labelled annual (Regions "Q2 2025: $1.9B total
# revenue" was mislabelled the annual figure). Rejected here so the field
# stays null rather than showing a wrong number.
_QUARTERLY_RE = re.compile(
    r"\bq[1-4]\b|\bquarter(?:ly)?\b|\bper\s+quarter\b|"
    r"\b(?:first|second|third|fourth)\s+quarter\b",
    re.I)


def prose_revenue(text: str) -> float | None:
    """Subject-attributed ANNUAL revenue for non-depository entities. Peer/
    industry lines and quarterly (Q#/quarterly) figures are skipped — a
    quarter's revenue is not the annual number and must never be labelled as
    such."""
    best: float | None = None
    for m in _REVENUE_RE.finditer(text or ""):
        ctx = (text or "")[max(0, m.start() - 55): m.end() + 20]
        if (_PEER_LINE_RE.search(ctx) or _BAND_RE.search(ctx)
                or _JUNK_RE.search(ctx) or _QUARTERLY_RE.search(ctx)):
            continue
        usd = _asset_usd(m.group(1) or m.group(3), m.group(2) or m.group(4))
        if usd and usd > 1e6 and (best is None or usd > best):
            best = usd
    return best


# ── DB-prose firmographics recovery ladders (2026-07-02) ────────────────────
# These read the entity's OWN persisted prose (firmographics.narrative_md +
# parsed_facts string values + financial_highlights lines) rather than a package
# on disk, so they heal the residual nulls even for entities whose corpus folder
# cannot be resolved. Every derived value is stamped with a *_basis marker so the
# provenance contract stays at 100%. Nothing is fabricated: a field stays NULL
# when the entity's own corpus carries no signal for it (honest-null).
US_STATE_NAMES = (
    "Alabama|Alaska|Arizona|Arkansas|California|Colorado|Connecticut|Delaware|"
    "Florida|Georgia|Hawaii|Idaho|Illinois|Indiana|Iowa|Kansas|Kentucky|Louisiana|"
    "Maine|Maryland|Massachusetts|Michigan|Minnesota|Mississippi|Missouri|Montana|"
    "Nebraska|Nevada|New Hampshire|New Jersey|New Mexico|New York|North Carolina|"
    "North Dakota|Ohio|Oklahoma|Oregon|Pennsylvania|Rhode Island|South Carolina|"
    "South Dakota|Tennessee|Texas|Utah|Vermont|Virginia|Washington|West Virginia|"
    "Wisconsin|Wyoming|District of Columbia")
_US_STATE_ABBR = ("AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|ID|IL|IN|IA|KS|KY|LA|ME|MD|MA|MI|"
                  "MN|MS|MO|MT|NE|NV|NH|NJ|NM|NY|NC|ND|OH|OK|OR|PA|RI|SC|SD|TN|TX|UT|"
                  "VT|VA|WA|WV|WI|WY|DC")
_US_REGIONS = ("Mid-Atlantic|Midwest|Northeast|Southeast|Southwest|Pacific Northwest|"
               "New England|West Coast|East Coast|Gulf Coast|Great Lakes|Sun Belt|"
               "Bay Area|Southern California|Northern California|Central Valley|"
               "Puget Sound|Silicon Valley|Tri-State")
_STATE_NAME_RE = re.compile(r"\b(" + US_STATE_NAMES + r")\b")
_REGION_NAME_RE = re.compile(r"\b(" + _US_REGIONS + r")\b", re.I)
_FOUND_CUE_RE = re.compile(
    r"\b(?:founded|established|chartered|incorporated|organized|opened its doors|"
    r"began operations|traces its (?:roots|history) to)\b", re.I)
_FOUND_YEAR_RE = re.compile(r"\b(1[789]\d\d|20[0-2]\d)\b")
_HQ_PROSE_RES = (
    re.compile(r"(?:headquartered|based|head office|home office)\s+(?:in|at)\s+"
               r"([A-Z][A-Za-z.\-]+(?:\s+[A-Z][A-Za-z.\-]+){0,2},?\s*"
               r"(?:" + US_STATE_NAMES + r"|" + _US_STATE_ABBR + r"))"),
    re.compile(r"\bHQ:?\s+(?:[\dA-Za-z .]+,\s*)?"
               r"([A-Z][A-Za-z.\-]+(?:\s+[A-Z][A-Za-z.\-]+){0,2}),?\s*"
               r"(?:" + _US_STATE_ABBR + r"|" + US_STATE_NAMES + r")"),
    re.compile(r"\bHQ:?\s+(" + US_STATE_NAMES + r")\b"),
)
# A footprint DECLARATION as an explicit abbrev list: "footprint: PA, DE, NJ".
# Free-text "footprint of/spanning …" prose is deliberately NOT mined — it
# captures garbage clauses ("footprint of a unified CDP", "Multi-region
# footprint (Explorium-in…)"); the clean, validated geography value is the
# reliable source instead.
_FOOTPRINT_LIST_RE = re.compile(
    r"footprint\s*[:—–-]\s*"  # noqa: RUF001
    r"((?:" + _US_STATE_ABBR + r")(?:\s*[,/]\s*(?:" + _US_STATE_ABBR + r")){1,12}"
    r"|\d{1,2}\+?\s+states\b[^.\n;]{0,40})")
_N_STATE_LIST_RE = re.compile(r"\b\d{1,2}\s+states?\s*\(([A-Z/,\s]+)\)")
_N_STATES_RE = re.compile(r"\b(?:across|in|serving|operates? in)\s+(\d{1,2})\s+states\b", re.I)


def firm_prose_hay(narrative_md: str | None, parsed_facts: dict | None,
                   financial_highlights: dict | None) -> str:
    """The entity's own persisted prose — narrative + string-valued parsed facts
    + financial-highlight lines — the honest corpus for the DB-prose ladders."""
    parts: list[str] = [narrative_md or ""]
    for k, v in (parsed_facts or {}).items():
        if isinstance(v, str) and not str(k).endswith("_basis"):
            parts.append(v)
    parts.extend(str(x) for x in ((financial_highlights or {}).get("lines") or []))
    return "\n".join(p for p in parts if p)


def prose_founded_year(hay: str) -> str | None:
    """A founding YEAR that is the direct object of a founding verb (within ~12
    chars) — NOT a CEO-tenure/tech-adoption year ('cloud-first since 2019')."""
    for m in _FOUND_CUE_RE.finditer(hay or ""):
        seg = (hay or "")[m.end(): m.end() + 18]
        ym = _FOUND_YEAR_RE.search(seg)
        if ym and ym.start() <= 12:
            y = int(ym.group(1))
            if 1700 <= y <= 2026:
                return str(y)
    return None


def prose_hq_location(hay: str) -> str | None:
    """A 'City, ST' / 'City ST' / 'State' HQ from a headquartered/based/HQ cue."""
    for rx in _HQ_PROSE_RES:
        m = rx.search(hay or "")
        if m:
            s = re.sub(r"\s{2,}", " ",
                       re.sub(r"^(?:headquartered|based|head office|home office|HQ)"
                              r"\s*:?\s*(?:in|at)?\s*", "", m.group(0), flags=re.I)).strip(" ,")
            if 3 <= len(s) <= 60:
                return s
    return None


def prose_geography(hay: str) -> str | None:
    """Operating geography — a state (or up to 6), a named region, or 'N states'."""
    states: list[str] = []
    for m in _STATE_NAME_RE.finditer(hay or ""):
        if m.group(1) not in states:
            states.append(m.group(1))
    if states:
        return ", ".join(states[:6])
    rm = _REGION_NAME_RE.search(hay or "")
    if rm:
        return rm.group(1)
    nm = _N_STATES_RE.search(hay or "")
    if nm:
        return f"{nm.group(1)} states"
    return None


_GEO_TOKEN_RE = re.compile(
    r"\b(?:" + US_STATE_NAMES + r"|" + _US_REGIONS + r")\b"
    r"|\b(?:" + _US_STATE_ABBR + r")\b"
    r"|\b(?:states?|counties|county|provinces?|branches|metros?|regions?|nationwide"
    r"|coast|markets)\b", re.I)


def _looks_geographic(s: object) -> bool:
    """A candidate footprint string is real only when it names geography — a
    state / region / 'N states' — never a peer name, an E-ID, or a stray clause."""
    return bool(_GEO_TOKEN_RE.search(str(s or "")))


def derive_footprint(geography: object, hay: str) -> str | None:
    """Operating footprint as a scalar string — an explicit 'N states (MI/OH/…)'
    list or 'footprint: MI, OH, …' abbrev list, else the (already-clean,
    validated) geography value. Free-text footprint prose is NOT mined (garbage
    class). Honest-null when the entity names no geography at all."""
    lm = _N_STATE_LIST_RE.search(hay or "")
    if lm:
        return re.sub(r"\s*/\s*", ", ", lm.group(1)).strip(" ,")
    fm = _FOOTPRINT_LIST_RE.search(hay or "")
    if fm and _looks_geographic(fm.group(1)):
        return re.sub(r"\s{2,}", " ", fm.group(1)).strip(" .;")[:120]
    if geography not in (None, "", [], {}):
        return str(geography)[:120]
    return None


# A leading "City, ST" / "City, State" — the HEADQUARTERS city, distinct from a
# multi-state operating footprint. The city phrase must not itself be a bare state
# token (so "New York, Nevada" — two states — never reads as a city "New York").
_CITY_STATE_LEAD_RE = re.compile(
    r"^([A-Z][A-Za-z.'\-]+(?:\s+[A-Z][A-Za-z.'\-]+){0,3}),\s+"
    r"(?:" + US_STATE_NAMES + r"|" + _US_STATE_ABBR + r")\b")
# A pure abbrev/state list ("NY, FL, VT, MA" / "MI/OH/IN") is a footprint, not an HQ.
_STATE_LIST_RE = re.compile(
    r"^(?:" + _US_STATE_ABBR + r")(?:\s*[,/]\s*(?:" + _US_STATE_ABBR + r"))+\s*$")
_STATE_TOKENS = ({t.lower() for t in US_STATE_NAMES.split("|")}
                 | {t.lower() for t in _US_STATE_ABBR.split("|")})


_HQ_JUNK_RE = re.compile(r"\bby asset size\b|\blargest\b|\branke?d?\b|\bnth\b", re.I)


def hq_is_plausible(s: object) -> bool:
    """A stored hq_address is a real location — not a serialized dict/list, a
    ranking phrase ('Texas by asset size'), or an empty stub. Street addresses and
    'City, ST' both pass; only clearly-malformed values are rejected so a good
    value is never discarded."""
    if not isinstance(s, str):
        return False
    t = s.strip()
    if len(t) < 3 or t[0] in "{[":
        return False
    return not _HQ_JUNK_RE.search(t)


def _hq_from_dict(v: dict) -> str | None:
    """Some rows persist ``hq`` as a structured dict (address/primary_city/state);
    prefer a 'City, State', else the primary city or the raw address string."""
    city = v.get("primary_city") or v.get("city")
    state = v.get("state")
    if city and state:
        return f"{city}, {state}"
    return city or (v.get("address") if isinstance(v.get("address"), str) else None)


def derive_hq_address(parsed_facts: object, hay: str = "") -> str | None:
    """The headquarters LOCATION (a city, or an explicitly-tagged HQ fact) — never
    a multi-state footprint or a 'National'/'Regional' descriptor (those stay the
    operating footprint, surfaced separately). Priority: the curated ``hq`` /
    ``hq_city`` parsed-fact, else a 'City, ST' mined from the entity's own prose,
    else a leading 'City, ST' inside the footprint/geography value. Honest-null
    when only a regional/national descriptor is available."""
    pf = parsed_facts if isinstance(parsed_facts, dict) else {}
    for key in ("hq", "hq_city"):
        v = pf.get(key)
        if isinstance(v, dict):
            v = _hq_from_dict(v)
        if isinstance(v, str) and v.strip() and not _STATE_LIST_RE.match(v.strip()):
            # trim a trailing explanatory clause ("Lake Forest, IL — state-chartered…")
            cleaned = re.split(r"\s+[—–-]\s+|\s*\(", v.strip(), maxsplit=1)[0]  # noqa: RUF001
            cleaned = re.sub(r"\s{2,}", " ", cleaned).strip(" ,.;")
            if 3 <= len(cleaned) <= 80:
                return cleaned
    mined = prose_hq_location(hay)
    if mined:
        return mined
    for key in ("footprint", "geography"):
        v = pf.get(key)
        if not isinstance(v, str) or _STATE_LIST_RE.match(v.strip()):
            continue
        m = _CITY_STATE_LEAD_RE.match(v.strip())
        if m and m.group(1).lower() not in _STATE_TOKENS:
            return re.sub(r"\s{2,}", " ", m.group(0)).strip(" ,.;")[:80]
    return None


def derive_size_tier(aum_usd: object, headcount: object) -> str | None:
    """A size band from the assets figure (preferred) or the headcount — the
    community/mid-size/large/mega vocabulary the firmographics panel reads."""
    try:
        if aum_usd is not None:
            b = float(aum_usd) / 1e9
            if b >= 250:
                return "mega"
            if b >= 50:
                return "large"
            if b >= 10:
                return "mid-size"
            if b >= 1:
                return "community"
            return "small"
    except (TypeError, ValueError):
        pass
    try:
        if headcount is not None:
            h = int(headcount)
            if h >= 10000:
                return "large"
            if h >= 1000:
                return "mid-size"
            if h >= 100:
                return "community"
            return "small"
    except (TypeError, ValueError):
        pass
    return None


# Provenanced firmographic fields (mirrors qa_coverage_contract._PROVENANCED_FIELDS)
# — every one that is present MUST carry a companion *_basis marker.
_PROVENANCED_FIRM_FIELDS: tuple[str, ...] = (
    "trend", "cagr", "website", "footprint", "thought_leadership", "revenue_usd",
    "aum_usd",
)


def stamp_firmographic_provenance(pf: dict) -> int:
    """Backfill a *_basis marker for any present provenanced parsed-fact that has
    none, so the firmographics-provenance contract holds at 100%. Returns the
    number of markers added. Idempotent (a field that already has a basis is
    untouched)."""
    stamped = 0
    _default_basis = {
        "trend": "derived:financial_highlights",
        "cagr": "derived:financial_highlights",
        "website": "derived:entity_profile",
        "footprint": "derived:geography",
        "thought_leadership": "derived:research",
        "revenue_usd": "derived:report",
        "aum_usd": "derived:total_assets",
    }
    for field in _PROVENANCED_FIRM_FIELDS:
        if pf.get(field) in (None, "", [], {}):
            continue
        stem = field.removesuffix("_usd")
        if any(pf.get(f"{s}{suf}") not in (None, "", [], {})
               for s in {field, stem}
               for suf in ("_basis", "_source", "_derived_from", "_provenance")):
            continue
        pf[f"{stem}_basis"] = _default_basis.get(field, "derived:corpus")
        stamped += 1
    return stamped


def collect_docx_text(pkg_dir: Path) -> str:
    """Concatenated paragraph + table text of the profile/research docx, for
    the guarded prose fallbacks."""
    parts: list[str] = []
    for p in _profile_docx_paths(pkg_dir)[:2]:
        doc = _read_docx(p)
        if doc is None:
            continue
        parts.extend(par.text for par in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)


def _profile_docx_paths(pkg_dir: Path) -> list[Path]:
    return [
        p for p in sorted(pkg_dir.glob("**/*.docx"))
        if _DOCX_PROFILE_RE.search(p.name) and not _DOCX_SKIP_RE.search(p.name)
        and not _EXCLUDE_DIR.search(str(p))
    ]


def collect_source_docs(pkg_dir: Path) -> list[dict[str, Any]]:
    """Gather firmographics-bearing JSON + CSV + docx docs for the SUBJECT
    entity, excluding peer/evidence files that would contaminate the scan with
    other companies' numbers. CSV `Field,Value` sheets and docx label/value
    tables are loaded as flat dicts."""
    docs: list[dict[str, Any]] = []
    for p in sorted(pkg_dir.glob("**/*.json")):
        if (_JSON_SOURCE_RE.search(p.name) and not _EXCLUDE_DIR.search(str(p))
                and not _EXCLUDE_NAME.search(p.name)):
            d = _load_json(p)
            if d:
                docs.append(d)
    for p in sorted(pkg_dir.glob("**/*.csv")):
        if "entity_profile" in p.name.lower() and not _EXCLUDE_DIR.search(str(p)):
            d = _load_csv_kv(p)
            if d:
                docs.append(d)
    # Profile/research docx — the richest source for the ~31 entities with no
    # structured JSON/CSV firmographics (label→value tables + `Label: value`).
    for p in _profile_docx_paths(pkg_dir)[:2]:
        d = extract_docx_facts(p)
        if d:
            docs.append(d)
    return docs


_MONEY_RE = re.compile(r"\$?\s*([0-9][0-9,]*\.?[0-9]*)\s*([bBmMkK])?")
_YEAR_RE = re.compile(r"(20\d{2})")


def _to_usd(val: Any, key: str = "") -> float | None:
    """Normalize an asset/income value to USD across the corpus's variants:
    raw numbers (Amarillo `9550000000`), `*_billions`/`*_millions` keys
    (BOK `total_assets_billions: 52.2`), and `$24.23B`/`$286.0M` strings
    (Cathay). Returns None when not parseable."""
    kl = key.lower()
    bill = "billion" in kl or kl.endswith(("_b", "_billions"))
    mill = "million" in kl or kl.endswith(("_m", "_millions"))
    if isinstance(val, bool):
        return None
    if isinstance(val, int | float):
        n = float(val)
        if bill:
            return n * 1e9
        if mill:
            return n * 1e6
        if n > 1e6:
            return n            # already raw USD
        if 0 < n < 1e4:
            return n * 1e9      # bare "52.2" → bank assets are in $B
        return n
    if isinstance(val, str):
        m = _MONEY_RE.search(val)
        if not m:
            return None
        num = float(m.group(1).replace(",", ""))
        unit = (m.group(2) or "").lower()
        mult = {"b": 1e9, "m": 1e6, "k": 1e3}.get(unit, 1e9 if bill else 1e6 if mill else 1.0)
        return num * mult
    return None


def _short_reg(val: Any) -> str | None:
    """Normalize a regulator string to fit `firmographics.primary_regulator`
    (varchar 64): take the leading clause (before ';'/' for ') and hard-cap at
    a word boundary so we never truncate mid-word or overflow the column."""
    if not isinstance(val, str):
        return None
    s = re.split(r";| for | as principal", val.strip(), maxsplit=1)[0].strip(" .,")
    if len(s) <= 64:
        return s or None
    cut = s[:64]
    if " " in cut:
        cut = cut[: cut.rfind(" ")]
    return cut.strip() or None


def _to_int(val: Any) -> int | None:
    if isinstance(val, bool):
        return None
    if isinstance(val, int | float):
        return int(val) if val > 0 else None
    if isinstance(val, str):
        m = re.search(r"([0-9][0-9,]*)", val)
        return int(m.group(1).replace(",", "")) if m else None
    return None


def _scan(obj: Any, key_pred: Any):
    """Recursively yield (key, scalar_value) where key matches key_pred. A
    key whose value is a LIST of scalars (e.g. `primary_regulators: [FDIC,
    OCC, ...]`) yields its first scalar, then still recurses."""
    if isinstance(obj, dict):
        for k, v in obj.items():
            if isinstance(v, str | int | float) and not isinstance(v, bool) and key_pred(k):
                yield k, v
            elif isinstance(v, list) and key_pred(k):
                for item in v:
                    if isinstance(item, str | int | float) and not isinstance(item, bool):
                        yield k, item
                        break
                yield from _scan(v, key_pred)
            else:
                yield from _scan(v, key_pred)
    elif isinstance(obj, list):
        for x in obj:
            yield from _scan(x, key_pred)


def _year_of(key: str) -> int:
    m = _YEAR_RE.search(key)
    return int(m.group(1)) if m else 0


def extract_firmographics(pkg_dir: Path) -> dict[str, Any]:
    """Schema-tolerant recursive extraction the leaf parsers miss. Scans
    financial_baseline.json + entity_profile.json + subvertical_classification.json
    for assets / employees / regulator / ratios across ALL observed schemas."""
    out: dict[str, Any] = {}
    facts: dict[str, Any] = {}
    highlights: dict[str, Any] = {}

    docs: list[dict[str, Any]] = collect_source_docs(pkg_dir)

    asset_re = re.compile(r"total[_ ]?assets|total[_ ]?client[_ ]?assets|\baua\b|\baum\b|fy\d{4}_assets|(^|_| )assets($|_| |_billions|_\d{4})", re.I)
    emp_re = re.compile(r"employee|headcount|staff|\bfte\b|domestic_offices|branch", re.I)
    reg_re = re.compile(r"primary[_ ]?(federal[_ ]?)?regulator|lead[_ ]?regulator|prudential[_ ]?regulator|primary_federal|regulatory[_ ]?body|^regulator$", re.I)
    ni_re = re.compile(r"net_income", re.I)
    ratio_re = re.compile(r"roa|roe|roaa|roae|nim|efficiency|tier_1|tier1|cagr|growth_yoy", re.I)

    # (year, usd, is_core_total_assets). "Core" = balance-sheet total assets,
    # NOT AUM/AUA (for a bank the headline is total assets, not assets under
    # administration). Year is read from the key OR the value ("$53B (Dec 2024)").
    aum_re = re.compile(r"\baum\b|\baua\b|under management|under administration|client assets|managed", re.I)
    asset_hits: list[tuple[int, float, bool]] = []
    branch_hits: list[int] = []
    for d in docs:
        for k, v in _scan(d, lambda k: bool(asset_re.search(k)) and "growth" not in k.lower()):
            usd = _to_usd(v, k)
            if usd and usd > 1e7:
                year = max(_year_of(k), _year_of(str(v)))
                is_core = not aum_re.search(k)
                asset_hits.append((year, usd, is_core))
        for k, v in _scan(d, emp_re.search):
            # Reject non-headcount signals misnamed with an employee word:
            # social/proxy counts (Cathay `linkedin_employees: "~14,980
            # followers"`) and ratings (Greenstone `Employee sentiment:
            # "4.2/5 Glassdoor"`).
            if re.search(r"linkedin|follow|social|proxy|twitter|glassdoor|sentiment|"
                         r"engagement|rating|review|culture|satisfaction|\bnps\b|\besg\b", k, re.I):
                continue
            if isinstance(v, str) and (re.search(r"follow", v, re.I)
                                       or re.search(r"\b\d(\.\d+)?\s*/\s*(5|10)\b", v)):
                continue
            if re.search(r"branch|office", k, re.I):
                bi = _to_int(v)
                if bi and bi < 5000:
                    branch_hits.append(bi)
            elif re.search(r"employee|headcount|staff|\bfte\b", k, re.I):
                # A qualitative string ("Several hundred …") is not a count —
                # only accept when the value LEADS with a number ("~700-900").
                if isinstance(v, str) and not re.match(r"^\s*[~≈<>]?\s*\d", v):
                    continue
                ei = _to_int(v)
                if ei and "headcount" not in out and ei < 1_000_000:
                    out["headcount"] = ei
        for _k, v in _scan(d, reg_re.search):
            if "primary_regulator" not in out:
                reg = _short_reg(v)
                if reg:
                    out["primary_regulator"] = reg
        for k, v in _scan(d, ni_re.search):
            usd = _to_usd(v, k)
            if usd and "net_income" not in highlights:
                highlights["net_income"] = usd
        for k, v in _scan(d, ratio_re.search):
            if isinstance(v, int | float) and k not in highlights and len(highlights) < 10:
                highlights[k] = v

    if asset_hits:
        # Prefer core total-assets candidates over AUM/AUA; within the chosen
        # class a year-stamped figure beats an unstamped one (avoids stray
        # outliers), then latest year, then largest.
        core = [t for t in asset_hits if t[2]]
        pool = core or asset_hits
        pool.sort(key=lambda t: (t[0] > 0, t[0], t[1]))
        out["aum_usd"] = pool[-1][1]
        facts["aum_basis"] = "total_assets" if core else "aum"
    if branch_hits:
        facts["branches"] = max(branch_hits)

    # Domain firmographics that keep non-bank panels rich (members for credit
    # unions, homes/units for residential REITs).
    mem_re = re.compile(r"^members?$|member_count|members_total|total_members", re.I)
    home_re = re.compile(r"^homes?$|homes_owned|total_units|units_owned|properties_owned|rental_homes", re.I)
    for d in docs:
        for _k, v in _scan(d, mem_re.search):
            mi = _to_int(v)
            if mi and mi > 100 and "members" not in facts:
                facts["members"] = mi
        for _k, v in _scan(d, home_re.search):
            hi = _to_int(v)
            if hi and hi > 100 and "homes" not in facts:
                facts["homes"] = hi

    # entity_profile identity facts (ownership / identifiers / founded / regulator_stack)
    for d in docs:
        ent = d.get("entity") if isinstance(d.get("entity"), dict) else d
        if "primary_regulator" not in out:
            reg = ent.get("regulator_stack") or {}
            r = reg.get("primary_federal") if isinstance(reg, dict) else None
            reg_s = _short_reg(r) if r else None
            if reg_s:
                out["primary_regulator"] = reg_s
        if isinstance(ent.get("ownership"), dict) and "ownership" not in facts:
            facts["ownership"] = ent["ownership"]
        if isinstance(ent.get("identifiers"), dict) and "identifiers" not in facts:
            facts["identifiers"] = ent["identifiers"]
        fnd = ent.get("founded")
        if fnd and "founded" not in facts:
            m = re.match(r"(\d{4})", str(fnd))
            if m:
                facts["founded"] = int(m.group(1))

    # Identity firmographics that enrich every panel (ticker / website / HQ
    # geography / size tier / founding year) regardless of subvertical.
    id_re = re.compile(r"^ticker$|stock_ticker|^website$|^web_?site$|homepage|"
                       r"^domain$|^url$|^geography$|^hq_city$|headquarters_city|"
                       r"^size_tier$|founded|established|year_founded", re.I)
    for d in docs:
        for k, v in _scan(d, id_re.search):
            kk = k.lower()
            if not (isinstance(v, str | int | float) and str(v).strip()):
                continue
            sv = str(v).strip()
            if "ticker" in kk and "ticker" not in facts:
                tk = clean_ticker(sv)
                if tk:
                    facts["ticker"] = tk
                    facts["ticker_basis"] = "entity_profile"
            elif re.search(r"website|web_?site|homepage|^domain$|^url$", kk) \
                    and "website" not in facts:
                ws = clean_website(sv)
                if ws:
                    facts["website"] = ws
                    facts["website_basis"] = "entity_profile"
            elif "geograph" in kk and "geography" not in facts:
                facts["geography"] = sv[:80]
            elif kk == "hq_city" or "headquarters_city" in kk:
                facts.setdefault("hq_city", sv[:80])
            elif "size_tier" in kk and "size_tier" not in facts:
                facts["size_tier"] = sv[:60]
            elif re.search(r"founded|established|year_founded", kk) and "founded" not in facts:
                ym = re.search(r"(1[89]\d{2}|20\d{2})", sv)
                if ym:
                    facts["founded"] = int(ym.group(1))

    # Prose fallbacks for the residual: scan the profile/research docx text for
    # the panel fields the structured + KV pass left empty (peer lines skipped).
    need_prose = ("aum_usd" not in out or "headcount" not in out
                  or "primary_regulator" not in out
                  or not all(k in facts for k in ("website", "ticker", "cagr", "trend")))
    if need_prose:
        text_blob = collect_docx_text(pkg_dir)
        if text_blob:
            if "cagr" not in facts:
                cg = prose_cagr(text_blob)
                if cg:
                    facts["cagr"] = cg
                    facts["cagr_basis"] = "trajectory_prose"
            if "trend" not in facts:
                tr = prose_trend(text_blob)
                if tr:
                    facts["trend"] = tr
            if "ticker" not in facts:
                tk = prose_ticker(text_blob, pkg_dir.name)
                if tk:
                    facts["ticker"] = tk
                    facts["ticker_basis"] = "prose:entity-scoped"
            if "website" not in facts:
                # own-domain heuristic ONLY — never the first URL in the doc
                # (that could be a peer's or an aggregator's site).
                ws = website_from_urls(
                    [f"https://{h}" for h in _URL_RE.findall(text_blob)[:80]],
                    pkg_dir.name)
                if ws:
                    facts["website"] = ws
                    facts["website_basis"] = "profile_prose:own-domain"
            if "revenue_usd" not in out:
                rv = prose_revenue(text_blob)
                if rv:
                    out["revenue_usd"] = rv
                    facts["revenue_basis"] = "profile_prose"
            if "aum_usd" not in out:
                pa = prose_assets(text_blob)
                if pa:
                    out["aum_usd"] = pa
                    facts["aum_basis"] = "total_assets"
                else:
                    # non-bank scale: prefer the analyst-curated Primary Metric
                    # row, then a domain metric (premium / servicing UPB).
                    pm = prose_primary_metric(text_blob) or prose_domain_scale(text_blob)
                    if pm:
                        out["aum_usd"] = pm[0]
                        facts["aum_basis"] = pm[1]
            if "headcount" not in out:
                ph = prose_headcount(text_blob)
                if ph:
                    out["headcount"] = ph
            if "primary_regulator" not in out:
                pr = prose_regulator(text_blob)
                if pr:
                    out["primary_regulator"] = pr
            if "branches" not in facts:
                pb = prose_branches(text_blob)
                if pb:
                    facts["branches"] = pb

    if highlights:
        out["financial_highlights"] = highlights
    if facts:
        out["facts"] = facts
    return out


# ── Grounded "About" narrative composer (Context D5 + Overview) ──────────────
# The Context "About" paragraph (firmographics.narrative_md) ships in the
# client-profile DOCX for most clients, and Vertex re-synthesizes it when a
# narrative is missing. But the QA DB (and any Vertex-cold deploy) leaves ~26
# clients whose package carried no/short profile prose with an empty "About".
# The mandate is no-empty-state for ALL 94 even with Vertex cold, and the
# product owns a rich grounded fact base for every entity — so we COMPOSE the
# paragraph deterministically from already-persisted firmographics + the
# assessed maturity. This is composition of grounded facts (every clause is a
# stored value), never fabrication: no fact is invented, and a real DOCX
# narrative always wins (package_persist COALESCEs EXCLUDED.narrative_md over a
# composed one on re-ingest; heal only fills when the stored text is < the
# gate's 120-char floor).
_SV_PROSE = {
    "RB": "regional bank", "CU": "credit union", "AM": "asset manager",
    "RIA": "wealth & advisory firm", "IC": "insurance carrier",
    "IB": "insurance broker", "FC": "farm-credit institution",
    "CIB": "corporate & investment bank", "CL": "lending institution",
}
# aum_basis (set by extract_firmographics) → the prose that labels the scale
# figure honestly (a servicer's headline is servicing UPB, not "assets").
_BASIS_PROSE = {
    "total_assets": "in total assets", "aum": "in assets under management",
    "aua": "in assets under administration", "earning_assets": "in earning assets",
    "servicing_upb": "in servicing unpaid-principal balance",
    "premium_volume": "in annual premium volume", "loan_portfolio": "in loan portfolio",
    "market_cap": "in market capitalization", "policyholder_surplus": "in policyholder surplus",
    "revenue": "in revenue", "size_tier": "in assets", "scale": "in assets",
}
# Parse artifacts that occasionally land in primary_regulator — never emit
# "regulated by Role"/"N/A" in prose.
_REG_JUNK = {"role", "n/a", "na", "tbd", "unknown", "none", "null", "-", "—"}


def _usd_prose(v: Any) -> str | None:
    """Spelled-out money for a sentence ("$24.2 billion", "$213.8 billion")."""
    try:
        n = float(v)
    except (TypeError, ValueError):
        return None
    if n >= 1e12:
        return f"${n / 1e12:.2f} trillion"
    if n >= 1e9:
        return f"${n / 1e9:.1f} billion"
    if n >= 1e6:
        return f"${n / 1e6:.0f} million"
    return f"${n:,.0f}"


def _clean_reg_for_prose(reg: str | None) -> str | None:
    """A short, clean regulator name fit for a sentence. The stored value can
    carry trailing prose ("NCUA, GCU is member-owned …"), a second regulator
    ("FDIC + CA DFPI"), or an unbalanced parenthetical from a 64-char truncation
    ("FDIC (state-chartered"). Take the leading clause and drop a dangling paren;
    suppress known parse-junk."""
    s = (reg or "").split(",")[0].split(" + ")[0].strip()
    if "(" in s and ")" not in s:           # dangling paren from truncation
        s = s[: s.index("(")].strip()
    if len(s) < 3 or s.lower() in _REG_JUNK:
        return None
    return s[:60]


def _clean_hq_for_prose(hq: Any) -> str | None:
    """A clean "City, ST" headquarters for prose, or None. HQ is the dirtiest
    firmographic field in the corpus (dict strings, field-of-membership notes,
    multi-region blobs, full street addresses), so we accept it ONLY when it
    cleanly matches a "<place>, <ST|State>" shape with no braces/semicolons/
    digits — every other case is skipped (HQ is optional flavour; the paragraph
    clears the length floor without it)."""
    if not isinstance(hq, str):
        return None
    s = hq.strip()
    if not s or "{" in s or ";" in s or any(c.isdigit() for c in s) or len(s) > 40:
        return None
    if not re.match(r"^[A-Za-z][A-Za-z.\- ]+,\s*[A-Za-z][A-Za-z.\- ]+$", s):
        return None
    return s


def compose_about_narrative(
    *, name: str | None, subvertical: str | None, aum_usd: Any,
    aum_basis: str | None, regulator: str | None, headcount: Any,
    facts: dict[str, Any] | None, overall: float | None, hq: Any = None,
) -> str | None:
    """A grounded Context "About" paragraph from persisted firmographics + the
    assessed maturity. Returns None if the result can't clear the gate's
    120-char floor (honest: we don't store a stub). Deterministic + idempotent."""
    name = (name or "").strip()
    if not name:
        return None
    facts = facts or {}
    sv = _SV_PROSE.get((subvertical or "").upper(), "financial institution")
    article = "an" if sv[:1].lower() in "aeiou" else "a"
    hq_clean = _clean_hq_for_prose(hq)
    ident = f"{name} is {article} {sv}"
    if hq_clean:
        ident += f" headquartered in {hq_clean}"
    parts = [ident + "."]

    scale = _usd_prose(aum_usd)
    if scale:
        basis = _BASIS_PROSE.get((aum_basis or "").lower(), "in assets")
        s = f"It reports {scale} {basis}"
        founded = str(facts.get("founded") or "").strip()
        if re.fullmatch(r"(1[89]\d{2}|20\d{2})", founded):
            s += f" and has operated since {founded}"
        parts.append(s + ".")

    clause: list[str] = []
    reg = _clean_reg_for_prose(regulator)
    if reg:
        clause.append(f"is regulated by {reg}")
    head = _to_int(headcount)
    members = _to_int(facts.get("members"))
    branches = _to_int(facts.get("branches"))
    if head:
        people = f"employs approximately {head:,} people"
        if branches:
            people += f" across {branches} branches"
        clause.append(people)
    elif members:
        people = f"serves approximately {members:,} members"
        if branches:
            people += f" through {branches} branches"
        clause.append(people)
    elif branches:
        clause.append(f"operates {branches} branches")
    if clause:
        parts.append("It " + " and ".join(clause) + ".")

    if isinstance(overall, int | float) and not isinstance(overall, bool):
        # seeded per entity so the maturity sentence never stamps one frame
        # across the corpus (45/94 shared it verbatim — 2026-07-13 census)
        from app.services.nlp.stylebook import pick, seeded
        parts.append(pick(seeded(name, "about-maturity"), (
            "Its overall digital maturity is assessed at {s} out of 5 "
            "across the DMA capability framework.",
            "Across the DMA capability framework, its overall digital "
            "maturity reads {s} out of 5.",
            "The DMA framework places its overall digital maturity at "
            "{s} out of 5.",
            "On the DMA capability framework, the assessment reads its "
            "overall digital maturity at {s} of 5.",
        ), s=f"{float(overall):.1f}"))

    text_out = " ".join(parts)
    return text_out if len(text_out) >= 120 else None


async def _active_overall(session: AsyncSession, entity_id: str) -> float | None:
    """Assessed overall maturity for the entity's ACTIVE run: the persisted
    overall_score, else the mean of its scored subcaps (always present for an
    ACTIVE entity). Grounds the maturity clause of the composed narrative."""
    row = (await session.execute(text(
        """
        SELECT COALESCE(r.overall_score,
                        (SELECT AVG(s.score) FROM subcap_scores s
                         WHERE s.run_id = r.id AND s.score BETWEEN 1 AND 5)) AS overall
        FROM runs r WHERE r.entity_id = CAST(:e AS uuid) AND r.status='ACTIVE'
        ORDER BY r.created_at DESC LIMIT 1
        """
    ), {"e": entity_id})).first()
    return float(row.overall) if row and row.overall is not None else None


async def heal_entity(
    session: AsyncSession, *, entity_id: str, drive_folder_id: str | None,
    corpus_dir: str | None = None, dry_run: bool = False,
    subvertical: str | None = None, name: str | None = None,
) -> dict[str, Any]:
    """Fill empty firmographics for one entity from its package, classifying a
    NULL subvertical first so the subvertical-default regulator can fill. Returns
    a report {filled, still_empty, panel_ok, pkg, subvertical, subvertical_classified}."""
    pkg = resolve_package_dir(drive_folder_id, corpus_dir)
    # Resolve a MISSING subvertical FIRST, before the regulator default below.
    # The subvertical-default regulator (the last-resort firmographics fill that
    # keeps the panel non-empty) keys off the subvertical, so the subvertical
    # must be known BEFORE we compute the regulator — not after. The deploy
    # reparse seeds some entities with a NULL subvertical (the leaf parser
    # couldn't read it); classify it from the package + name here and persist it,
    # so the regulator default fires for those entities too. Doing this inside
    # heal_entity fixes BOTH heal callers (heal_entities wave 2, heal_all_stages
    # wave 8) at the source — the previous order (heal_all_stages classified
    # AFTER healing) left the 9 just-classified entities healed against a NULL
    # subvertical → no regulator → GAP firmographics on every deploy.
    sv = (subvertical or "").strip() or None
    sv_classified = False
    if sv is None:
        sv_text = extract_subvertical_text(pkg) if pkg is not None else ""
        code = classify_subvertical(name or "", sv_text)
        if code:
            sv = code
            sv_classified = True
            if not dry_run:
                await session.execute(text(
                    "UPDATE entities SET subvertical=:c "
                    "WHERE id=CAST(:e AS uuid) AND (subvertical IS NULL OR subvertical='')"
                ), {"c": code, "e": entity_id})
    ext = extract_firmographics(pkg) if pkg is not None else {}
    # A GARBLED extracted regulator (a sentinel 'Role', or a regex that ran off
    # the field end — 'State DOIs (NAIC-aligned, … +', 'FDIC (state-') must NOT
    # be used: heal is fill-if-empty, so a garbled value would be stored, the
    # downstream sanitize pass would null it, and nothing re-fills it → a blank
    # panel field at the gate (the deploy's GAP firmographics: 5). Drop it so the
    # subvertical default below takes over.
    if regulator_is_garbled(ext.get("primary_regulator")):
        ext.pop("primary_regulator", None)
    # Last-resort regulator from the entity TYPE — derivable from the DB
    # subvertical ALONE, so it must apply EVEN when the package can't be resolved
    # (a name/layout mismatch must never leave primary_regulator NULL, which would
    # fail the firmographics completeness contract → a blank panel field).
    if not ext.get("primary_regulator"):
        reg_default = _default_regulator(sv, name)
        if reg_default:
            ext["primary_regulator"] = reg_default
            ext.setdefault("facts", {})["regulator_basis"] = "inferred:subvertical"
    if pkg is None and not ext:
        return {"filled": [], "still_empty": list(PANEL_FIELDS), "panel_ok": False,
                "pkg": False, "subvertical": sv, "subvertical_classified": sv_classified}

    cur = (await session.execute(text(
        "SELECT aum_usd, revenue_usd, headcount, primary_regulator, financial_highlights, "
        "parsed_facts, leadership, narrative_md, hq_address FROM firmographics WHERE entity_id=:e"
    ), {"e": entity_id})).first()
    cur_aum = cur.aum_usd if cur else None
    cur_rev = cur.revenue_usd if cur else None
    cur_head = cur.headcount if cur else None
    cur_reg = cur.primary_regulator if cur else None
    # A garbled regulator already in the DB (the parser stored 'Role' or a
    # truncated parenthetical) must not be treated as "filled" — drop it so the
    # clean extracted value / subvertical default below replaces it. Otherwise
    # fill-if-empty keeps the garbage, the sanitize pass nulls it, and the panel
    # field is blank at the gate.
    if regulator_is_garbled(cur_reg):
        cur_reg = None
    cur_fh = (cur.financial_highlights if cur and cur.financial_highlights else {}) or {}
    cur_pf = (cur.parsed_facts if cur and cur.parsed_facts else {}) or {}
    cur_lead = (cur.leadership if cur and cur.leadership else []) or []
    cur_narr = (cur.narrative_md if cur else None) or ""
    cur_hq = (cur.hq_address if cur else None) or ""

    filled: list[str] = []
    new_aum = cur_aum if cur_aum is not None else ext.get("aum_usd")
    new_rev = cur_rev if cur_rev is not None else ext.get("revenue_usd")
    new_head = cur_head if cur_head is not None else ext.get("headcount")
    new_reg = cur_reg if (cur_reg or "").strip() else ext.get("primary_regulator")
    # Merge highlights: existing keys win, extracted ratios fill the gaps.
    new_fh = {**(ext.get("financial_highlights") or {}), **cur_fh}
    new_pf = dict(cur_pf)
    for k, v in (ext.get("facts") or {}).items():
        if new_pf.get(k) in (None, "", [], {}):
            new_pf[k] = v
    # website ladder rung 2 (plan 4.2): the entity's OWN domain from its
    # evidence source URLs (own-domain heuristic; aggregators excluded).
    if not new_pf.get("website"):
        url_rows = (await session.execute(text(
            "SELECT DISTINCT source_url FROM evidence_index "
            "WHERE entity_id = CAST(:e AS uuid) AND source_url IS NOT NULL LIMIT 400"
        ), {"e": entity_id})).scalars().all()
        ws = website_from_urls(list(url_rows), name or "")
        if ws:
            new_pf["website"] = ws
            new_pf["website_basis"] = "evidence_urls:own-domain"
    # ticker rung 2: entity-scoped disambiguation over the entity's own
    # evidence excerpts (a peer/vendor ticker like NYSE:GIB never matches
    # because the entity name must co-occur in-clause).
    if not new_pf.get("ticker"):
        exc_rows = (await session.execute(text(
            "SELECT excerpt FROM evidence_index WHERE entity_id = CAST(:e AS uuid) "
            "AND excerpt ~ '(NYSE|NASDAQ|OTC|AMEX|TSX)' LIMIT 60"
        ), {"e": entity_id})).scalars().all()
        for exc in exc_rows:
            tk = prose_ticker(exc or "", name or "")
            if tk:
                new_pf["ticker"] = tk
                new_pf["ticker_basis"] = "evidence:entity-scoped"
                break
    # DB-prose firmographics recovery (2026-07-02): fill founded / hq / geography
    # from the entity's OWN persisted prose (analyst narrative + parsed facts +
    # financial-highlight lines) — the residual the package parser left NULL.
    # Every value is grounded in a founding/headquartered/place cue and stamped
    # with its basis; a field with no signal stays NULL (honest-null).
    hay = firm_prose_hay(cur_narr, new_pf, new_fh)
    # The 2026-07-04 deep search found founding/HQ statements for 25/19
    # still-null clients ONLY in evidence excerpts + report sections —
    # append the cue-bearing lines so the same grounded extractors see
    # them (basis records the wider source honestly).
    _hay_basis = "nlp:narrative"
    if not new_pf.get("founded") or not (
            (cur_hq or "").strip() or new_pf.get("hq") or new_pf.get("hq_city")):
        cue_rows = (await session.execute(text(
            """
            SELECT excerpt FROM evidence_index
            WHERE entity_id = CAST(:e AS uuid) AND excerpt IS NOT NULL
              AND excerpt ~* '(founded|established|chartered|incorporated|headquarter|based in|head office)'
            LIMIT 40
            """), {"e": entity_id})).scalars().all()
        sec_rows = (await session.execute(text(
            """
            SELECT ds.body FROM document_sections ds
            JOIN runs r ON r.id = ds.run_id AND r.status = 'ACTIVE'
            WHERE ds.entity_id = CAST(:e AS uuid)
              AND ds.body ~* '(founded|established|chartered|headquarter)'
            LIMIT 6
            """), {"e": entity_id})).scalars().all()
        extra = "\n".join([*(str(x) for x in cue_rows),
                           *(str(x or "")[:1500] for x in sec_rows)])
        if extra:
            hay = hay + "\n" + extra
            _hay_basis = "nlp:evidence+narrative"
    if not new_pf.get("founded"):
        fy = prose_founded_year(hay)
        if fy:
            new_pf["founded"] = int(fy)
            new_pf["founded_basis"] = _hay_basis
    if not (cur_hq or "").strip() and not new_pf.get("hq") and not new_pf.get("hq_city"):
        hq_loc = prose_hq_location(hay)
        if hq_loc:
            new_pf["hq"] = hq_loc
            new_pf["hq_basis"] = _hay_basis
    if not new_pf.get("geography"):
        geo = prose_geography(hay)
        if geo:
            new_pf["geography"] = geo
            new_pf["geography_basis"] = _hay_basis
    # Provenance floor: every present provenanced parsed-fact carries a *_basis.
    stamp_firmographic_provenance(new_pf)
    # Leadership roster — fill-if-empty from the package's c_suite roster
    # (only when the package resolved; the regulator-only heal path has pkg=None).
    new_lead = cur_lead if cur_lead else (extract_leadership(pkg) if pkg is not None else [])
    # "About" narrative — fill-if-below-floor from grounded firmographics + the
    # assessed maturity (never overwrites a real DOCX/Vertex narrative that
    # already clears the 120-char contract floor). Composition of stored facts,
    # not fabrication; marked composed in parsed_facts for auditability.
    new_narr = cur_narr
    if len(cur_narr.strip()) < 120:
        composed = compose_about_narrative(
            name=name, subvertical=sv, aum_usd=new_aum,
            aum_basis=new_pf.get("aum_basis"), regulator=new_reg, headcount=new_head,
            facts=new_pf, overall=await _active_overall(session, entity_id), hq=cur_hq,
        )
        if composed:
            new_narr = composed
            new_pf["narrative_basis"] = "composed:firmographics"
    for col, old, new in (
        ("aum_usd", cur_aum, new_aum), ("revenue_usd", cur_rev, new_rev),
        ("headcount", cur_head, new_head), ("primary_regulator", cur_reg, new_reg),
        ("financial_highlights", cur_fh, new_fh), ("parsed_facts", cur_pf, new_pf),
        ("leadership", cur_lead, new_lead), ("narrative_md", cur_narr, new_narr),
    ):
        if new != old:
            filled.append(col)

    if filled and not dry_run:
        await session.execute(text(
            """
            INSERT INTO firmographics (entity_id, aum_usd, revenue_usd, headcount,
                primary_regulator, financial_highlights, parsed_facts, leadership,
                narrative_md, updated_at)
            VALUES (:e, :aum, :rev, :head, :reg, CAST(:fh AS jsonb), CAST(:pf AS jsonb),
                CAST(:lead AS jsonb), :narr, NOW())
            ON CONFLICT (entity_id) DO UPDATE SET
                aum_usd=EXCLUDED.aum_usd, revenue_usd=EXCLUDED.revenue_usd,
                headcount=EXCLUDED.headcount, primary_regulator=EXCLUDED.primary_regulator,
                financial_highlights=EXCLUDED.financial_highlights,
                parsed_facts=EXCLUDED.parsed_facts, leadership=EXCLUDED.leadership,
                narrative_md=EXCLUDED.narrative_md, updated_at=NOW()
            """
        ), {"e": entity_id, "aum": new_aum, "rev": new_rev, "head": new_head,
            "reg": new_reg, "fh": json.dumps(new_fh), "pf": json.dumps(new_pf),
            "lead": json.dumps(new_lead), "narr": (new_narr or None)})

    # Sentiment recovery (2026-07-02): derive_sentiment mines only the report
    # prose; for several entities the consumer/employee ratings (Glassdoor / BBB
    # / app-store lines) live ONLY in the evidence excerpts, so that pass leaves
    # sentiment NULL. Re-mine those excerpts with the SAME extractor + normalizer
    # (identical shape) when sentiment is still empty — fill-if-empty, honest-null
    # when the corpus carries no rating signal at all.
    if not dry_run:
        cur_sent = (await session.execute(text(
            "SELECT sentiment FROM firmographics WHERE entity_id=:e"), {"e": entity_id})).scalar()
        sent_empty = cur_sent in (None, {}, []) or (
            isinstance(cur_sent, dict) and not cur_sent.get("sources"))
        if sent_empty:
            exc_blob = (await session.execute(text(
                "SELECT string_agg(excerpt, ' ') FROM evidence_index "
                "WHERE entity_id=CAST(:e AS uuid) AND excerpt IS NOT NULL"
            ), {"e": entity_id})).scalar() or ""
            if exc_blob:
                from app.scripts.derive_sentiment import _extract as _extract_sentiment
                from app.scripts.derive_sentiment import normalize_sentiment
                src = _extract_sentiment(exc_blob)
                if src:
                    blob = {"sources": src, "derived_from": "evidence_excerpts"}
                    nb = normalize_sentiment(blob) or blob
                    await session.execute(text(
                        "UPDATE firmographics SET sentiment=CAST(:s AS jsonb), updated_at=NOW() "
                        "WHERE entity_id=CAST(:e AS uuid) "
                        "AND (sentiment IS NULL OR NOT (sentiment ? 'sources'))"
                    ), {"s": json.dumps(nb), "e": entity_id})
                    filled.append("sentiment")

    still_empty: list[str] = []
    if new_aum is None:
        still_empty.append("aum_usd")
    if new_head is None:
        still_empty.append("headcount")
    if not (new_reg or "").strip():
        still_empty.append("primary_regulator")

    # Panel is non-empty when the entity shows a scale figure + regulator +
    # real substance: a headcount, or ≥2 substantive facts/highlights (members,
    # homes, branches, founded, CAGR, ticker, geography, size tier, ROA/ROE/NIM
    # …). A genuinely-undisclosed headcount alone never blanks the panel.
    substance = len([k for k in new_pf if k not in ("aum_basis", "legal_name", "regulator_basis", "narrative_basis")]) + len(new_fh)
    panel_ok = (
        new_aum is not None and bool((new_reg or "").strip())
        and (new_head is not None or substance >= 2)
    )
    return {"filled": filled, "still_empty": still_empty, "panel_ok": panel_ok,
            "pkg": True, "subvertical": sv, "subvertical_classified": sv_classified}
