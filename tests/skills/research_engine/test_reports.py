"""Two reports, curated from the workbook, that refuse rather than mislead."""
import re

import pytest
from docx import Document

from engine import ledger as L, report_spec as RS, reports
from engine.reports import ReportRefused

from fixtures import (CAT, bank_evidence, good_synthesis, new_run,  # noqa: F401
                      scored_run, section_record, sign_off_sections, synthesise)

def _narrate(wb, spec, *, words_per_section=None, cite=None, rec_cards=None,
             ic_cards=None):
    """Rows for every section of `spec`, in the PINNED template's shape —
    blocks, control-block checks, one card per pillar in scope, the Doc's
    minimum of REC cards — appended directly so each test can break exactly
    one thing the renderer must catch (a dead citation, a short section, a
    card floor). The writer's own path is tested in test_report_structure."""
    from engine import narrative as N
    eids = [cite] if cite else list(wb.evidence_index())[:10]
    for sec in spec.sections:
        if sec.kind == "pillar":
            cards = sorted({c[:2] for c in wb.selected_subcaps()})
        elif sec.is_card:
            n = rec_cards if rec_cards is not None else N.card_floor_for(wb, sec)
            cards = [f"{sec.card_prefix}{i + 1:02d}" for i in range(n)]
        else:
            cards = [None]
        for card in cards:
            rec = section_record(sec.id, eids, report=spec.key)
            body = rec["Body"]
            if words_per_section is not None:
                body = " ".join(body.split()[:words_per_section])
            if ic_cards is not None and spec.key == "client_research" and sec.id == "5":
                body = re.sub(r"\bIC-(\d{3})\b",
                              lambda m: m.group(0) if int(m.group(1)) <= ic_cards else "",
                              body)
            row = {"Report": spec.key, "Section_ID": sec.id,
                   "Heading": (f"{card}: {sec.heading}" if card else sec.heading),
                   "Body": body, "Evidence_IDs": ", ".join(eids), "Kind": sec.kind,
                   "Card_ID": card or "", "Author": "test",
                   "Written_At": "2026-08-29T00:00:00Z"}
            for k in ("Weighing", "Assumptions", "Bias_Notes", "Inference_Tags",
                      "Absence_Basis"):
                row[k] = rec.get(k, "")
            wb.append("Report_Narrative", row, save=False)
    wb.save()


def _run_with_content(tmp_path, n=6):
    """A researched, gated and SCORED run — the assessment report's §3 reads
    Cap_Triggers / Subcap_Scores / Caps_Applied_Log, which only exist once
    the scoring stage has run; a research-only run has no source for it."""
    run, wb, cells, ev = scored_run(tmp_path, n=n)
    return run, wb


# ── AUD-0003 · both artefacts have a producer, and an ingestible name ─────

def test_both_reports_render_to_docx(tmp_path):
    run, wb = _run_with_content(tmp_path)
    out = []
    for spec in RS.SPECS.values():
        _narrate(wb, spec)
        sign_off_sections(wb)
        out.append(reports.render(wb, spec, run.deliverables))
    assert all(r["path"].endswith(".docx") for r in out)
    assert all(Document(r["path"]).paragraphs for r in out)


def test_the_filenames_are_the_ones_the_app_classifies(tmp_path):
    """AUD-0003: the produced report was `client_profile.md`, and
    classify() returns None for it, so it was uningestable."""
    import sys
    sys.path.insert(0, str(__import__("pathlib").Path(__file__)
                           .resolve().parents[3] / "apps" / "worker"))
    from dma_worker.classification import classify
    run, wb = _run_with_content(tmp_path)
    for spec in RS.SPECS.values():
        _narrate(wb, spec)
        sign_off_sections(wb)
        r = reports.render(wb, spec, run.deliverables)
        c = classify(r["path"].rsplit("/", 1)[-1])
        assert c is not None, f"{r['path']} is unclassifiable"
    assert classify("client_profile.md") is None   # the old output


# ── AUD-0033 · a dead citation refuses the render ────────────────────────

def test_one_unresolvable_citation_refuses_the_whole_report(tmp_path):
    run, wb = _run_with_content(tmp_path)
    spec = RS.SPECS["assessment"]
    _narrate(wb, spec, cite="E-999")
    with pytest.raises(ReportRefused, match="do not resolve"):
        reports.render(wb, spec, run.deliverables)


def test_the_citation_list_at_the_back_resolves_to_the_register(tmp_path):
    run, wb = _run_with_content(tmp_path)
    spec = RS.SPECS["assessment"]
    _narrate(wb, spec)
    sign_off_sections(wb)
    r = reports.render(wb, spec, run.deliverables)
    assert r["unresolved"] == []
    text = "\n".join(p.text for p in Document(r["path"]).paragraphs)
    for e in set(re.findall(r"\[(E-\d+)\]", text)):
        assert f"[{e}]" in text and "https://" in text


# ── AUD-0105 · the word minimums are measured ────────────────────────────

def test_a_short_section_refuses_and_says_which(tmp_path):
    run, wb = _run_with_content(tmp_path)
    spec = RS.SPECS["assessment"]
    _narrate(wb, spec, words_per_section=30)
    with pytest.raises(ReportRefused) as e:
        reports.render(wb, spec, run.deliverables)
    assert "blocking minimum" in str(e.value)
    assert "Report_Narrative" in str(e.value)


# ── AUD-0145 · the card floors are the template's, not 3 ─────────────────

def test_four_recommendation_cards_is_not_enough(tmp_path):
    """The Doc's §8 carries five to eight REC-NN cards."""
    run, wb = _run_with_content(tmp_path)
    spec = RS.SPECS["assessment"]
    _narrate(wb, spec, rec_cards=4)
    with pytest.raises(ReportRefused, match=r"recommendation\(s\) against the"):
        reports.render(wb, spec, run.deliverables)


def test_three_insight_cards_the_old_floor_is_refused(tmp_path):
    """The profile's insight cards are IC-NNN ids inside §5, eight or more;
    the old floor of three is refused by the countable check."""
    run, wb = _run_with_content(tmp_path)
    spec = RS.SPECS["client_research"]
    _narrate(wb, spec, ic_cards=3)
    with pytest.raises(ReportRefused, match="insight cards"):
        reports.render(wb, spec, run.deliverables)


# ── AUD-0052 · the numbers in the report ARE the numbers in the sheets ───

def test_the_coverage_table_is_read_from_the_workbook_at_render_time(tmp_path):
    run, wb = _run_with_content(tmp_path)
    spec = RS.SPECS["assessment"]
    _narrate(wb, spec)
    sign_off_sections(wb)
    r = reports.render(wb, spec, run.deliverables)
    doc = Document(r["path"])
    cov = wb.rows("Coverage_Map")[0]
    from engine import contract as C
    seen = [t for t in doc.tables
            if [c.text for c in t.rows[0].cells] == list(C.COVERAGE_MAP_COLUMNS)]
    assert seen, "the coverage table must be in the document"
    body = [c.text for row in seen[0].rows for c in row.cells]
    assert str(cov["subcaps"]) in body and cov["category_id"] in body


def test_changing_the_workbook_changes_the_report(tmp_path):
    """The property that makes 'one substrate' true rather than asserted."""
    run, wb = _run_with_content(tmp_path)
    spec = RS.SPECS["assessment"]
    _narrate(wb, spec)
    sign_off_sections(wb)
    before = reports.render(wb, spec, run.deliverables)["citations"]
    L.append_evidence(wb, source_name="NCUA call report 2025",
                      source_url="https://ncua.example/cr25", tier="T1",
                      excerpt="Digital channel volumes for Acme Credit Union "
                              "rose 31 percent year over year in 2025 filings.",
                      subcaps=[wb.selected_subcaps()[0]], published="2025-09-30")
    r = reports.render(wb, spec, run.deliverables)
    doc = Document(r["path"])
    text = "\n".join(p.text for p in doc.paragraphs) + "".join(
        c.text for t in doc.tables for row in t.rows for c in row.cells)
    assert "NCUA call report 2025" in text
    assert before == r["citations"]  # the narrative's citations, unchanged


# ── AUD-0107 · a section with no source says so, it does not render empty ─

def test_a_section_whose_inputs_are_all_empty_is_named_in_the_refusal(tmp_path):
    """AUD-0107: thirteen sheets named as INPUTS by the templates do not
    exist, 'leaving §3.2 and §3.3 with no source at all'. A section whose
    every input is empty must say so."""
    # prelim=False: PRELIM fills Entity_Timeline, and this test is about a
    # section whose every named input is genuinely empty.
    run = new_run(tmp_path, n=2, prelim=False)
    wb = run.open()
    spec = RS.from_json({
        "key": "assessment", "title": "Probe", "min_words": 10,
        "filename": "DMA_Assessment_Report_{entity}_{date}.docx",
        "sections": [{"id": "9", "heading": "Acquisition history",
                      "min_words": 10, "inputs": ["Entity_Timeline"],
                      "requires_citation": False}],
    })
    wb.append("Report_Narrative", {
        "Report": "assessment", "Section_ID": "9", "Heading": "Acquisition",
        "Body": "Acme Credit Union completed no acquisitions in the period "
                "under review, on the evidence gathered.",
        "Kind": "section"})
    with pytest.raises(ReportRefused) as e:
        reports.render(wb, spec, run.deliverables)
    assert "no source at all" in str(e.value)
    assert "Entity_Timeline" in str(e.value)


def test_a_focused_engagement_states_its_scope_instead_of_refusing(tmp_path):
    """The opposite error to AUD-0107's: a run scoped to one pillar
    legitimately leaves three pillar sheets empty, and refusing that would
    reject a correct run."""
    run, wb = _run_with_content(tmp_path)
    spec = RS.SPECS["assessment"]
    _narrate(wb, spec)
    sign_off_sections(wb)
    r = reports.render(wb, spec, run.deliverables)
    text = "\n".join(p.text for p in Document(r["path"]).paragraphs)
    assert "Not in this engagement's scope" in text
    assert "P4_Subcap_Scoring" in text


def test_forcing_marks_the_gaps_in_the_document_rather_than_hiding_them(tmp_path):
    run = new_run(tmp_path, n=2)
    wb = run.open()
    spec = RS.SPECS["assessment"]
    r = reports.render(wb, spec, run.deliverables, force=True)
    assert r["forced"] and r["problems"]
    text = "\n".join(p.text for p in Document(r["path"]).paragraphs)
    assert "NO SOURCE" in text
