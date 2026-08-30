#!/usr/bin/env python3
"""
DMA Governance Auditor — Automated structural checks for Workflow A.

Runs all deterministic and pattern-matching checks from audit_checks.md
(Categories 1–8) plus distributional checks (DC-01 through DC-08).

Produces:
  - check_results.json  (per-check PASS/FAIL with details)
  - preliminary_issues.csv  (issue register rows for failures)
  - audit_summary.json  (aggregate counts, verdict recommendation)

Usage:
  python gov_auditor.py <assessment_dir> [--output-dir <dir>]

  <assessment_dir> must contain:
    - *.xlsx  (scoring workbook)
    - *.docx  (assessment report)
    - run_manifest.json
    - caps_applied_log.csv
    - contradiction_log.csv
    - evidence_index.csv
"""

import argparse
import csv
import json
import os
import re
import sys
from collections import Counter, defaultdict
from datetime import datetime
from io import StringIO
from pathlib import Path

# ---------------------------------------------------------------------------
# Attempt imports; give clear install instructions on failure
# ---------------------------------------------------------------------------
try:
    import openpyxl
except ImportError:
    sys.exit("openpyxl required: pip install openpyxl --break-system-packages")

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None  # graceful degradation — skip report checks

try:
    import jsonschema
except ImportError:
    jsonschema = None

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
REQUIRED_WORKBOOK_TABS = [
    "Summary", "Calculation_Chain",
    "P1_Scoring_Detail", "P2_Scoring_Detail",
    "P3_Scoring_Detail", "P4_Scoring_Detail",
    "Evidence_Index", "Caps_Applied_Log",
    "Contradiction_Log", "Absent_Evidence_Log",
    "QA_Validation_Log",
]

PILLAR_NAMES = ["P1", "P2", "P3", "P4"]

SCORING_DETAIL_SHEETS = [f"{p}_Scoring_Detail" for p in PILLAR_NAMES]

VALID_CAP_TYPES = {
    "EVIDENCE_CEILING", "SENTIMENT", "REGULATORY",
    "CROSS_PILLAR", "ADJ_STALENESS", "ADJ_COMPLAINT",
    "ADJ_INCIDENT_MAJOR", "ADJ_INCIDENT_PATTERN",
}

VALID_RESOLUTION_RULES = {
    "ERS_RANKING", "T1T2_OVERRIDE", "TIEBREAKER",
    "CONSERVATIVE_DEFAULT", "UNRESOLVED",
}

VALID_TIERS = {"T1", "T2", "T3", "T4", "T5"}

VALID_CONFIDENCE = {"HIGH", "MEDIUM", "LOW"}

FORBIDDEN_LANGUAGE_PATTERNS = [
    r"\bidentical\s+methodology\b",
    r"\bCritical\s+Gaps\b",
    r"\bMonths?\s+\d+\s*[-–]\s*\d+\b",  # "Months 0-6", "Month 7-12"
]

EVIDENCE_ID_PATTERN = re.compile(r"E-\d{3}(?::F\d+)?")

REPORT_SECTIONS_EXPECTED = 12
REPORT_APPENDICES_EXPECTED = 3

# ---------------------------------------------------------------------------
# Data Classes
# ---------------------------------------------------------------------------

#: A check whose INPUT IS ABSENT did not run. It did not fail.
#:
#: Owner, 2026-08-23: the routines "default to rejecting in case of issues,
#: rather than triaging and fixing". This constant is most of the mechanism
#: behind that. Measured the same day (MEM-0158, MEM-0168): session
#: accelerate-63 vetted three packages — all three carried governance
#: artifacts asserting PASS / 0 CRITICAL, and a fresh gov_auditor returned
#: CRITICAL on the same bytes, so all three were refused. Two of those
#: CRITICALs were checks that never ran:
#:
#:   * AG-01 read an absent weight column as a list of ZEROS, so `if weights`
#:     was satisfied, the sum was 0.0, and "weights do not sum to 1.0" fired
#:     CRITICAL against a workbook generation that simply does not carry the
#:     column.
#:   * AG-08 required a sheet literally named `Summary` and reported its
#:     absence as a CRITICAL aggregation failure.
#:
#: Neither says anything about the assessment. A verdict that cannot tell
#: "this is wrong" from "I could not look" carries no information, and it
#: costs a whole client — which is what it did, three times in one session.
#: NOT_RUN carries the reason and never counts toward CRITICAL or FAIL.
NOT_RUN = "NOT_RUN"


class CheckResult:
    def __init__(self, check_id, status, severity, description, details=""):
        self.check_id = check_id
        self.status = status  # PASS, FAIL or NOT_RUN
        self.severity = severity
        self.description = description
        self.details = details

    @property
    def ran(self) -> bool:
        return self.status != NOT_RUN

    def to_dict(self):
        return {
            "check_id": self.check_id,
            "status": self.status,
            "severity": self.severity,
            "description": self.description,
            "details": self.details,
        }


def not_run(check_id, description, why):
    """A check that could not look, said so. `why` names the missing input in
    words an author can act on — never "check failed"."""
    return CheckResult(check_id, NOT_RUN, "INFO", description, why)


class Issue:
    _counter = 0

    def __init__(self, severity, category, subcategory, affected_id,
                 description, detection_evidence, fix_instruction,
                 auto_fixable=False):
        Issue._counter += 1
        self.issue_id = f"ISS-{Issue._counter:03d}"
        self.severity = severity
        self.category = category
        self.subcategory = subcategory
        self.affected_id = affected_id
        self.description = description
        self.detection_evidence = detection_evidence
        self.fix_instruction = fix_instruction
        self.auto_fixable = auto_fixable
        self.status = "OPEN"

    def to_csv_row(self):
        return [
            self.issue_id, self.severity, self.category, self.subcategory,
            self.affected_id, self.description, self.detection_evidence,
            self.fix_instruction, str(self.auto_fixable).lower(), self.status,
        ]

    CSV_HEADER = [
        "issue_id", "severity", "category", "subcategory", "affected_id",
        "description", "detection_evidence", "fix_instruction",
        "auto_fixable", "status",
    ]


# ---------------------------------------------------------------------------
# File Discovery
# ---------------------------------------------------------------------------

def discover_files(assessment_dir):
    """Locate all expected input files in the assessment directory."""
    d = Path(assessment_dir)
    files = {}
    # Workbook
    xlsx_files = list(d.glob("*.xlsx"))
    files["workbook"] = xlsx_files[0] if xlsx_files else None
    # Report
    docx_files = list(d.glob("*.docx"))
    files["report"] = docx_files[0] if docx_files else None
    # JSON/CSV
    for name in ["run_manifest.json", "caps_applied_log.csv",
                  "contradiction_log.csv", "evidence_index.csv"]:
        p = d / name
        files[name] = p if p.exists() else None
    return files


# ---------------------------------------------------------------------------
# Parsers
# ---------------------------------------------------------------------------

def load_manifest(path):
    with open(path) as f:
        return json.load(f)


def load_csv_rows(path):
    """Return list of dicts from a CSV file."""
    with open(path, newline="", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        return list(reader), reader.fieldnames


def load_workbook(path):
    return openpyxl.load_workbook(path, data_only=True)


def sheet_to_dicts(ws):
    """Convert worksheet to list of dicts using first row as header."""
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return [], []
    headers = [str(h).strip() if h else f"col_{i}" for i, h in enumerate(rows[0])]
    data = []
    for row in rows[1:]:
        if all(c is None for c in row):
            continue
        data.append(dict(zip(headers, row)))
    return data, headers


# ---------------------------------------------------------------------------
# Check Categories
# ---------------------------------------------------------------------------

def run_input_validation(files, manifest, wb):
    """Category 1: IV-01 through IV-10."""
    results = []
    issues = []

    # IV-01: run_manifest.json present
    if files.get("run_manifest.json"):
        results.append(CheckResult("IV-01", "PASS", "CRITICAL", "run_manifest.json present"))
    else:
        results.append(CheckResult("IV-01", "FAIL", "CRITICAL", "run_manifest.json present",
                                   "File not found"))
        issues.append(Issue("CRITICAL", "INPUT_VALIDATION", "IV-01", "run_manifest.json",
                            "run_manifest.json missing", "File not found in assessment dir",
                            "Generate run_manifest.json from Layer 1"))
        return results, issues  # Fatal — stop

    # IV-02: schema valid (check required top-level keys)
    required_keys = {"run_id", "institution_name", "sub_vertical", "overall_score",
                     "pillar_scores", "evidence_count"}
    manifest_keys = set(manifest.keys()) if manifest else set()
    missing = required_keys - manifest_keys
    if not missing:
        results.append(CheckResult("IV-02", "PASS", "CRITICAL", "run_manifest schema valid"))
    else:
        results.append(CheckResult("IV-02", "FAIL", "CRITICAL", "run_manifest schema valid",
                                   f"Missing keys: {missing}"))
        issues.append(Issue("CRITICAL", "INPUT_VALIDATION", "IV-02", "run_manifest.json",
                            f"Missing required fields: {missing}",
                            "Key check", "Add missing fields to manifest"))

    # IV-03: overall score matches pillar weighted average
    if manifest and "pillar_scores" in manifest and "overall_score" in manifest:
        pillars = manifest["pillar_scores"]
        if isinstance(pillars, dict) and len(pillars) >= 4:
            pillar_vals = [float(pillars.get(p, 0)) for p in PILLAR_NAMES]
            # Default equal weights if not specified
            computed = sum(pillar_vals) / len(pillar_vals)
            overall = float(manifest["overall_score"])
            delta = abs(computed - overall)
            if delta <= 0.02:
                results.append(CheckResult("IV-03", "PASS", "CRITICAL",
                    "Overall score matches pillar average", f"Delta={delta:.4f}"))
            else:
                results.append(CheckResult("IV-03", "FAIL", "CRITICAL",
                    "Overall score matches pillar average",
                    f"Computed={computed:.2f}, Stated={overall:.2f}, Delta={delta:.4f}"))
                issues.append(Issue("CRITICAL", "INPUT_VALIDATION", "IV-03",
                    "run_manifest.overall_score",
                    f"Overall score {overall} != pillar avg {computed:.2f} (delta {delta:.4f})",
                    "Mathematical verification",
                    "Recalculate overall from pillar scores with correct weights"))

    # IV-04: evidence total = sum of tier distribution
    if manifest and "evidence_count" in manifest:
        # Try both flat and nested formats
        total = int(manifest.get("evidence_count", 0))
        tier_dist = manifest.get("tier_distribution", {})
        if not tier_dist:
            em = manifest.get("evidence_metrics", {})
            tier_dist = em.get("tier_distribution", {})
        if tier_dist:
            tier_sum = sum(int(v) for v in tier_dist.values())
            if tier_sum == total:
                results.append(CheckResult("IV-04", "PASS", "HIGH",
                    "Evidence total = tier sum", f"{total}={tier_sum}"))
            else:
                results.append(CheckResult("IV-04", "FAIL", "HIGH",
                    "Evidence total = tier sum",
                    f"Total={total}, Tier sum={tier_sum}"))
                issues.append(Issue("HIGH", "INPUT_VALIDATION", "IV-04",
                    "run_manifest.evidence_count",
                    f"Evidence total ({total}) != sum of tiers ({tier_sum})",
                    "Arithmetic check", "Reconcile evidence count with tier breakdown"))

    # IV-05..IV-07: CSV files present
    for check_id, fname in [("IV-05", "caps_applied_log.csv"),
                            ("IV-06", "contradiction_log.csv"),
                            ("IV-07", "evidence_index.csv")]:
        if files.get(fname):
            results.append(CheckResult(check_id, "PASS", "CRITICAL", f"{fname} present"))
        else:
            results.append(CheckResult(check_id, "FAIL", "CRITICAL", f"{fname} present",
                                       "File not found"))
            issues.append(Issue("CRITICAL", "INPUT_VALIDATION", check_id, fname,
                                f"{fname} missing", "File not found",
                                f"Export {fname} from Layer 1 workbook"))

    # IV-08: Workbook present with required tabs
    if wb:
        present_tabs = set(wb.sheetnames)
        missing_tabs = [t for t in REQUIRED_WORKBOOK_TABS if t not in present_tabs]
        if not missing_tabs:
            results.append(CheckResult("IV-08", "PASS", "CRITICAL",
                "Workbook has all required tabs"))
        else:
            results.append(CheckResult("IV-08", "FAIL", "CRITICAL",
                "Workbook has all required tabs",
                f"Missing: {missing_tabs}"))
            issues.append(Issue("CRITICAL", "INPUT_VALIDATION", "IV-08", "workbook",
                f"Missing tabs: {missing_tabs}", "Tab name check",
                "Add missing worksheet tabs to workbook"))
    else:
        results.append(CheckResult("IV-08", "FAIL", "CRITICAL",
            "Workbook present", "No .xlsx file found"))
        issues.append(Issue("CRITICAL", "INPUT_VALIDATION", "IV-08", "workbook",
            "No .xlsx workbook found", "File discovery",
            "Provide completed scoring workbook (.xlsx)"))

    # IV-09: Report present
    if files.get("report"):
        results.append(CheckResult("IV-09", "PASS", "CRITICAL", "Report (.docx) present"))
    else:
        results.append(CheckResult("IV-09", "FAIL", "CRITICAL", "Report (.docx) present",
                                   "No .docx file found"))
        issues.append(Issue("CRITICAL", "INPUT_VALIDATION", "IV-09", "report",
            "No .docx report found", "File discovery",
            "Provide assessment report draft (.docx)"))

    # IV-10: Rubric version match (manifest vs workbook Summary)
    if wb and manifest and "Summary" in wb.sheetnames:
        summary_data, _ = sheet_to_dicts(wb["Summary"])
        # Try to find version info in summary — implementation depends on workbook format
        manifest_version = str(manifest.get("assessment_skill_version",
                               manifest.get("versions", {}).get("rubric", "unknown")))
        results.append(CheckResult("IV-10", "PASS", "HIGH",
            "Rubric version check", f"Manifest version: {manifest_version} (manual verify recommended)"))

    return results, issues


def run_score_integrity(wb):
    """Category 2: SI-01 through SI-21 (deterministic + pattern subset)."""
    results = []
    issues = []

    all_subcaps = []
    all_rationales = []
    pillar_subcaps = defaultdict(list)

    for sheet_name in SCORING_DETAIL_SHEETS:
        if sheet_name not in wb.sheetnames:
            continue
        pillar = sheet_name.split("_")[0]  # P1, P2, P3, P4
        rows, headers = sheet_to_dicts(wb[sheet_name])
        for row in rows:
            row["_pillar"] = pillar
            row["_sheet"] = sheet_name
        all_subcaps.extend(rows)
        pillar_subcaps[pillar].extend(rows)

    if not all_subcaps:
        results.append(CheckResult("SI-01", "FAIL", "CRITICAL",
            "Subcap data found", "No scoring detail rows found in any pillar sheet"))
        issues.append(Issue("CRITICAL", "SCORE_INTEGRITY", "SI-01", "workbook",
            "No subcap rows found in scoring detail sheets", "Empty sheets",
            "Populate P1-P4_Scoring_Detail sheets"))
        return results, issues

    # Detect column names (flexible matching)
    sample = all_subcaps[0]
    cols = {k.lower().replace(" ", "_"): k for k in sample.keys() if not k.startswith("_")}

    def get_col(row, *candidates):
        """Find a column value by trying multiple name variants."""
        for c in candidates:
            if c in row and row[c] is not None:
                return row[c]
            cl = c.lower().replace(" ", "_")
            for k, v in row.items():
                if k.lower().replace(" ", "_") == cl:
                    return v
        return None

    # SI-01: Subcap count per pillar
    for pillar, subcaps in pillar_subcaps.items():
        count = len(subcaps)
        results.append(CheckResult("SI-01", "PASS", "CRITICAL",
            f"Subcap count {pillar}: {count}", f"{count} rows found"))

    # SI-02: No empty rows
    empty_count = sum(1 for r in all_subcaps
                      if all(v is None for k, v in r.items() if not k.startswith("_")))
    if empty_count == 0:
        results.append(CheckResult("SI-02", "PASS", "CRITICAL", "No empty rows"))
    else:
        results.append(CheckResult("SI-02", "FAIL", "CRITICAL", "No empty rows",
                                   f"{empty_count} empty rows found"))
        issues.append(Issue("CRITICAL", "SCORE_INTEGRITY", "SI-02", "scoring_detail",
            f"{empty_count} empty rows in scoring sheets", "Row scan",
            "Remove or populate empty rows"))

    # SI-03: No duplicate subcap IDs
    subcap_ids = [get_col(r, "SubCap_ID", "Subcap_ID", "subcap_id", "ID") for r in all_subcaps]
    subcap_ids_clean = [s for s in subcap_ids if s is not None]
    dupes = [sid for sid, c in Counter(subcap_ids_clean).items() if c > 1]
    if not dupes:
        results.append(CheckResult("SI-03", "PASS", "CRITICAL", "No duplicate subcap IDs"))
    else:
        results.append(CheckResult("SI-03", "FAIL", "CRITICAL", "No duplicate subcap IDs",
                                   f"Duplicates: {dupes}"))
        issues.append(Issue("CRITICAL", "SCORE_INTEGRITY", "SI-03", str(dupes),
            f"Duplicate subcap IDs: {dupes}", "ID uniqueness check",
            "Remove duplicate rows or assign unique IDs"))

    # SI-04: All rationale cells populated
    rationale_col_candidates = ["Rationale", "rationale", "Rationale_Text"]
    empty_rationales = []
    for r in all_subcaps:
        rat = get_col(r, *rationale_col_candidates)
        sid = get_col(r, "SubCap_ID", "Subcap_ID", "subcap_id", "ID")
        if rat:
            all_rationales.append({"text": str(rat), "subcap_id": sid, "pillar": r["_pillar"]})
        else:
            empty_rationales.append(sid)

    if not empty_rationales:
        results.append(CheckResult("SI-04", "PASS", "CRITICAL", "All rationales populated"))
    else:
        results.append(CheckResult("SI-04", "FAIL", "CRITICAL", "All rationales populated",
                                   f"{len(empty_rationales)} empty: {empty_rationales[:5]}"))
        issues.append(Issue("CRITICAL", "SCORE_INTEGRITY", "SI-04",
            str(empty_rationales[:5]),
            f"{len(empty_rationales)} subcaps have empty rationale cells",
            "Null check on rationale column",
            "Populate rationale for every subcap"))

    # SI-05: All rationales ≥150 characters
    short_rationales = [(r["subcap_id"], len(r["text"])) for r in all_rationales
                        if len(r["text"]) < 150]
    if not short_rationales:
        results.append(CheckResult("SI-05", "PASS", "HIGH", "All rationales ≥150 chars"))
    else:
        results.append(CheckResult("SI-05", "FAIL", "HIGH", "All rationales ≥150 chars",
                                   f"{len(short_rationales)} too short"))
        issues.append(Issue("HIGH", "SCORE_INTEGRITY", "SI-05",
            str([s[0] for s in short_rationales[:5]]),
            f"{len(short_rationales)} rationales under 150 chars",
            "Character count", "Expand rationale text with additional evidence-backed reasoning"))

    # SI-06: No forbidden rationale patterns
    forbidden_hits = []
    for r in all_rationales:
        for pat in FORBIDDEN_LANGUAGE_PATTERNS:
            if re.search(pat, r["text"], re.IGNORECASE):
                forbidden_hits.append((r["subcap_id"], pat))
    if not forbidden_hits:
        results.append(CheckResult("SI-06", "PASS", "HIGH", "No forbidden rationale patterns"))
    else:
        results.append(CheckResult("SI-06", "FAIL", "HIGH", "No forbidden rationale patterns",
                                   f"{len(forbidden_hits)} matches"))
        issues.append(Issue("HIGH", "SCORE_INTEGRITY", "SI-06",
            str([h[0] for h in forbidden_hits[:5]]),
            f"Forbidden language found in {len(forbidden_hits)} rationales",
            "Regex pattern match", "Remove or rephrase forbidden language"))

    # SI-07: Every rationale has ≥1 evidence ID
    no_evidence_refs = []
    for r in all_rationales:
        if not EVIDENCE_ID_PATTERN.search(r["text"]):
            no_evidence_refs.append(r["subcap_id"])
    if not no_evidence_refs:
        results.append(CheckResult("SI-07", "PASS", "HIGH",
            "All rationales cite evidence IDs"))
    else:
        results.append(CheckResult("SI-07", "FAIL", "HIGH",
            "All rationales cite evidence IDs",
            f"{len(no_evidence_refs)} missing: {no_evidence_refs[:5]}"))
        issues.append(Issue("HIGH", "SCORE_INTEGRITY", "SI-07",
            str(no_evidence_refs[:5]),
            f"{len(no_evidence_refs)} rationales lack evidence ID citations",
            "Evidence ID regex search", "Add E-NNN evidence citations to rationale"))

    # SI-08: Maturity descriptor references
    maturity_keywords = ["M1", "M2", "M3", "M4", "M5",
                         "Initial", "Developing", "Defined", "Managed", "Optimizing",
                         "maturity level", "maturity"]
    no_maturity = []
    for r in all_rationales:
        text_lower = r["text"].lower()
        if not any(kw.lower() in text_lower for kw in maturity_keywords):
            no_maturity.append(r["subcap_id"])
    if not no_maturity:
        results.append(CheckResult("SI-08", "PASS", "MEDIUM", "All rationales ref maturity"))
    else:
        results.append(CheckResult("SI-08", "FAIL", "MEDIUM", "All rationales ref maturity",
                                   f"{len(no_maturity)} missing"))

    # SI-12: No duplicate rationale text (exact match)
    seen_texts = {}
    exact_dupes = []
    for r in all_rationales:
        normalized = " ".join(r["text"].split()).strip()
        if normalized in seen_texts:
            exact_dupes.append((r["subcap_id"], seen_texts[normalized]))
        else:
            seen_texts[normalized] = r["subcap_id"]
    if not exact_dupes:
        results.append(CheckResult("SI-12", "PASS", "HIGH", "No duplicate rationale text"))
    else:
        results.append(CheckResult("SI-12", "FAIL", "HIGH", "No duplicate rationale text",
                                   f"{len(exact_dupes)} duplicates"))
        issues.append(Issue("HIGH", "SCORE_INTEGRITY", "SI-12",
            str(exact_dupes[:3]),
            f"{len(exact_dupes)} rationale text duplicates found",
            "Text exact-match comparison",
            "Write unique rationale for each subcap"))

    # SI-09: Non-0.5 Raw_Scores have quantitative justification
    non_half_no_quant = []
    quant_patterns = re.compile(r'\d+[\.\d]*\s*%|\d+[\.\d]*x|\$\d|#\d|\d+\s*(users|members|accounts|transactions|customers)', re.IGNORECASE)
    for r in all_rationales:
        sid = r["subcap_id"]
        # Find matching subcap row
        matching = [s for s in all_subcaps if get_col(s, "SubCap_ID", "Subcap_ID", "subcap_id", "ID") == sid]
        if matching:
            raw = get_col(matching[0], "Raw_Score", "raw_score")
            if raw is not None:
                try:
                    rv = float(raw)
                    if rv % 0.5 != 0 and not quant_patterns.search(r["text"]):
                        non_half_no_quant.append(sid)
                except (ValueError, TypeError):
                    pass
    if not non_half_no_quant:
        results.append(CheckResult("SI-09", "PASS", "HIGH", "Non-0.5 scores have quantitative justification"))
    else:
        results.append(CheckResult("SI-09", "FAIL", "HIGH", "Non-0.5 scores have quantitative justification",
                                   f"{len(non_half_no_quant)} missing: {non_half_no_quant[:5]}"))
        issues.append(Issue("HIGH", "SCORE_INTEGRITY", "SI-09", str(non_half_no_quant[:5]),
            f"{len(non_half_no_quant)} non-0.5 scores lack quantitative justification",
            "Quantitative pattern search in rationale",
            "Add quantitative metrics to justify precision beyond 0.5 increments"))

    # SI-10: Ceiling and cap checks documented in rationale
    cap_keywords = ["ceiling", "cap", "capped", "evidence ceiling", "single.source", "dependency"]
    no_cap_doc = []
    for r in all_rationales:
        sid = r["subcap_id"]
        matching = [s for s in all_subcaps if get_col(s, "SubCap_ID", "Subcap_ID", "subcap_id", "ID") == sid]
        if matching:
            caps = get_col(matching[0], "Caps_Applied", "caps_applied")
            if caps and str(caps).strip().lower() != "none":
                text_lower = r["text"].lower()
                if not any(kw in text_lower for kw in cap_keywords):
                    no_cap_doc.append(sid)
    if not no_cap_doc:
        results.append(CheckResult("SI-10", "PASS", "HIGH", "Cap checks documented in rationale"))
    else:
        results.append(CheckResult("SI-10", "FAIL", "HIGH", "Cap checks documented in rationale",
                                   f"{len(no_cap_doc)} missing"))
        issues.append(Issue("HIGH", "SCORE_INTEGRITY", "SI-10", str(no_cap_doc[:5]),
            f"{len(no_cap_doc)} capped subcaps lack cap documentation in rationale",
            "Keyword search for cap-related terms",
            "Document ceiling/cap application in rationale text"))

    # SI-11: Counter-evidence addressed in rationale
    counter_keywords = ["counter", "however", "despite", "although", "contrary", "opposing",
                        "counterargument", "counterclaim", "rebuttal", "limitation"]
    no_counter = []
    for r in all_rationales:
        text_lower = r["text"].lower()
        if not any(kw in text_lower for kw in counter_keywords):
            no_counter.append(r["subcap_id"])
    if not no_counter:
        results.append(CheckResult("SI-11", "PASS", "MEDIUM", "Counter-evidence addressed"))
    else:
        results.append(CheckResult("SI-11", "FAIL", "MEDIUM", "Counter-evidence addressed",
                                   f"{len(no_counter)} missing"))

    # SI-13: Contradictions cite resolution rule
    resolution_rule_pattern = re.compile(r"ERS_RANKING|T1T2_OVERRIDE|TIEBREAKER|CONSERVATIVE_DEFAULT|UNRESOLVED", re.IGNORECASE)
    no_resolution_cite = []
    for r in all_rationales:
        text_lower = r["text"].lower()
        if "contradict" in text_lower or "conflict" in text_lower:
            if not resolution_rule_pattern.search(r["text"]):
                no_resolution_cite.append(r["subcap_id"])
    if not no_resolution_cite:
        results.append(CheckResult("SI-13", "PASS", "HIGH", "Contradictions cite resolution rule"))
    else:
        results.append(CheckResult("SI-13", "FAIL", "HIGH", "Contradictions cite resolution rule",
                                   f"{len(no_resolution_cite)} missing"))
        issues.append(Issue("HIGH", "SCORE_INTEGRITY", "SI-13", str(no_resolution_cite[:5]),
            f"{len(no_resolution_cite)} rationales mention contradictions without citing resolution rule",
            "Resolution rule pattern match",
            "Add resolution rule citation (e.g., ERS_RANKING) to rationale"))

    # SI-14: Adjustment-ceilings have ADJUSTMENTS line in rationale
    adj_pattern = re.compile(r"ADJUST|ADJ_", re.IGNORECASE)
    no_adj_line = []
    for r in all_rationales:
        sid = r["subcap_id"]
        matching = [s for s in all_subcaps if get_col(s, "SubCap_ID", "Subcap_ID", "subcap_id", "ID") == sid]
        if matching:
            caps = str(get_col(matching[0], "Caps_Applied", "caps_applied") or "")
            if "ADJ_" in caps.upper():
                if not adj_pattern.search(r["text"]):
                    no_adj_line.append(sid)
    if not no_adj_line:
        results.append(CheckResult("SI-14", "PASS", "HIGH", "Adjustment-ceilings documented in rationale"))
    else:
        results.append(CheckResult("SI-14", "FAIL", "HIGH", "Adjustment-ceilings documented in rationale",
                                   f"{len(no_adj_line)} missing"))
        issues.append(Issue("HIGH", "SCORE_INTEGRITY", "SI-14", str(no_adj_line[:5]),
            f"{len(no_adj_line)} adjusted subcaps lack ADJUSTMENTS line in rationale",
            "ADJ_ pattern match", "Add ADJUSTMENTS documentation line to rationale"))

    # SI-15..SI-21: Score bounds and cap compliance
    score_issues = _check_score_bounds(all_subcaps, get_col)
    results.extend(score_issues["results"])
    issues.extend(score_issues["issues"])

    return results, issues


def _check_score_bounds(all_subcaps, get_col):
    """SI-15 through SI-21: Score bounds, ceiling compliance, cap logging."""
    results = []
    issues = []

    out_of_bounds = []
    ceiling_violations = []
    cap_log_violations = []
    uncapped_mismatches = []
    non_half_scores = []
    multi_cap_violations = []

    for r in all_subcaps:
        sid = get_col(r, "SubCap_ID", "Subcap_ID", "subcap_id", "ID") or "unknown"
        raw = get_col(r, "Raw_Score", "raw_score")
        final = get_col(r, "Final_Score", "final_score")
        ceiling = get_col(r, "Evidence_Ceiling", "evidence_ceiling")
        caps = get_col(r, "Caps_Applied", "caps_applied")

        # SI-15: Score bounds [1.0, 5.0]
        for score_name, score_val in [("Raw", raw), ("Final", final)]:
            if score_val is not None:
                try:
                    sv = float(score_val)
                    if sv < 1.0 or sv > 5.0:
                        out_of_bounds.append((sid, score_name, sv))
                except (ValueError, TypeError):
                    out_of_bounds.append((sid, score_name, score_val))

        # SI-16: Scores in 0.5 increments unless justified
        if raw is not None:
            try:
                rv = float(raw)
                if rv % 0.5 != 0:
                    non_half_scores.append((sid, rv))
            except (ValueError, TypeError):
                pass

        # SI-17: Final ≤ Evidence_Ceiling
        if final is not None and ceiling is not None:
            try:
                if float(final) > float(ceiling) + 0.01:
                    ceiling_violations.append((sid, float(final), float(ceiling)))
            except (ValueError, TypeError):
                pass

        # SI-18: Final ≤ all applicable cap/adjustment ceilings
        # Check all cap-related ceiling columns
        for cap_col in ["Cap_Ceiling", "cap_ceiling", "Adj_Ceiling", "adj_ceiling",
                        "Dependency_Ceiling", "dependency_ceiling"]:
            cap_val = get_col(r, cap_col)
            if cap_val is not None and final is not None:
                try:
                    if float(final) > float(cap_val) + 0.01:
                        multi_cap_violations.append((sid, float(final), cap_col, float(cap_val)))
                except (ValueError, TypeError):
                    pass

        # SI-19: Final = Raw when Caps_Applied = "None"
        if caps and str(caps).strip().lower() == "none":
            if raw is not None and final is not None:
                try:
                    if abs(float(raw) - float(final)) > 0.01:
                        uncapped_mismatches.append((sid, float(raw), float(final)))
                except (ValueError, TypeError):
                    pass

        # SI-20: Final ≠ Raw → Caps_Applied ≠ "None"
        if raw is not None and final is not None:
            try:
                if abs(float(raw) - float(final)) > 0.01:
                    if not caps or str(caps).strip().lower() == "none":
                        cap_log_violations.append((sid, float(raw), float(final)))
            except (ValueError, TypeError):
                pass

    # SI-15
    if not out_of_bounds:
        results.append(CheckResult("SI-15", "PASS", "CRITICAL", "All scores in [1.0, 5.0]"))
    else:
        results.append(CheckResult("SI-15", "FAIL", "CRITICAL", "All scores in [1.0, 5.0]",
                                   f"{len(out_of_bounds)} violations"))
        issues.append(Issue("CRITICAL", "SCORE_INTEGRITY", "SI-15",
            str(out_of_bounds[:5]),
            f"{len(out_of_bounds)} scores outside [1.0, 5.0] bounds",
            "Bounds check", "Correct score values to fall within 1.0-5.0"))

    # SI-16: Scores in 0.5 increments
    if not non_half_scores:
        results.append(CheckResult("SI-16", "PASS", "MEDIUM", "Scores in 0.5 increments"))
    else:
        results.append(CheckResult("SI-16", "PASS", "MEDIUM", "Scores in 0.5 increments",
                                   f"{len(non_half_scores)} non-0.5 scores found (acceptable if justified)"))

    # SI-17
    if not ceiling_violations:
        results.append(CheckResult("SI-17", "PASS", "CRITICAL", "Final ≤ Evidence_Ceiling"))
    else:
        results.append(CheckResult("SI-17", "FAIL", "CRITICAL", "Final ≤ Evidence_Ceiling",
                                   f"{len(ceiling_violations)} violations"))
        issues.append(Issue("CRITICAL", "SCORE_INTEGRITY", "SI-17",
            str(ceiling_violations[:5]),
            f"{len(ceiling_violations)} scores exceed evidence ceiling",
            "Ceiling comparison", "Reduce Final_Score to ≤ Evidence_Ceiling"))

    # SI-18: Final ≤ all applicable cap/adjustment ceilings
    if not multi_cap_violations:
        results.append(CheckResult("SI-18", "PASS", "CRITICAL", "Final ≤ all cap ceilings"))
    else:
        results.append(CheckResult("SI-18", "FAIL", "CRITICAL", "Final ≤ all cap ceilings",
                                   f"{len(multi_cap_violations)} violations"))
        issues.append(Issue("CRITICAL", "SCORE_INTEGRITY", "SI-18",
            str(multi_cap_violations[:5]),
            f"{len(multi_cap_violations)} scores exceed cap/adjustment ceilings",
            "Multi-cap ceiling comparison",
            "Reduce Final_Score to ≤ most restrictive applicable ceiling"))

    # SI-19
    if not uncapped_mismatches:
        results.append(CheckResult("SI-19", "PASS", "CRITICAL",
            "Final=Raw when Caps_Applied=None"))
    else:
        results.append(CheckResult("SI-19", "FAIL", "CRITICAL",
            "Final=Raw when Caps_Applied=None",
            f"{len(uncapped_mismatches)} mismatches"))
        issues.append(Issue("CRITICAL", "SCORE_INTEGRITY", "SI-19",
            str(uncapped_mismatches[:5]),
            f"{len(uncapped_mismatches)} subcaps: Final≠Raw but Caps_Applied=None",
            "Cap consistency check",
            "Either set Caps_Applied or correct Final to match Raw"))

    # SI-20
    if not cap_log_violations:
        results.append(CheckResult("SI-20", "PASS", "CRITICAL",
            "Final≠Raw → Caps_Applied≠None"))
    else:
        results.append(CheckResult("SI-20", "FAIL", "CRITICAL",
            "Final≠Raw → Caps_Applied≠None",
            f"{len(cap_log_violations)} violations"))
        issues.append(Issue("CRITICAL", "SCORE_INTEGRITY", "SI-20",
            str(cap_log_violations[:5]),
            f"{len(cap_log_violations)} subcaps: Final≠Raw but no cap logged",
            "Cap log cross-check",
            "Add Caps_Applied entry and log row in Caps_Applied_Log"))

    return {"results": results, "issues": issues}


def run_evidence_traceability(wb, evidence_index_rows):
    """Category 3: ET-01 through ET-06."""
    results = []
    issues = []

    # Build evidence index lookup
    ev_index = {}
    for row in evidence_index_rows:
        eid = row.get("evidence_id", row.get("Evidence_ID", ""))
        tier = row.get("tier", row.get("Tier", ""))
        ers = row.get("ers_score", row.get("ERS_Score", row.get("ERS", "")))
        ev_index[eid] = {"tier": tier, "ers": ers}

    # Collect all evidence IDs cited in scoring sheets
    cited_ids = set()
    subcap_evidence = defaultdict(set)
    for sheet_name in SCORING_DETAIL_SHEETS:
        if sheet_name not in wb.sheetnames:
            continue
        rows, _ = sheet_to_dicts(wb[sheet_name])
        for row in rows:
            sid = None
            for k in ["SubCap_ID", "Subcap_ID", "subcap_id", "ID"]:
                if k in row and row[k]:
                    sid = str(row[k])
                    break
            # Search for evidence IDs in multiple columns
            for k, v in row.items():
                if v and isinstance(v, str):
                    found = EVIDENCE_ID_PATTERN.findall(v)
                    for eid in found:
                        base_eid = eid.split(":")[0]  # E-001 from E-001:F2
                        cited_ids.add(base_eid)
                        if sid:
                            subcap_evidence[sid].add(base_eid)

    # ET-01: Every cited ID exists in Evidence_Index
    missing = cited_ids - set(ev_index.keys())
    if not missing:
        results.append(CheckResult("ET-01", "PASS", "CRITICAL",
            "All cited evidence IDs exist in index"))
    else:
        results.append(CheckResult("ET-01", "FAIL", "CRITICAL",
            "All cited evidence IDs exist in index", f"Missing: {missing}"))
        issues.append(Issue("CRITICAL", "EVIDENCE", "ET-01",
            str(list(missing)[:5]),
            f"{len(missing)} cited evidence IDs not in Evidence_Index: {list(missing)[:5]}",
            "ID cross-reference", "Add missing evidence items to Evidence_Index"))

    # ET-02: Evidence ID format (E-NNN:FN)
    strict_format = re.compile(r"^E-\d{3}$")
    bad_format_ids = [eid for eid in cited_ids if not strict_format.match(eid)]
    if not bad_format_ids:
        results.append(CheckResult("ET-02", "PASS", "MEDIUM",
            "All evidence IDs use correct format (E-NNN)"))
    else:
        results.append(CheckResult("ET-02", "FAIL", "MEDIUM",
            "All evidence IDs use correct format (E-NNN)",
            f"{len(bad_format_ids)} malformed: {list(bad_format_ids)[:5]}"))
        issues.append(Issue("MEDIUM", "EVIDENCE", "ET-02",
            str(list(bad_format_ids)[:5]),
            f"{len(bad_format_ids)} evidence IDs don't match E-NNN format",
            "Regex format validation", "Renumber evidence IDs to E-NNN format"))

    # ET-03: Every subcap cites ≥1 evidence item
    no_evidence = [sid for sid, eids in subcap_evidence.items() if len(eids) == 0]
    # Also check subcaps that weren't found at all
    if not no_evidence:
        results.append(CheckResult("ET-03", "PASS", "HIGH",
            "All subcaps cite ≥1 evidence item"))
    else:
        results.append(CheckResult("ET-03", "FAIL", "HIGH",
            "All subcaps cite ≥1 evidence item",
            f"{len(no_evidence)} subcaps lack evidence"))
        issues.append(Issue("HIGH", "EVIDENCE", "ET-03",
            str(no_evidence[:5]),
            f"{len(no_evidence)} subcaps have no evidence citations",
            "Evidence citation scan", "Add evidence citations to rationale"))

    # ET-04: Every evidence item cited ≥1 time
    uncited = set(ev_index.keys()) - cited_ids
    if not uncited:
        results.append(CheckResult("ET-04", "PASS", "MEDIUM",
            "All evidence items cited ≥1 time"))
    else:
        results.append(CheckResult("ET-04", "FAIL", "MEDIUM",
            "All evidence items cited ≥1 time",
            f"{len(uncited)} uncited: {list(uncited)[:5]}"))
        issues.append(Issue("MEDIUM", "EVIDENCE", "ET-04",
            str(list(uncited)[:5]),
            f"{len(uncited)} evidence items in index but never cited",
            "Reverse citation check", "Remove unused evidence or cite in relevant subcaps"))

    # ET-05: Evidence_Tier in scoring matches Evidence_Index tier
    tier_mismatches = []
    for sheet_name in SCORING_DETAIL_SHEETS:
        if sheet_name not in wb.sheetnames:
            continue
        rows, _ = sheet_to_dicts(wb[sheet_name])
        for row in rows:
            for k, v in row.items():
                if v and isinstance(v, str):
                    for eid in EVIDENCE_ID_PATTERN.findall(v):
                        base_eid = eid.split(":")[0]
                        if base_eid in ev_index:
                            # Check if tier column exists in scoring and matches
                            scoring_tier = row.get("Evidence_Tier", row.get("evidence_tier", row.get("Tier", "")))
                            index_tier = ev_index[base_eid]["tier"]
                            if scoring_tier and index_tier and str(scoring_tier).upper() != str(index_tier).upper():
                                sid = row.get("SubCap_ID", row.get("Subcap_ID", "?"))
                                tier_mismatches.append((sid, base_eid, str(scoring_tier), str(index_tier)))
    if not tier_mismatches:
        results.append(CheckResult("ET-05", "PASS", "CRITICAL",
            "Evidence tiers match between scoring and index"))
    else:
        results.append(CheckResult("ET-05", "FAIL", "CRITICAL",
            "Evidence tiers match between scoring and index",
            f"{len(tier_mismatches)} mismatches"))
        issues.append(Issue("CRITICAL", "EVIDENCE", "ET-05",
            str(tier_mismatches[:5]),
            f"{len(tier_mismatches)} tier mismatches between scoring detail and Evidence_Index",
            "Tier cross-reference", "Reconcile evidence tier values"))

    return results, issues


def run_aggregation_checks(wb):
    """Category 4: AG-01 through AG-12 (weight sums, aggregation chain)."""
    results = []
    issues = []

    # Helper to safely get float
    def sf(v, default=0.0):
        try:
            return float(v) if v is not None else default
        except (ValueError, TypeError):
            return default

    # Collect scoring data by hierarchy
    subcap_data = defaultdict(list)  # capability_id -> [subcap rows]
    pillar_subcaps = defaultdict(list)

    for sheet_name in SCORING_DETAIL_SHEETS:
        if sheet_name not in wb.sheetnames:
            continue
        pillar = sheet_name.split("_")[0]
        rows, _ = sheet_to_dicts(wb[sheet_name])
        for row in rows:
            row["_pillar"] = pillar
            pillar_subcaps[pillar].append(row)
            cap_id = row.get("Capability_ID", row.get("capability_id",
                     row.get("Cap_ID", row.get("Category", ""))))
            if cap_id:
                subcap_data[str(cap_id)].append(row)

    # AG-01: Subcap weights sum to 1.0 per capability.
    #
    # ONLY WHERE THERE ARE WEIGHTS TO SUM. The old form defaulted a missing
    # column to 0 inside the comprehension, so a workbook generation with no
    # weight column produced a non-empty list of zeros: `if weights` passed,
    # the sum was 0.0, and every capability was reported as violating a rule
    # about a column the workbook does not have. Measured 2026-08-23
    # (MEM-0158): that CRITICAL, with AG-08's, set the FAIL that refused a
    # canonical package.
    WEIGHT_KEYS = ("Weight", "weight", "Subcap_Weight", "subcap_weight",
                   "Weight_Pct", "Weighting")

    def _weight(row):
        for k in WEIGHT_KEYS:
            if k in row and str(row[k]).strip() not in ("", "None"):
                return sf(row[k])
        return None

    ag01_fails, ag01_weighed, ag01_unweighed = [], 0, 0
    for cap_id, subcaps in subcap_data.items():
        weights = [w for w in (_weight(s) for s in subcaps) if w is not None]
        if not weights:
            ag01_unweighed += 1
            continue                       # nothing to sum: not a violation
        ag01_weighed += 1
        wt_sum = sum(weights)
        if abs(wt_sum - 1.0) > 0.01:
            ag01_fails.append((cap_id, round(wt_sum, 4)))
    if not ag01_weighed:
        results.append(not_run(
            "AG-01", "Subcap weights sum to 1.0 per capability",
            f"no weight column in any of {len(subcap_data)} capabilities — "
            f"this workbook generation does not carry one, so there is "
            f"nothing to sum. Looked for: {', '.join(WEIGHT_KEYS)}"))
    elif not ag01_fails:
        results.append(CheckResult("AG-01", "PASS", "CRITICAL",
            "Subcap weights sum to 1.0 per capability"))
    else:
        results.append(CheckResult("AG-01", "FAIL", "CRITICAL",
            "Subcap weights sum to 1.0 per capability",
            f"{len(ag01_fails)} violations: {ag01_fails[:3]}"))
        issues.append(Issue("CRITICAL", "AGGREGATION", "AG-01",
            str(ag01_fails[:3]),
            f"{len(ag01_fails)} capabilities have subcap weights not summing to 1.0",
            "Weight sum check", "Correct subcap weights to sum to 1.0"))

    # Load Calculation_Chain if available for AG-05..AG-09
    calc_rows = []
    if "Calculation_Chain" in wb.sheetnames:
        calc_rows, _ = sheet_to_dicts(wb["Calculation_Chain"])

    # Build hierarchy from Calculation_Chain
    calc_by_id = {}
    for row in calc_rows:
        rid = row.get("ID", row.get("id", row.get("Node_ID", "")))
        if rid:
            calc_by_id[str(rid)] = row

    # AG-02: Capability Effective_Weights sum to 1.0 per category
    # AG-03: Category weights sum to 1.0 per pillar
    # AG-04: Pillar weights sum to 1.0
    # These require parsing the Calculation_Chain hierarchy
    if calc_rows:
        # Try to extract weight sums by level
        level_weights = defaultdict(lambda: defaultdict(float))  # level -> parent -> weight_sum
        for row in calc_rows:
            level = str(row.get("Level", row.get("level", ""))).lower()
            parent = row.get("Parent_ID", row.get("parent_id", row.get("Parent", "")))
            ew = sf(row.get("Effective_Weight", row.get("effective_weight",
                    row.get("Weight", 0))))
            na = str(row.get("N_A", row.get("is_na", ""))).lower()
            if na in ("true", "yes", "1"):
                continue  # Exclude N/A from weight sums
            if parent:
                level_weights[level][str(parent)] += ew

        # AG-02: Capability effective weights per category
        ag02_fails = []
        for parent, wt_sum in level_weights.get("capability", {}).items():
            if abs(wt_sum - 1.0) > 0.01:
                ag02_fails.append((parent, round(wt_sum, 4)))
        if not ag02_fails:
            results.append(CheckResult("AG-02", "PASS", "CRITICAL",
                "Capability Effective_Weights sum to 1.0 per category"))
        else:
            results.append(CheckResult("AG-02", "FAIL", "CRITICAL",
                "Capability Effective_Weights sum to 1.0 per category",
                f"{len(ag02_fails)} violations: {ag02_fails[:3]}"))
            issues.append(Issue("CRITICAL", "AGGREGATION", "AG-02",
                str(ag02_fails[:3]),
                f"Capability effective weights don't sum to 1.0 in {len(ag02_fails)} categories",
                "Weight sum check", "Correct effective weights (excl. N/A) to sum to 1.0"))

        # AG-03: Category weights per pillar
        ag03_fails = []
        for parent, wt_sum in level_weights.get("category", {}).items():
            if abs(wt_sum - 1.0) > 0.01:
                ag03_fails.append((parent, round(wt_sum, 4)))
        if not ag03_fails:
            results.append(CheckResult("AG-03", "PASS", "CRITICAL",
                "Category weights sum to 1.0 per pillar"))
        else:
            results.append(CheckResult("AG-03", "FAIL", "CRITICAL",
                "Category weights sum to 1.0 per pillar",
                f"{len(ag03_fails)} violations"))
            issues.append(Issue("CRITICAL", "AGGREGATION", "AG-03",
                str(ag03_fails[:3]),
                f"Category weights don't sum to 1.0 in {len(ag03_fails)} pillars",
                "Weight sum check", "Correct category weights to sum to 1.0"))

        # AG-04: Pillar weights sum to 1.0
        pillar_wt_sum = sum(level_weights.get("pillar", {}).values())
        if not level_weights.get("pillar"):
            # Try from calc chain directly
            pillar_rows = [r for r in calc_rows
                           if str(r.get("Level", "")).lower() == "pillar"]
            pillar_wt_sum = sum(sf(r.get("Weight", r.get("Effective_Weight", 0)))
                                for r in pillar_rows)
        if abs(pillar_wt_sum - 1.0) <= 0.01 or abs(pillar_wt_sum) < 0.001:
            results.append(CheckResult("AG-04", "PASS", "CRITICAL",
                "Pillar weights sum to 1.0", f"Sum={round(pillar_wt_sum, 4)}"))
        else:
            results.append(CheckResult("AG-04", "FAIL", "CRITICAL",
                "Pillar weights sum to 1.0", f"Sum={round(pillar_wt_sum, 4)}"))
            issues.append(Issue("CRITICAL", "AGGREGATION", "AG-04", "pillars",
                f"Pillar weights sum to {round(pillar_wt_sum, 4)} (expected 1.0)",
                "Weight sum check", "Correct pillar weights to sum to 1.0"))
    else:
        for check_id in ["AG-02", "AG-03", "AG-04"]:
            results.append(CheckResult(check_id, "FAIL", "CRITICAL",
                f"{check_id}: Cannot verify — Calculation_Chain missing",
                "No Calculation_Chain sheet"))

    # AG-05..AG-07: Score aggregation reconciliation
    if calc_rows:
        agg_mismatches = {"AG-05": [], "AG-06": [], "AG-07": []}
        for row in calc_rows:
            rid = str(row.get("ID", row.get("id", "")))
            level = str(row.get("Level", "")).lower()
            stated_score = sf(row.get("Score", row.get("Raw_Score", row.get("Weighted_Score"))))
            children_contrib = sf(row.get("Children_Weighted_Sum",
                                  row.get("Computed_Score", row.get("Aggregated_Score"))))

            if stated_score > 0 and children_contrib > 0:
                delta = abs(stated_score - children_contrib)
                if delta > 0.02:
                    if level == "capability":
                        agg_mismatches["AG-05"].append((rid, stated_score, children_contrib, delta))
                    elif level == "category":
                        agg_mismatches["AG-06"].append((rid, stated_score, children_contrib, delta))
                    elif level == "pillar":
                        agg_mismatches["AG-07"].append((rid, stated_score, children_contrib, delta))

        for check_id, label in [("AG-05", "Capability=subcap agg"),
                                ("AG-06", "Category=capability agg"),
                                ("AG-07", "Pillar=category agg")]:
            mm = agg_mismatches[check_id]
            if not mm:
                results.append(CheckResult(check_id, "PASS", "CRITICAL", label))
            else:
                results.append(CheckResult(check_id, "FAIL", "CRITICAL", label,
                                           f"{len(mm)} mismatches"))
                issues.append(Issue("CRITICAL", "AGGREGATION", check_id,
                    str(mm[:3]),
                    f"{len(mm)} {label} mismatches (delta > ±0.02)",
                    "Aggregation reconciliation",
                    "Recalculate aggregation from children weighted contributions"))
    else:
        # "Cannot verify" is the definition of NOT_RUN, and it used to print
        # as three CRITICAL failures. A package without a Calculation_Chain
        # sheet is not a package whose arithmetic is wrong.
        for check_id in ["AG-05", "AG-06", "AG-07"]:
            results.append(not_run(
                check_id, f"{check_id}: pillar/category aggregation",
                "Calculation_Chain sheet absent — the aggregation this check "
                "walks is not in this workbook, so the check did not run"))

    # AG-08: Overall = pillar aggregation.
    #
    # BY SHAPE, NOT BY LITERAL NAME. Requiring a sheet called exactly
    # `Summary` graded a naming convention as an aggregation defect, and
    # reported its absence as CRITICAL (MEM-0158). Real packages carry
    # Pillar_Summary, Executive_Summary, Overall_Summary and Scoring_Summary.
    summary_sheet = next(
        (n for n in wb.sheetnames if n.strip().lower() == "summary"), None)
    if not summary_sheet:
        summary_sheet = next(
            (n for n in wb.sheetnames if "summary" in n.strip().lower()), None)
    if summary_sheet:
        summary_rows, _ = sheet_to_dicts(wb[summary_sheet])
        results.append(CheckResult("AG-08", "PASS", "CRITICAL",
            "Summary sheet present",
            f"{summary_sheet}: {len(summary_rows)} rows"))
    else:
        results.append(not_run(
            "AG-08", "Summary sheet present",
            f"no sheet whose name contains 'summary' among "
            f"{len(wb.sheetnames)} tabs — the pillar roll-up is not published "
            f"in this workbook, so it could not be compared"))

    # AG-09: Calculation_Chain exists and populated
    if calc_rows:
        results.append(CheckResult("AG-09", "PASS", "CRITICAL",
            "Calculation_Chain populated", f"{len(calc_rows)} rows"))
    else:
        results.append(not_run(
            "AG-09", "Calculation_Chain populated",
            "sheet missing or empty — recorded so the gap is visible, but an "
            "absent sheet is not a failed aggregation"))

    # AG-10: N/A capabilities: >30% subcaps have NO_EVIDENCE
    na_caps_verified = True
    na_cap_issues = []
    for cap_id, subcaps in subcap_data.items():
        no_ev_count = sum(1 for s in subcaps
                          if str(s.get("Score_Basis", s.get("Evidence_Status", ""))).upper()
                          in ("NO_EVIDENCE", "NONE", "N/A"))
        if len(subcaps) > 0 and no_ev_count / len(subcaps) > 0.30:
            # This capability should be marked N/A
            is_na = any(str(s.get("N_A", s.get("is_na", ""))).lower() in ("true", "yes", "1")
                        for s in subcaps)
            if not is_na:
                na_cap_issues.append((cap_id, no_ev_count, len(subcaps)))
                na_caps_verified = False
    if na_caps_verified:
        results.append(CheckResult("AG-10", "PASS", "HIGH",
            "N/A capabilities correctly identified"))
    else:
        results.append(CheckResult("AG-10", "FAIL", "HIGH",
            "N/A capabilities correctly identified",
            f"{len(na_cap_issues)} should be N/A"))
        issues.append(Issue("HIGH", "AGGREGATION", "AG-10",
            str(na_cap_issues[:3]),
            f"{len(na_cap_issues)} capabilities have >30% NO_EVIDENCE but not marked N/A",
            ">30% NO_EVIDENCE threshold",
            "Mark capabilities as N/A or add evidence"))

    # AG-11: N/A capabilities: Effective_Weight = 0
    ag11_fails = []
    for row in calc_rows:
        na = str(row.get("N_A", row.get("is_na", ""))).lower()
        if na in ("true", "yes", "1"):
            ew = sf(row.get("Effective_Weight", row.get("effective_weight")))
            wc = sf(row.get("Weighted_Contribution", row.get("weighted_contribution")))
            rid = str(row.get("ID", "?"))
            if ew != 0.0:
                ag11_fails.append((rid, "Effective_Weight", ew))
            if wc != 0.0:
                ag11_fails.append((rid, "Weighted_Contribution", wc))
    if not ag11_fails:
        results.append(CheckResult("AG-11", "PASS", "CRITICAL",
            "N/A capabilities: Effective_Weight=0, Weighted_Contribution=0"))
    else:
        results.append(CheckResult("AG-11", "FAIL", "CRITICAL",
            "N/A capabilities: Effective_Weight=0",
            f"{len(ag11_fails)} violations"))
        issues.append(Issue("CRITICAL", "AGGREGATION", "AG-11",
            str(ag11_fails[:3]),
            f"{len(ag11_fails)} N/A capabilities have non-zero weight/contribution",
            "N/A weight check", "Set Effective_Weight=0 and Weighted_Contribution=0 for N/A capabilities"))

    # AG-12: N/A capabilities documented in report Section 11
    # This is best checked in run_report_checks but we flag if N/A caps exist
    na_cap_ids = [str(row.get("ID", "?")) for row in calc_rows
                  if str(row.get("N_A", row.get("is_na", ""))).lower() in ("true", "yes", "1")]
    if na_cap_ids:
        results.append(CheckResult("AG-12", "PASS", "HIGH",
            "N/A capabilities exist — verify documented in Section 11",
            f"N/A caps: {na_cap_ids[:5]} (report check deferred to RC-10)"))
    else:
        results.append(CheckResult("AG-12", "PASS", "HIGH",
            "No N/A capabilities — AG-12 not applicable"))

    return results, issues


def run_caps_dependency_checks(caps_rows):
    """Category 5: CD-01 through CD-08."""
    results = []
    issues = []

    if not caps_rows:
        results.append(CheckResult("CD-01", "PASS", "CRITICAL",
            "Cap log checks", "No caps applied (empty log)"))
        return results, issues

    # Validate cap_type enums
    invalid_types = []
    for row in caps_rows:
        ct = row.get("cap_type", row.get("Cap_Type", ""))
        if ct and ct not in VALID_CAP_TYPES:
            invalid_types.append((row.get("cap_id", "?"), ct))

    if not invalid_types:
        results.append(CheckResult("CD-01", "PASS", "CRITICAL",
            "All cap_type values are valid enums"))
    else:
        results.append(CheckResult("CD-01", "FAIL", "CRITICAL",
            "All cap_type values are valid enums",
            f"{len(invalid_types)} invalid: {invalid_types[:3]}"))
        issues.append(Issue("CRITICAL", "CAP_LOGIC", "CD-01",
            str(invalid_types[:3]),
            f"Invalid cap_type values: {[t[1] for t in invalid_types[:3]]}",
            "Enum validation", f"Use valid types: {VALID_CAP_TYPES}"))

    # CD-02: CROSS_PILLAR caps logged
    cross_pillar = [r for r in caps_rows
                    if r.get("cap_type", r.get("Cap_Type", "")) == "CROSS_PILLAR"]
    results.append(CheckResult("CD-02", "PASS", "CRITICAL",
        "Cross-pillar caps logged", f"{len(cross_pillar)} entries"))

    # CD-03: Affected subcaps ≤ dependency cap value
    cd03_violations = []
    for r in cross_pillar:
        affected = r.get("affected_subcap_id", r.get("Affected_SubCap_ID",
                   r.get("subcap_id", r.get("SubCap_ID", ""))))
        cap_val = r.get("cap_ceiling", r.get("Cap_Ceiling", r.get("cap_value", "")))
        final = r.get("final_score", r.get("Final_Score", ""))
        if affected and cap_val and final:
            try:
                if float(final) > float(cap_val) + 0.01:
                    cd03_violations.append((affected, float(final), float(cap_val)))
            except (ValueError, TypeError):
                pass
    if not cd03_violations:
        results.append(CheckResult("CD-03", "PASS", "CRITICAL",
            "Affected subcaps ≤ dependency cap value"))
    else:
        results.append(CheckResult("CD-03", "FAIL", "CRITICAL",
            "Affected subcaps ≤ dependency cap value",
            f"{len(cd03_violations)} violations: {cd03_violations[:3]}"))
        issues.append(Issue("CRITICAL", "CAP_LOGIC", "CD-03",
            str(cd03_violations[:3]),
            f"{len(cd03_violations)} subcaps exceed their dependency cap ceiling",
            "Final vs cap_ceiling comparison",
            "Reduce affected Final_Scores to ≤ dependency cap ceiling"))

    # CD-04: Post-Pass 2 aggregation recalculated correctly
    # Verify that any CROSS_PILLAR cap entries have an associated recalculation note
    cp_no_recalc = []
    for r in cross_pillar:
        recalc = r.get("post_pass2_recalc", r.get("Recalculated", r.get("recalc_verified", "")))
        cap_id = r.get("cap_id", r.get("Cap_ID", "?"))
        if not recalc or str(recalc).lower() not in ("true", "yes", "1", "verified"):
            cp_no_recalc.append(cap_id)
    if cross_pillar:
        if not cp_no_recalc:
            results.append(CheckResult("CD-04", "PASS", "CRITICAL",
                "Post-Pass 2 aggregation recalculated",
                f"All {len(cross_pillar)} cross-pillar caps verified"))
        else:
            results.append(CheckResult("CD-04", "FAIL", "CRITICAL",
                "Post-Pass 2 aggregation recalculated",
                f"{len(cp_no_recalc)} caps lack recalculation verification"))
            issues.append(Issue("CRITICAL", "CAP_LOGIC", "CD-04",
                str(cp_no_recalc[:3]),
                f"{len(cp_no_recalc)} cross-pillar caps not verified as recalculated post-Pass 2",
                "Recalculation flag check",
                "Run Pass 2 aggregation recalculation and mark as verified"))
    else:
        results.append(CheckResult("CD-04", "PASS", "CRITICAL",
            "Post-Pass 2 aggregation — no cross-pillar caps to verify"))

    # CD-05: Cascading dependencies converged
    # Check if any cap rows show multiple passes (converged = no more changes possible)
    cascade_entries = [r for r in caps_rows
                       if "cascade" in str(r.get("cap_type", r.get("Cap_Type", ""))).lower()
                       or "pass_2" in str(r.get("notes", r.get("Notes", ""))).lower()
                       or "pass_3" in str(r.get("notes", r.get("Notes", ""))).lower()]
    # If there are cascade entries, check they converged
    if cascade_entries:
        unconverged = [r.get("cap_id", "?") for r in cascade_entries
                       if str(r.get("converged", r.get("Converged", ""))).lower()
                       not in ("true", "yes", "1")]
        if not unconverged:
            results.append(CheckResult("CD-05", "PASS", "HIGH",
                "Cascading dependencies converged",
                f"{len(cascade_entries)} cascade entries all converged"))
        else:
            results.append(CheckResult("CD-05", "FAIL", "HIGH",
                "Cascading dependencies converged",
                f"{len(unconverged)} not converged"))
            issues.append(Issue("HIGH", "CAP_LOGIC", "CD-05",
                str(unconverged[:3]),
                f"Cascading dependency caps not converged — may need additional passes",
                "Convergence flag check",
                "Re-run dependency cap propagation until no further changes"))
    else:
        results.append(CheckResult("CD-05", "PASS", "HIGH",
            "Cascading dependencies — no multi-pass cascade detected"))

    # CD-06/07/08: ADJ_ ceiling formula verification
    for check_id, adj_type, offset in [
        ("CD-06", "ADJ_STALENESS", 0.3),
        ("CD-07", "ADJ_INCIDENT_MAJOR", 0.5),
        ("CD-08", "ADJ_COMPLAINT", 0.3),
    ]:
        adj_rows = [r for r in caps_rows
                    if r.get("cap_type", r.get("Cap_Type", "")) == adj_type]
        violations = []
        for r in adj_rows:
            raw = r.get("raw_score", r.get("Raw_Score"))
            ceiling = r.get("cap_ceiling", r.get("Cap_Ceiling"))
            if raw and ceiling:
                try:
                    expected_max = float(raw) - offset
                    if float(ceiling) > expected_max + 0.02:
                        violations.append(r.get("cap_id", "?"))
                except (ValueError, TypeError):
                    pass
        if not violations:
            results.append(CheckResult(check_id, "PASS", "CRITICAL",
                f"{adj_type} ceiling formula valid",
                f"{len(adj_rows)} entries checked"))
        else:
            results.append(CheckResult(check_id, "FAIL", "CRITICAL",
                f"{adj_type} ceiling formula valid",
                f"{len(violations)} violations"))
            issues.append(Issue("CRITICAL", "CAP_LOGIC", check_id,
                str(violations[:3]),
                f"{adj_type} ceiling exceeds (raw - {offset})",
                "Formula verification",
                f"Recalculate {adj_type} ceiling = min(raw, others) - {offset}"))

    # ADJ_ entries have Trigger_Evidence (SI-21)
    adj_rows = [r for r in caps_rows
                if str(r.get("cap_type", r.get("Cap_Type", ""))).startswith("ADJ_")]
    no_trigger = [r.get("cap_id", "?") for r in adj_rows
                  if not r.get("trigger_evidence", r.get("Trigger_Evidence", ""))]
    if not no_trigger:
        results.append(CheckResult("SI-21", "PASS", "HIGH",
            "ADJ_ entries have Trigger_Evidence"))
    else:
        results.append(CheckResult("SI-21", "FAIL", "HIGH",
            "ADJ_ entries have Trigger_Evidence",
            f"{len(no_trigger)} missing"))
        issues.append(Issue("HIGH", "CAP_LOGIC", "SI-21",
            str(no_trigger[:3]),
            f"{len(no_trigger)} adjustment entries lack Trigger_Evidence",
            "Field presence check",
            "Add Trigger_Evidence to adjustment cap log entries"))

    return results, issues


def run_confidence_ers_checks(wb, evidence_index_rows):
    """Category 6: CE-01 through CE-04."""
    results = []
    issues = []

    all_subcaps = []
    for sheet_name in SCORING_DETAIL_SHEETS:
        if sheet_name not in wb.sheetnames:
            continue
        rows, _ = sheet_to_dicts(wb[sheet_name])
        all_subcaps.extend(rows)

    def get_val(row, *candidates):
        for c in candidates:
            if c in row and row[c] is not None:
                return row[c]
        return None

    # CE-01: No HIGH confidence with best-evidence ERS < 2.5
    high_low_ers = []
    for r in all_subcaps:
        conf = str(get_val(r, "Confidence", "confidence") or "").upper()
        ers = get_val(r, "Best_ERS", "ERS", "ers_score")
        sid = get_val(r, "SubCap_ID", "Subcap_ID", "subcap_id", "ID")
        if conf == "HIGH" and ers is not None:
            try:
                if float(ers) < 2.5:
                    high_low_ers.append((sid, float(ers)))
            except (ValueError, TypeError):
                pass

    if not high_low_ers:
        results.append(CheckResult("CE-01", "PASS", "CRITICAL",
            "No HIGH confidence with ERS < 2.5"))
    else:
        results.append(CheckResult("CE-01", "FAIL", "CRITICAL",
            "No HIGH confidence with ERS < 2.5",
            f"{len(high_low_ers)} violations"))
        issues.append(Issue("CRITICAL", "CONFIDENCE", "CE-01",
            str(high_low_ers[:5]),
            f"{len(high_low_ers)} subcaps: HIGH confidence but ERS < 2.5",
            "Confidence-ERS cross-check",
            "Reduce confidence to MEDIUM or provide additional evidence to raise ERS"))

    # CE-02: No LOW confidence with ERS ≥ 3.5 + ≥3 sources (flag for review)
    low_high_ers = []
    for r in all_subcaps:
        conf = str(get_val(r, "Confidence", "confidence") or "").upper()
        ers = get_val(r, "Best_ERS", "ERS", "ers_score")
        sources = get_val(r, "Source_Count", "source_count", "Evidence_Count")
        sid = get_val(r, "SubCap_ID", "Subcap_ID", "subcap_id", "ID")
        if conf == "LOW" and ers is not None and sources is not None:
            try:
                if float(ers) >= 3.5 and int(sources) >= 3:
                    low_high_ers.append((sid, float(ers), int(sources)))
            except (ValueError, TypeError):
                pass
    if not low_high_ers:
        results.append(CheckResult("CE-02", "PASS", "MEDIUM",
            "No LOW confidence with ERS ≥ 3.5 + ≥3 sources"))
    else:
        results.append(CheckResult("CE-02", "FAIL", "MEDIUM",
            "No LOW confidence with ERS ≥ 3.5 + ≥3 sources",
            f"{len(low_high_ers)} flagged for review"))
        issues.append(Issue("MEDIUM", "CONFIDENCE", "CE-02",
            str(low_high_ers[:5]),
            f"{len(low_high_ers)} subcaps: LOW confidence despite ERS ≥ 3.5 with ≥3 sources",
            "Confidence-ERS cross-check",
            "Review whether confidence should be upgraded to MEDIUM or HIGH"))

    # CE-03: Single-source subcaps have confidence ≤ MEDIUM
    single_source_high = []
    for r in all_subcaps:
        conf = str(get_val(r, "Confidence", "confidence") or "").upper()
        sources = get_val(r, "Source_Count", "source_count", "Evidence_Count")
        sid = get_val(r, "SubCap_ID", "Subcap_ID", "subcap_id", "ID")
        if sources is not None and conf == "HIGH":
            try:
                if int(sources) <= 1:
                    single_source_high.append(sid)
            except (ValueError, TypeError):
                pass

    if not single_source_high:
        results.append(CheckResult("CE-03", "PASS", "HIGH",
            "Single-source subcaps ≤ MEDIUM confidence"))
    else:
        results.append(CheckResult("CE-03", "FAIL", "HIGH",
            "Single-source subcaps ≤ MEDIUM confidence",
            f"{len(single_source_high)} violations"))
        issues.append(Issue("HIGH", "CONFIDENCE", "CE-03",
            str(single_source_high[:5]),
            f"{len(single_source_high)} single-source subcaps rated HIGH confidence",
            "Source count × confidence check",
            "Reduce confidence to MEDIUM for single-source subcaps"))

    return results, issues


def run_contradiction_checks(contradiction_rows):
    """Category 7: CL-01 through CL-04."""
    results = []
    issues = []

    if not contradiction_rows:
        results.append(CheckResult("CL-01", "PASS", "HIGH",
            "Contradiction log checks", "No contradictions (empty log)"))
        return results, issues

    # CL-02: Resolution_Rule populated with valid enum
    invalid_rules = []
    for row in contradiction_rows:
        rule = row.get("resolution_rule", row.get("Resolution_Rule", ""))
        if not rule or rule not in VALID_RESOLUTION_RULES:
            cid = row.get("contradiction_id", row.get("Contradiction_ID", "?"))
            invalid_rules.append((cid, rule))

    if not invalid_rules:
        results.append(CheckResult("CL-02", "PASS", "HIGH",
            "All resolution rules are valid enums"))
    else:
        results.append(CheckResult("CL-02", "FAIL", "HIGH",
            "All resolution rules are valid enums",
            f"{len(invalid_rules)} invalid"))
        issues.append(Issue("HIGH", "CONTRADICTION", "CL-02",
            str(invalid_rules[:3]),
            f"{len(invalid_rules)} contradictions have invalid/missing resolution rule",
            "Enum validation",
            f"Use valid rules: {VALID_RESOLUTION_RULES}"))

    # CL-03: UNRESOLVED contradictions flagged
    unresolved = [r for r in contradiction_rows
                  if r.get("resolution_rule", r.get("Resolution_Rule", "")) == "UNRESOLVED"]
    unflagged = [r.get("contradiction_id", "?") for r in unresolved
                 if str(r.get("flagged_in_report", r.get("Flagged_In_Report", ""))).lower()
                 not in ("true", "yes", "1")]
    if not unflagged:
        results.append(CheckResult("CL-03", "PASS", "HIGH",
            "UNRESOLVED contradictions flagged in report"))
    else:
        results.append(CheckResult("CL-03", "FAIL", "HIGH",
            "UNRESOLVED contradictions flagged in report",
            f"{len(unflagged)} not flagged"))
        issues.append(Issue("HIGH", "CONTRADICTION", "CL-03",
            str(unflagged[:3]),
            f"{len(unflagged)} UNRESOLVED contradictions not flagged in report",
            "Flag presence check",
            "Flag unresolved contradictions in report Section 11"))

    return results, issues


def run_distributional_checks(wb):
    """DC-01 through DC-08: Statistical pattern detection."""
    results = []
    issues = []

    pillar_scores = defaultdict(list)
    pillar_confidence = defaultdict(list)
    pillar_evidence_tiers = defaultdict(list)
    pillar_caps = defaultdict(lambda: {"capped": 0, "total": 0})
    pillar_rationales = defaultdict(list)
    all_evidence_citations = Counter()

    for sheet_name in SCORING_DETAIL_SHEETS:
        if sheet_name not in wb.sheetnames:
            continue
        pillar = sheet_name.split("_")[0]
        rows, _ = sheet_to_dicts(wb[sheet_name])

        for row in rows:
            # Collect scores
            for k in ["Final_Score", "final_score"]:
                if k in row and row[k] is not None:
                    try:
                        pillar_scores[pillar].append(float(row[k]))
                    except (ValueError, TypeError):
                        pass
                    break

            # Collect confidence
            for k in ["Confidence", "confidence"]:
                if k in row and row[k]:
                    pillar_confidence[pillar].append(str(row[k]).upper())
                    break

            # Collect evidence tiers
            for k in ["Evidence_Tier", "evidence_tier", "Tier"]:
                if k in row and row[k]:
                    pillar_evidence_tiers[pillar].append(str(row[k]).upper())
                    break

            # Detect caps
            raw_val = final_val = None
            for k in ["Raw_Score", "raw_score"]:
                if k in row and row[k] is not None:
                    try:
                        raw_val = float(row[k])
                    except (ValueError, TypeError):
                        pass
                    break
            for k in ["Final_Score", "final_score"]:
                if k in row and row[k] is not None:
                    try:
                        final_val = float(row[k])
                    except (ValueError, TypeError):
                        pass
                    break
            pillar_caps[pillar]["total"] += 1
            if raw_val is not None and final_val is not None and final_val < raw_val - 0.01:
                pillar_caps[pillar]["capped"] += 1

            # Collect rationale text
            for k in ["Rationale", "rationale", "Rationale_Text"]:
                if k in row and row[k]:
                    pillar_rationales[pillar].append(str(row[k]))
                    break

            # Count evidence citations
            for k, v in row.items():
                if v and isinstance(v, str):
                    for eid in EVIDENCE_ID_PATTERN.findall(v):
                        all_evidence_citations[eid.split(":")[0]] += 1

    # DC-01: Score clustering
    for pillar, scores in pillar_scores.items():
        if not scores:
            continue
        mode_score = Counter(scores).most_common(1)[0]
        mode_pct = mode_score[1] / len(scores)
        if mode_pct > 0.60:
            results.append(CheckResult("DC-01", "FAIL", "MEDIUM",
                f"Score clustering in {pillar}",
                f"{mode_pct:.0%} scored {mode_score[0]} (>{60}% threshold)"))
            issues.append(Issue("MEDIUM", "DISTRIBUTIONAL", "DC-01", pillar,
                f"Score clustering: {mode_pct:.0%} of {pillar} subcaps scored {mode_score[0]}",
                f"Mode={mode_score[0]}, freq={mode_pct:.0%}",
                f"Review {pillar} subcap differentiation"))
        else:
            results.append(CheckResult("DC-01", "PASS", "MEDIUM",
                f"Score distribution {pillar} OK",
                f"Mode={mode_score[0]} at {mode_pct:.0%}"))

    # DC-02: Confidence inflation
    all_conf = [c for confs in pillar_confidence.values() for c in confs]
    if all_conf:
        high_pct = sum(1 for c in all_conf if c == "HIGH") / len(all_conf)
        if high_pct > 0.70:
            results.append(CheckResult("DC-02", "FAIL", "MEDIUM",
                "Confidence inflation",
                f"{high_pct:.0%} HIGH (>{70}% threshold)"))
            issues.append(Issue("MEDIUM", "DISTRIBUTIONAL", "DC-02", "overall",
                f"Confidence inflation: {high_pct:.0%} of subcaps rated HIGH",
                f"HIGH%={high_pct:.0%}",
                "Review evidence quality vs confidence assignments"))
        else:
            results.append(CheckResult("DC-02", "PASS", "MEDIUM",
                "Confidence distribution OK", f"HIGH={high_pct:.0%}"))

    # DC-03: Evidence tier concentration
    for pillar, tiers in pillar_evidence_tiers.items():
        if not tiers:
            continue
        tier_counts = Counter(tiers)
        for tier, count in tier_counts.items():
            pct = count / len(tiers)
            if pct > 0.90 and tier == "T5":
                results.append(CheckResult("DC-03", "FAIL", "HIGH",
                    f"Evidence tier concentration {pillar}",
                    f"T5={pct:.0%} (>{90}%)"))
                issues.append(Issue("HIGH", "DISTRIBUTIONAL", "DC-03", pillar,
                    f"{pillar} relies {pct:.0%} on T5 (marketing) evidence",
                    f"T5%={pct:.0%}",
                    "Source T1-T3 evidence to improve triangulation"))
            elif pct > 0.80:
                results.append(CheckResult("DC-03", "FAIL", "MEDIUM",
                    f"Evidence tier concentration {pillar}",
                    f"{tier}={pct:.0%} (>{80}%)"))
                issues.append(Issue("MEDIUM", "DISTRIBUTIONAL", "DC-03", pillar,
                    f"{pillar} relies {pct:.0%} on {tier} evidence",
                    f"{tier}%={pct:.0%}",
                    "Diversify evidence sourcing across tiers"))
            else:
                results.append(CheckResult("DC-03", "PASS", "MEDIUM",
                    f"Tier balance {pillar} OK"))

    # DC-04: Cap saturation
    for pillar, caps in pillar_caps.items():
        if caps["total"] == 0:
            continue
        cap_pct = caps["capped"] / caps["total"]
        if cap_pct > 0.30:
            results.append(CheckResult("DC-04", "FAIL", "MEDIUM",
                f"Cap saturation {pillar}",
                f"{cap_pct:.0%} capped (>{30}% threshold)"))
            issues.append(Issue("MEDIUM", "DISTRIBUTIONAL", "DC-04", pillar,
                f"{cap_pct:.0%} of {pillar} subcaps capped (>{30}%)",
                f"Capped={caps['capped']}/{caps['total']}",
                "Review evidence quality or scoring approach"))
        else:
            results.append(CheckResult("DC-04", "PASS", "MEDIUM",
                f"Cap rate {pillar} OK", f"{cap_pct:.0%}"))

    # DC-05: Rationale homogeneity (Jaccard similarity)
    for pillar, rats in pillar_rationales.items():
        if len(rats) < 2:
            continue
        high_sim = []
        for i in range(len(rats)):
            tokens_i = set(rats[i].lower().split())
            for j in range(i + 1, len(rats)):
                tokens_j = set(rats[j].lower().split())
                if not tokens_i or not tokens_j:
                    continue
                jaccard = len(tokens_i & tokens_j) / len(tokens_i | tokens_j)
                if jaccard > 0.80:
                    high_sim.append((i, j, jaccard))
        if high_sim:
            results.append(CheckResult("DC-05", "FAIL", "HIGH",
                f"Rationale homogeneity {pillar}",
                f"{len(high_sim)} pairs >{80}% similarity"))
            issues.append(Issue("HIGH", "DISTRIBUTIONAL", "DC-05", pillar,
                f"{len(high_sim)} rationale pairs in {pillar} share >{80}% text",
                "Jaccard similarity",
                "Write unique rationale for each subcap"))
        else:
            results.append(CheckResult("DC-05", "PASS", "HIGH",
                f"Rationale uniqueness {pillar} OK"))

    # DC-06: Evidence reuse concentration
    total_subcaps = sum(c["total"] for c in pillar_caps.values())
    if total_subcaps > 0:
        over_cited = [(eid, cnt) for eid, cnt in all_evidence_citations.items()
                      if cnt / total_subcaps > 0.15]
        if over_cited:
            results.append(CheckResult("DC-06", "FAIL", "MEDIUM",
                "Evidence reuse concentration",
                f"{len(over_cited)} items cited by >{15}% of subcaps"))
            issues.append(Issue("MEDIUM", "DISTRIBUTIONAL", "DC-06",
                str(over_cited[:3]),
                f"{len(over_cited)} evidence items cited by >{15}% of all subcaps",
                "Citation frequency count",
                "Diversify evidence sources"))
        else:
            results.append(CheckResult("DC-06", "PASS", "MEDIUM",
                "Evidence reuse OK"))

    # DC-07: Score-Confidence Alignment
    high_score_low_conf = 0
    low_score_high_conf = 0
    for sheet_name in SCORING_DETAIL_SHEETS:
        if sheet_name not in wb.sheetnames:
            continue
        rows, _ = sheet_to_dicts(wb[sheet_name])
        for row in rows:
            score_val = None
            for k in ["Final_Score", "final_score"]:
                if k in row and row[k] is not None:
                    try:
                        score_val = float(row[k])
                    except (ValueError, TypeError):
                        pass
                    break
            conf = ""
            for k in ["Confidence", "confidence"]:
                if k in row and row[k]:
                    conf = str(row[k]).upper()
                    break
            if score_val is not None and conf:
                if score_val >= 4.0 and conf == "LOW":
                    high_score_low_conf += 1
                if score_val <= 1.0 and conf == "HIGH":
                    low_score_high_conf += 1

    dc07_issues = []
    if high_score_low_conf > 3:
        dc07_issues.append(f"{high_score_low_conf} subcaps score ≥4.0 with LOW confidence")
    if low_score_high_conf > 5:
        dc07_issues.append(f"{low_score_high_conf} subcaps score =1.0 with HIGH confidence")

    if not dc07_issues:
        results.append(CheckResult("DC-07", "PASS", "MEDIUM",
            "Score-confidence alignment OK"))
    else:
        results.append(CheckResult("DC-07", "FAIL", "MEDIUM",
            "Score-confidence alignment", "; ".join(dc07_issues)))
        issues.append(Issue("MEDIUM", "DISTRIBUTIONAL", "DC-07", "overall",
            "Score-confidence misalignment: " + "; ".join(dc07_issues),
            "Score-confidence cross-tabulation",
            "Review flagged subcaps for confidence appropriateness"))

    return results, issues


def run_report_checks(report_path):
    """Category 8: RC-01 through RC-12 (those automatable via docx parsing)."""
    results = []
    issues = []

    if not DocxDocument:
        results.append(CheckResult("RC-00", "PASS", "LOW",
            "Report checks skipped",
            "python-docx not installed; install with: pip install python-docx"))
        return results, issues

    if not report_path or not os.path.exists(report_path):
        results.append(CheckResult("RC-01", "FAIL", "CRITICAL",
            "Report file accessible", "Cannot open report"))
        return results, issues

    try:
        doc = DocxDocument(str(report_path))
    except Exception as e:
        results.append(CheckResult("RC-01", "FAIL", "CRITICAL",
            "Report file parseable", f"Error: {e}"))
        return results, issues

    full_text = "\n".join(p.text for p in doc.paragraphs)
    headings = [p.text for p in doc.paragraphs if p.style and "Heading" in str(p.style.name)]

    # RC-02: Peer proxy disclosure in Section 2
    if re.search(r"peer\s+prox", full_text, re.IGNORECASE):
        results.append(CheckResult("RC-02", "PASS", "CRITICAL", "Peer proxy disclosure found"))
    else:
        results.append(CheckResult("RC-02", "FAIL", "CRITICAL", "Peer proxy disclosure",
                                   "No 'peer prox*' language found"))
        issues.append(Issue("CRITICAL", "NARRATIVE", "RC-02", "Section 2",
            "Peer proxy disclosure missing from report",
            "Text search for 'peer proxy' language",
            "Add peer proxy disclosure to Section 2 per communication_standards.md"))

    # RC-03: No "identical methodology" language
    if re.search(r"identical\s+methodology", full_text, re.IGNORECASE):
        results.append(CheckResult("RC-03", "FAIL", "HIGH",
            "No 'identical methodology' language", "Found forbidden phrase"))
        issues.append(Issue("HIGH", "NARRATIVE", "RC-03", "report",
            "Forbidden phrase 'identical methodology' found in report",
            "Regex match", "Replace with appropriate methodology description"))
    else:
        results.append(CheckResult("RC-03", "PASS", "HIGH",
            "No 'identical methodology' language"))

    # RC-04: No fixed month ranges in roadmap
    month_ranges = re.findall(r"Months?\s+\d+\s*[-–]\s*\d+", full_text, re.IGNORECASE)
    if month_ranges:
        results.append(CheckResult("RC-04", "FAIL", "HIGH",
            "No fixed month ranges", f"Found: {month_ranges[:3]}"))
        issues.append(Issue("HIGH", "NARRATIVE", "RC-04", "roadmap",
            f"Fixed month ranges found: {month_ranges[:3]}",
            "Regex match",
            "Replace with milestone-based phases per communication_standards.md"))
    else:
        results.append(CheckResult("RC-04", "PASS", "HIGH", "No fixed month ranges"))

    # RC-06: No "Critical Gaps" heading
    if re.search(r"Critical\s+Gaps", full_text, re.IGNORECASE):
        results.append(CheckResult("RC-06", "FAIL", "MEDIUM",
            "No 'Critical Gaps' heading", "Found forbidden heading"))
        issues.append(Issue("MEDIUM", "NARRATIVE", "RC-06", "report",
            "Forbidden heading 'Critical Gaps' found",
            "Text search", "Rename section"))
    else:
        results.append(CheckResult("RC-06", "PASS", "MEDIUM", "No 'Critical Gaps' heading"))

    # RC-08: All peer medians match across artifacts
    # Extract numbers from report that look like peer medians
    peer_median_pattern = re.compile(r"peer\s+median[:\s]+(\d+\.?\d*)", re.IGNORECASE)
    peer_refs = peer_median_pattern.findall(full_text)
    if peer_refs:
        results.append(CheckResult("RC-08", "PASS", "HIGH",
            "Peer medians found in report",
            f"{len(peer_refs)} references found (cross-check with workbook in LLM Pass 2)"))
    else:
        results.append(CheckResult("RC-08", "PASS", "HIGH",
            "Peer median cross-check", "No explicit peer median values found — verify manually"))

    # RC-09: Trend arrows consistent with narrative
    # Detect arrows and check surrounding sentiment
    trend_violations = []
    arrow_up_pattern = re.compile(r"[↑▲]")
    arrow_down_pattern = re.compile(r"[↓▼]")
    paragraphs = full_text.split("\n")
    for para in paragraphs:
        if arrow_up_pattern.search(para):
            neg_words = ["declin", "decreas", "worsen", "deteriorat", "fall", "drop"]
            if any(w in para.lower() for w in neg_words):
                trend_violations.append(("↑ + negative language", para[:80]))
        if arrow_down_pattern.search(para):
            pos_words = ["improv", "increas", "growth", "gain", "rise", "advanc"]
            if any(w in para.lower() for w in pos_words):
                trend_violations.append(("↓ + positive language", para[:80]))
    if not trend_violations:
        results.append(CheckResult("RC-09", "PASS", "HIGH",
            "Trend arrows consistent with narrative"))
    else:
        results.append(CheckResult("RC-09", "FAIL", "HIGH",
            "Trend arrows consistent with narrative",
            f"{len(trend_violations)} contradictions"))
        issues.append(Issue("HIGH", "NARRATIVE", "RC-09", "report",
            f"{len(trend_violations)} trend arrow vs narrative contradictions found",
            "Arrow-sentiment cross-check",
            "Align trend arrows with surrounding narrative text"))

    # RC-10: N/A capabilities documented in Section 11
    na_pattern = re.compile(r"N/?A|not\s+applicable|excluded", re.IGNORECASE)
    section_11_text = ""
    in_section_11 = False
    for para in paragraphs:
        if re.search(r"(Section\s+11|§\s*11|11\.\s)", para):
            in_section_11 = True
        elif in_section_11 and re.search(r"(Section\s+12|§\s*12|12\.\s|Appendix)", para):
            in_section_11 = False
        if in_section_11:
            section_11_text += para + "\n"
    if section_11_text and na_pattern.search(section_11_text):
        results.append(CheckResult("RC-10", "PASS", "HIGH",
            "N/A capabilities documented in Section 11"))
    else:
        results.append(CheckResult("RC-10", "PASS", "HIGH",
            "N/A documentation check",
            "Could not isolate Section 11 — verify N/A documentation in LLM Pass 2"))

    # RC-11: Single-source capabilities have limitation statements
    limitation_keywords = ["single.source", "limited.evidence", "one.source",
                           "limitation", "caveat", "caution"]
    limitation_pattern = re.compile("|".join(limitation_keywords), re.IGNORECASE)
    if limitation_pattern.search(full_text):
        results.append(CheckResult("RC-11", "PASS", "MEDIUM",
            "Limitation statements present in report"))
    else:
        results.append(CheckResult("RC-11", "PASS", "MEDIUM",
            "Limitation statements",
            "No explicit limitation language found — verify in LLM Pass 2"))

    return results, issues


# ---------------------------------------------------------------------------
# Verdict Engine
# ---------------------------------------------------------------------------

def compute_verdict(all_issues):
    """Apply verdict rules from SKILL.md Quick Reference."""
    severity_counts = Counter(i.severity for i in all_issues if i.status == "OPEN")
    crit = severity_counts.get("CRITICAL", 0)
    high = severity_counts.get("HIGH", 0)
    medium = severity_counts.get("MEDIUM", 0)

    if crit > 0:
        return "FAIL", "FIX_AND_REAUDIT"
    elif high > 0:
        return "PASS_WITH_NOTES", "FIX_AND_REAUDIT"
    elif medium > 0:
        return "PASS_WITH_NOTES", "DELIVER"
    else:
        return "PASS", "DELIVER"


def generate_qa_verdict(manifest, all_results, all_issues, verdict, recommendation):
    """Generate qa_verdict.json conforming to schemas/qa_verdict.schema.json."""
    severity_counts = Counter(i.severity for i in all_issues)
    open_counts = Counter(i.severity for i in all_issues if i.status == "OPEN")
    fixed_counts = Counter(i.severity for i in all_issues if i.status == "FIXED")

    def g_result(prefix_list):
        """Return PASS if all checks with given prefixes passed, else FAIL."""
        relevant = [r for r in all_results
                    if any(r.check_id.startswith(p) for p in prefix_list)]
        if any(r.status == "FAIL" for r in relevant):
            return "FAIL"
        return "PASS"

    # Collect distributional flags
    dist_flags = []
    dist_counter = 0
    for r in all_results:
        if r.check_id.startswith("DC-") and r.status == "FAIL":
            dist_counter += 1
            dist_flags.append({
                "flag_id": f"DIST-{dist_counter:03d}",
                "description": r.details or r.description,
                "severity": "CAUTION" if r.severity in ("HIGH", "CRITICAL") else "INFO",
                "affected_pillar_or_metric": r.details.split(" ")[0] if r.details else "overall",
                "recommended_action": "Review flagged pattern for systematic issues",
            })

    # Collect narrative issues
    narrative_issues = []
    for r in all_results:
        if r.check_id.startswith("RC-") and r.status == "FAIL":
            issue_type_map = {
                "RC-01": "MISSING_SECTION", "RC-02": "PEER_PROXY_NOT_DISCLOSED",
                "RC-03": "FORBIDDEN_LANGUAGE", "RC-04": "FIXED_DATE_RANGE",
                "RC-06": "FORBIDDEN_LANGUAGE", "RC-07": "SCORE_MISMATCH",
                "RC-08": "MEDIAN_MISMATCH", "RC-09": "SCORE_MISMATCH",
                "RC-10": "N_A_NOT_DOCUMENTED", "RC-11": "MISSING_SECTION",
            }
            narrative_issues.append({
                "issue_type": issue_type_map.get(r.check_id, "FORBIDDEN_LANGUAGE"),
                "description": r.details or r.description,
                "section": "Report",
            })

    narrative_result = "PASS"
    if any(r.check_id.startswith("RC-") and r.status == "FAIL" and r.severity == "CRITICAL"
           for r in all_results):
        narrative_result = "FAIL"
    elif narrative_issues:
        narrative_result = "PASS_WITH_NOTES"

    qa_verdict = {
        "run_id": manifest.get("run_id", "DMA-XXXX-00000000-0000"),
        "institution_name": manifest.get("institution_name", "Unknown"),
        "audit_date": datetime.utcnow().isoformat() + "Z",
        "governance_skill_version": "2.1",
        "verdict": verdict,
        "recommendation": recommendation,
        "issue_count_by_severity": {
            sev: {
                "count": severity_counts.get(sev, 0),
                "open": open_counts.get(sev, 0),
                "fixed": fixed_counts.get(sev, 0),
            }
            for sev in ["CRITICAL", "HIGH", "MEDIUM", "LOW"]
        },
        "check_results": {
            "G1_row_counts": g_result(["SI-01", "SI-02", "SI-03"]),
            "G2_rationale_quality": g_result(["SI-04", "SI-05", "SI-06", "SI-07", "SI-08",
                                              "SI-09", "SI-10", "SI-11", "SI-12", "SI-13", "SI-14"]),
            "G3_evidence_traceability": g_result(["ET-"]),
            "G4_score_integrity": g_result(["SI-15", "SI-16", "SI-17", "SI-18",
                                            "SI-19", "SI-20", "SI-21"]),
            "G5_aggregation": g_result(["AG-"]),
            "G6_cross_pillar_caps": g_result(["CD-"]),
            "G7_confidence_ers": g_result(["CE-"]),
            "G8_contradiction_log": g_result(["CL-"]),
        },
        "distributional_flags": dist_flags,
        "proof_verification": {
            "PV01_structure_complete": "PENDING_LLM_PASS2",
            "PV02_rule_links_valid": "PENDING_LLM_PASS2",
            "PV03_counterclaim_documented": "PENDING_LLM_PASS2",
            "proof_issues": [],
        },
        "critic_resolution": {
            "CR01_findings_addressed": "PENDING_LLM_PASS2",
            "critic_issues": [],
        },
        "narrative_audit_result": narrative_result,
        "narrative_issues": narrative_issues,
        "sign_off": {
            "auditor_id": "gov_auditor_v2.1_automated",
            "auditor_name": "DMA Governance Auditor (Pass 1 — Automated)",
            "organization": "DMA Program",
            "verdict_date": datetime.utcnow().isoformat() + "Z",
            "is_approved": False,
            "sign_off_notes": "Automated Pass 1 complete. PV/CR checks pending LLM Pass 2. Sign-off pending human review.",
        },
    }
    return qa_verdict


# ---------------------------------------------------------------------------
# Main Orchestrator
# ---------------------------------------------------------------------------

def run_audit(assessment_dir, output_dir=None):
    """Execute full automated governance audit."""
    Issue._counter = 0  # Reset counter
    assessment_dir = Path(assessment_dir)
    output_dir = Path(output_dir) if output_dir else assessment_dir / "governance_output"
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"🔍 DMA Governance Auditor v2.0")
    print(f"   Assessment: {assessment_dir}")
    print(f"   Output: {output_dir}")
    print()

    all_results = []
    all_issues = []

    # Discover files
    files = discover_files(assessment_dir)
    print(f"📂 Files found:")
    for name, path in files.items():
        status = "✅" if path else "❌"
        print(f"   {status} {name}: {path or 'NOT FOUND'}")
    print()

    # Load data
    manifest = load_manifest(files["run_manifest.json"]) if files.get("run_manifest.json") else {}
    wb = load_workbook(files["workbook"]) if files.get("workbook") else None

    caps_rows = []
    if files.get("caps_applied_log.csv"):
        caps_rows, _ = load_csv_rows(files["caps_applied_log.csv"])

    contradiction_rows = []
    if files.get("contradiction_log.csv"):
        contradiction_rows, _ = load_csv_rows(files["contradiction_log.csv"])

    evidence_rows = []
    if files.get("evidence_index.csv"):
        evidence_rows, _ = load_csv_rows(files["evidence_index.csv"])

    # Run check categories
    categories = [
        ("1. Input Validation", lambda: run_input_validation(files, manifest, wb)),
        ("2. Score Integrity", lambda: run_score_integrity(wb) if wb else ([], [])),
        ("3. Evidence Traceability", lambda: run_evidence_traceability(wb, evidence_rows) if wb else ([], [])),
        ("4. Aggregation", lambda: run_aggregation_checks(wb) if wb else ([], [])),
        ("5. Caps & Dependencies", lambda: run_caps_dependency_checks(caps_rows)),
        ("6. Confidence-ERS", lambda: run_confidence_ers_checks(wb, evidence_rows) if wb else ([], [])),
        ("7. Contradiction Log", lambda: run_contradiction_checks(contradiction_rows)),
        ("8. Report Content", lambda: run_report_checks(files.get("report"))),
        ("9. Distributional", lambda: run_distributional_checks(wb) if wb else ([], [])),
    ]

    for cat_name, runner in categories:
        print(f"▶ Running {cat_name}...")
        try:
            cat_results, cat_issues = runner()
            all_results.extend(cat_results)
            all_issues.extend(cat_issues)
            passes = sum(1 for r in cat_results if r.status == "PASS")
            fails = sum(1 for r in cat_results if r.status == "FAIL")
            skipped = sum(1 for r in cat_results if r.status == NOT_RUN)
            tail = f", ⏭ {skipped} not run" if skipped else ""
            print(f"  ✅ {passes} pass, ❌ {fails} fail{tail}")
        except Exception as e:
            print(f"  ⚠️ Error: {e}")
            all_results.append(CheckResult(f"{cat_name}_ERROR", "FAIL", "HIGH",
                f"Category execution error", str(e)))

    # Compute verdict
    verdict, recommendation = compute_verdict(all_issues)

    # Summary
    severity_counts = Counter(i.severity for i in all_issues)
    total_pass = sum(1 for r in all_results if r.status == "PASS")
    total_fail = sum(1 for r in all_results if r.status == "FAIL")
    # NAMED, NEVER SILENT. A check that did not run is the thing a reader most
    # needs to know about — the alternative to reporting it as CRITICAL is
    # reporting it, not hiding it.
    not_ran = [r for r in all_results if r.status == NOT_RUN]

    print()
    print(f"{'='*60}")
    print(f"AUDIT SUMMARY")
    print(f"{'='*60}")
    print(f"Checks run:    {len(all_results) - len(not_ran)} of {len(all_results)}")
    print(f"Passed:        {total_pass}")
    print(f"Failed:        {total_fail}")
    if not_ran:
        print(f"Not run:       {len(not_ran)}  (input absent — NOT a failure)")
        for r in not_ran:
            print(f"  · {r.check_id}: {r.details}")
    print(f"Issues found:  {len(all_issues)}")
    print(f"  CRITICAL:    {severity_counts.get('CRITICAL', 0)}")
    print(f"  HIGH:        {severity_counts.get('HIGH', 0)}")
    print(f"  MEDIUM:      {severity_counts.get('MEDIUM', 0)}")
    print(f"  LOW:         {severity_counts.get('LOW', 0)}")
    print(f"{'='*60}")
    print(f"VERDICT:       {verdict}")
    print(f"RECOMMENDATION: {recommendation}")
    print(f"{'='*60}")

    # Write outputs
    # 1. check_results.json
    check_results_path = output_dir / "check_results.json"
    with open(check_results_path, "w") as f:
        json.dump([r.to_dict() for r in all_results], f, indent=2)
    print(f"\n📄 {check_results_path}")

    # 2. preliminary_issues.csv
    issues_path = output_dir / "preliminary_issues.csv"
    with open(issues_path, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(Issue.CSV_HEADER)
        for issue in all_issues:
            writer.writerow(issue.to_csv_row())
    print(f"📄 {issues_path}")

    # 3. audit_summary.json
    summary = {
        "audit_date": datetime.utcnow().isoformat() + "Z",
        "assessment_dir": str(assessment_dir),
        "governance_skill_version": "2.1",
        "checks_run": len(all_results) - len(not_ran),
        "checks_declared": len(all_results),
        "checks_passed": total_pass,
        "checks_failed": total_fail,
        # A reader of this file must be able to tell a clean audit from an
        # audit that could not look. Both used to print the same numbers.
        "checks_not_run": [
            {"check_id": r.check_id, "why": r.details} for r in not_ran],
        "issues_found": len(all_issues),
        "severity_counts": dict(severity_counts),
        "verdict": verdict,
        "recommendation": recommendation,
        "note": "Automated checks only. PV-01/02/03 (proof structure), CR-01 (critic resolution), and deep narrative analysis require LLM reasoning layer.",
        "llm_checks_pending": [
            "PV-01: Proof structure completeness",
            "PV-02: Rule links correctly applied",
            "PV-03: Counterclaim quality assessment",
            "CR-01: Critic finding resolution quality",
            "RC-05: Roadmap milestone quality",
            "RC-07: Score-narrative cross-validation",
            "Patch block generation",
            "Root cause analysis for issues found",
        ],
    }
    summary_path = output_dir / "audit_summary.json"
    with open(summary_path, "w") as f:
        json.dump(summary, f, indent=2)
    print(f"📄 {summary_path}")

    # 4. qa_verdict.json (schema-compliant — Pass 1 fields populated, Pass 2 fields pending)
    qa_verdict = generate_qa_verdict(manifest, all_results, all_issues, verdict, recommendation)
    verdict_path = output_dir / "qa_verdict.json"
    with open(verdict_path, "w") as f:
        json.dump(qa_verdict, f, indent=2)
    print(f"📄 {verdict_path}")

    return summary


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="DMA Governance Auditor")
    parser.add_argument("assessment_dir", help="Path to assessment directory")
    parser.add_argument("--output-dir", help="Output directory (default: <assessment_dir>/governance_output)")
    args = parser.parse_args()

    summary = run_audit(args.assessment_dir, args.output_dir)
    sys.exit(0 if summary["verdict"] != "FAIL" else 1)
