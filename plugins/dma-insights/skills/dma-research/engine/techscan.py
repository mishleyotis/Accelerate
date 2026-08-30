#!/usr/bin/env python3
"""The technographic scan — the fourth deliverable, from the Tech_Register.

    python3 -m engine.techscan record --run R --product ... --vendor ... \
        --layer OPS --status CONFIRMED --method public_document [...]
    python3 -m engine.techscan render --run R [--out DIR]
    python3 -m engine.techscan status --run R

WHAT IT IS. The client package's final outputs are four: the scoring
workbook, the research report, the assessment report, and THIS — the
technographic scan, the register of what technology the client demonstrably
runs, layer by layer, with the basis for every row. The serving-tier
techstack page is produced later by the connector agents; this scan is the
research-stage record they and the assessment draw on, and it ships to the
client folder as its own document.

THE VOCABULARY IS THE CHARTER'S, NOT THE PROTOTYPE'S. Four layers —
OPS · CUST · DATA · INFRA (never L2–L5, which collide with evidence levels)
— and four statuses with CLAIMED present and REQUIRED per row:

    CONFIRMED   independently evidenced (a technographic scan hit AND a
                second source, or a first-party artefact naming it live)
    INFERRED    one indirect signal (a job posting, an integration mention)
    CLAIMED     the client or the vendor says so, nobody else yet
    ABSENT      looked for and established missing — with the search that
                establishes it, because an unsearched estate is not absent
                (the AUD-0115 lesson: 'no register row' and 'confirmed
                absent' are different facts and conflating them over-
                recommended by 28 fit points)

Every row records HOW it was detected (`Detection_Method`), what the
detection rests on (`Detection_Basis`, one clause), and — for CONFIRMED —
the evidence ids that resolve in this run's register. The workbook stays
the substrate: `record` writes the Tech_Register sheet, `render` curates
the .docx and .json FROM it, and the two cannot disagree because there is
only one of them.
"""
from __future__ import annotations

# Runnable both ways: -m engine.techscan, or by path for --help.
if __package__ in (None, ""):  # noqa: E402
    import os as _os
    import sys as _sys
    _sys.path.insert(0, _os.path.dirname(_os.path.dirname(
        _os.path.abspath(__file__))))
    __package__ = "engine"

import argparse
import datetime as _dt
import json
import re
import sys
from collections import Counter
from pathlib import Path

from . import contract as C
from . import runstate
from .workbook import RunWorkbook, _split_ids

#: The deliverable's filenames. The .docx is what a person reads; the .json
#: is what the app ingests (both are classified — the docx by name, the json
#: by the app's package_structured registry).
DOCX_NAME = "Technographic_Scan_{entity}_{date}.docx"
JSON_NAME = "technographic_scan.json"


class ScanRefused(SystemExit):
    pass


def _utcnow() -> str:
    return _dt.datetime.now(_dt.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _slug(s: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "_", (s or "")).strip("_") or "client"


# ── recording, with the vocabulary enforced at the write ────────────────

def record(wb: RunWorkbook, *, product: str, vendor: str | None, layer: str,
           status: str, method: str, basis: str, providers=None, subcaps=None,
           evidence_ids=None, source_urls=None,
           as_of: str | None = None, impact: str | None = None) -> str:
    if layer not in C.TECH_LAYERS:
        raise ScanRefused(f"layer {layer!r} not in {C.TECH_LAYERS} — the "
                          f"prototype's L2-L5 keys collide with evidence "
                          f"levels and are refused by charter")
    if status not in C.TECH_STATUS:
        raise ScanRefused(f"status {status!r} not in {C.TECH_STATUS}")
    if method not in C.TECH_METHODS:
        raise ScanRefused(f"method {method!r} not in {C.TECH_METHODS}")
    if not str(product or "").strip():
        raise ScanRefused("a register row names a PRODUCT; a bare vendor or "
                          "a category is the CG-20 defect")
    if len(str(basis or "").strip()) < 15:
        raise ScanRefused("Detection_Basis is one real clause — what was "
                          "seen, where — not a token")
    provs = [str(x).strip().lower() for x in (providers or []) if str(x).strip()]
    if not provs:
        raise ScanRefused(
            "every register row names the PROVIDER(S) that produced it — "
            f"one or more of {C.TECH_PROVIDERS}. The app's techstack facet "
            "declares its sources as {explorium, clay}; a row that cannot "
            "say who saw it cannot be reconciled against that, and 'a web "
            "search found it' is a provider too (`web`), not an exemption.")
    bad = [x for x in provs if x not in C.TECH_PROVIDERS]
    if bad:
        raise ScanRefused(f"provider(s) {bad} not in {C.TECH_PROVIDERS}")
    eids = [e.strip() for e in (evidence_ids or []) if str(e).strip()]
    if status == "CONFIRMED":
        register = wb.evidence_index()
        if not eids:
            raise ScanRefused(
                "CONFIRMED requires evidence ids that resolve in this run — "
                "a confirmation nobody can open is a claim wearing a stronger "
                "word. Register the source first, or record the row as "
                "INFERRED/CLAIMED, which is what it currently is.")
        dead = [e for e in eids if e.split(":")[0] not in register]
        if dead:
            raise ScanRefused(f"CONFIRMED cites {dead}, which do not resolve "
                              f"in Evidence_Detail")
        if set(provs) <= set(C.TECH_BROKERS):
            raise ScanRefused(
                f"providers are {provs} — brokers only. A broker asserts a "
                f"deployment with no primary source behind it, which is "
                f"exactly CLAIMED. Two brokers agreeing is not corroboration "
                f"either: they resell one crawl. Record this as CLAIMED, or "
                f"add the non-broker provider that saw it independently "
                f"(a vendor announcement, a job posting, the client's own "
                f"page) and cite it.")
    if status == "ABSENT" and "search" not in str(basis).lower() and \
            "scan" not in str(basis).lower() and "0 hits" not in str(basis):
        raise ScanRefused(
            "ABSENT must state the search that establishes the absence — "
            "'no register row' and 'confirmed absent' are different facts, "
            "and conflating them over-recommends the estate (AUD-0115)")
    n = 1 + sum(1 for r in wb.rows("Tech_Register"))
    ts_id = f"TS-{n:03d}"
    wb.append("Tech_Register", {
        "TS_ID": ts_id, "Product": product.strip(),
        "Vendor": (vendor or "").strip() or None,
        "Layer": layer, "Status": status,
        "Evidence_Level": ("L1" if status == "CONFIRMED" else
                           "L2" if status == "INFERRED" else
                           "L3" if status == "CLAIMED" else "L4"),
        "Detection_Basis": basis.strip(), "Detection_Method": method,
        "Providers": ", ".join(dict.fromkeys(provs)),
        "SubCap_IDs": ", ".join(subcaps or []) or None,
        "Evidence_IDs": ", ".join(eids) or None,
        "Source_URLs": ", ".join(source_urls or []) or None,
        "As_Of": as_of or _utcnow()[:10],
        "DMA_Impact": _checked_impact(impact) if impact else None,
    })
    return ts_id


def _providers_of(row) -> list[str]:
    return [x.strip().lower() for x in str(row.get("Providers") or "").split(",")
            if x.strip()]


#: The served contract's own band for `dma_impact` (T3's headline card).
#: Under it the card is a label; over it, nobody reads it.
IMPACT_MIN_WORDS, IMPACT_MAX_WORDS = 40, 90


def _checked_impact(text: str) -> str:
    text = str(text or "").strip()
    n = len(text.split())
    if not (IMPACT_MIN_WORDS <= n <= IMPACT_MAX_WORDS):
        raise ScanRefused(
            f"DMA_Impact is {n} words; the served contract's band is "
            f"{IMPACT_MIN_WORDS}-{IMPACT_MAX_WORDS}. This is the drilldown's "
            f"headline card — the answer to 'so what does running this do to "
            f"the assessment'. Name the pathway: which capability it lifts or "
            f"caps, and what would have to change for that to move.")
    return text


def set_impact(wb: RunWorkbook, ts_id: str, text: str) -> dict:
    """Attach the T3 impact narrative to a row already in the register."""
    checked = _checked_impact(text)
    rows = wb.rows("Tech_Register")
    if not any(str(r["TS_ID"]) == ts_id for r in rows):
        raise ScanRefused(
            f"{ts_id} is not in this run's Tech_Register — "
            f"{', '.join(str(r['TS_ID']) for r in rows) or 'it is empty'}")
    wb.update_row("Tech_Register", "TS_ID", ts_id, {"DMA_Impact": checked})
    return {"ts_id": ts_id, "words": len(checked.split())}


#: THREE answers, not two. The served contract asks for "one row per peer
#: INCLUDING the peers you could not establish (deployed: null)", and AG-04's
#: arithmetic counts only `deployed is True` against the rows it was given.
#: A binary yes/no transcribes a peer the run searched and could not settle
#: as a positive "not deployed" — a fabricated finding about a named
#: institution, which is the class AG-04 exists to stop, and a sentinel that
#: looks like data, which invariant 9 forbids.
PEER_DEPLOYED = {"yes": True, "no": False, "unknown": None}


def peer_record(wb: RunWorkbook, *, ts_id: str, peer: str,
                deployed: bool | None, basis: str,
                source_url: str | None = None,
                as_of: str | None = None) -> dict:
    """One peer, one product, one of THREE answers — with what settles it.

    `deployed=None` is "looked and could not establish", and it is a real
    answer: it keeps the peer in the cohort without asserting anything about
    it, and it keeps `peer_coverage` honest by staying out of the
    denominator.

    T3's peer card is the second reason a reader opens the drawer: not
    "what do they run" but "is running it normal here". The producer used
    to have to research this inside the synthesis session, which is the
    work that gets skipped first under a turn budget — so it is captured
    HERE, where the run is already looking at peers.
    """
    rows = wb.rows("Tech_Register")
    if not any(str(r["TS_ID"]) == ts_id for r in rows):
        raise ScanRefused(f"{ts_id} is not in this run's Tech_Register")
    if len(str(basis or "").strip()) < 15:
        raise ScanRefused(
            "a peer row states the BASIS in a clause — what was seen and "
            "where. 'Deployed: yes' with nothing behind it is the shape "
            "that makes a peer comparison unfalsifiable.")
    if deployed and not source_url:
        raise ScanRefused(
            "a peer recorded as DEPLOYED carries the source that says so. "
            "AG-04 refuses the served row without it, so a peer claim with "
            "no url never reaches the page anyway — record it as not "
            "deployed, or find the source.")
    word = {True: "yes", False: "no", None: "unknown"}[deployed]
    wb.append("Tech_Peer_Deployments", {
        "TS_ID": ts_id, "Peer": str(peer).strip(),
        "Deployed": word,
        "Basis": str(basis).strip(),
        "Source_URL": source_url or None,
        "As_Of": as_of or _utcnow()[:10],
    })
    return peer_state(wb, ts_id)


def peer_state(wb: RunWorkbook, ts_id: str | None = None) -> dict:
    """Peer coverage, COMPUTED — deployed over ESTABLISHED, never stored.

    An unestablished peer stays out of the denominator. Counting it as a
    negative would put a number on the page that reads "one of three peers
    runs this" when the truth is "one does, one does not, and one could not
    be settled" — and invariant 9 is that a derived value is computed or
    null, never a default that looks like data.
    """
    rows = [r for r in wb.rows("Tech_Peer_Deployments")
            if ts_id is None or str(r["TS_ID"]) == ts_id]
    states = [PEER_DEPLOYED.get(str(r["Deployed"]).strip().lower())
              for r in rows]
    established = [x for x in states if x is not None]
    deployed = sum(1 for x in established if x)
    return {"ts_id": ts_id, "peers_examined": len(rows),
            "peers_established": len(established),
            "peers_unknown": len(states) - len(established),
            "peers_deployed": deployed,
            "peer_coverage": (round(deployed / len(established), 3)
                              if established else None)}


def peer_rows(wb: RunWorkbook, ts_id: str) -> list[dict]:
    return [{"peer": r["Peer"],
             "deployed": PEER_DEPLOYED.get(
                 str(r["Deployed"]).strip().lower()),
             "basis": r["Basis"], "source_url": r["Source_URL"],
             "as_of": str(r["As_Of"])[:10] if r["As_Of"] else None}
            for r in wb.rows("Tech_Peer_Deployments")
            if str(r["TS_ID"]) == ts_id]


def scan_state(wb: RunWorkbook) -> dict:
    rows = wb.rows("Tech_Register")
    by_layer = Counter(str(r["Layer"]) for r in rows)
    by_status = Counter(str(r["Status"]) for r in rows)
    by_provider = Counter(p for r in rows for p in _providers_of(r))
    broker_rows = [r for r in rows
                   if set(_providers_of(r)) & set(C.TECH_BROKERS)]
    return {
        "rows": len(rows),
        "by_layer": {l: by_layer.get(l, 0) for l in C.TECH_LAYERS},
        "by_status": {s: by_status.get(s, 0) for s in C.TECH_STATUS},
        "by_provider": {p: by_provider.get(p, 0) for p in C.TECH_PROVIDERS},
        # The owner's requirement, as a number rather than an intention: what
        # share of the register the two contracted brokers actually produced.
        "broker_rows": len(broker_rows),
        "broker_share": (round(len(broker_rows) / len(rows), 3)
                         if rows else None),
        "providers_never_run": [p for p in C.TECH_BROKERS
                                if not by_provider.get(p)],
        # The drilldown's own coverage. On the last measured run before
        # 2026-08-30 both of these were 0 of 32, which is why a click opened
        # onto three empty states.
        "rows_with_impact": sum(1 for r in rows if r.get("DMA_Impact")),
        "rows_with_peers": len({str(r["TS_ID"])
                                for r in wb.rows("Tech_Peer_Deployments")}),
        "layers_never_looked_at": [l for l in C.TECH_LAYERS
                                   if by_layer.get(l, 0) == 0],
    }


# ── the two contracted providers ─────────────────────────────────────────
#
# The app's techstack facet names its sources as exactly {explorium, clay}
# (apps/api/dma_api/computed.py; apps/mcp/server.py's `record_enrichment`
# source vocabulary), and the owner's instruction is that the scan lean on
# those two. They are not symmetric, and pretending they were is how a scan
# reports detections it never made:
#
#   clay       REACHABLE from a session. `find-and-enrich-company` takes a
#              `Tech Stack` data point and `Open Jobs`; the scanner carries
#              those tools. `CLAY_PLAN` fixes the call sequence so it is the
#              same every run and the credit cost is bounded.
#   explorium  THREE DOORS, and only one of them needs a key.
#              (1) The Vibe Prospecting MCP connector, authenticated AT THE
#                  SESSION — no key, no Secret Manager — and already in the
#                  plugin's auto-approve list. Measured 2026-08-23 across
#                  three promoted clients it returned 392 / 357 / 147 named
#                  technologies. This is the door to try first.
#              (2) The INGEST scan, which is a different path and genuinely
#                  dark: apps/worker/dma_worker/enrichment.py records that it
#                  "has no key in Secret Manager, so a scheduled job cannot
#                  reach either". Its darkness says nothing about (1), and
#                  conflating the two cost every run its technographics once.
#              (3) An EXPORT the owner drops in the client folder — the
#                  `*_Explorium_Tech_Stack.xlsx` shape the app's own package
#                  parser has read across five client packages — which
#                  `import_explorium` parses.
#              A row can also arrive through `record --provider explorium`
#              when a session read the connector directly.
#
# A run with neither provider is not a scan with a gap; it is a scan that did
# not happen on its contracted sources, and `scan_state.providers_never_run`
# says so by name.

#: The Clay call plan, fixed so it is the same every run. Ordered: the
#: identifying call first (it creates the task), then the second data point
#: against the entity ids it returned.
CLAY_PLAN = (
    ("mcp__Clay__find-and-enrich-company",
     'companyIdentifier=<registrable domain>, '
     'companyDataPoints=[{"type": "Tech Stack"}]',
     "the register's spine: Clay's own technographic rows for this domain"),
    ("mcp__Clay__add-company-data-points",
     'taskId=<from the call above>, entityIds=[<the company>], '
     'dataPoints=[{"type": "Open Jobs"}]',
     "job postings are the highest-yield DATA and INFRA signal, and they "
     "are INFERRED evidence, never CONFIRMED"),
    ("mcp__Clay__get-task-context",
     "taskId=<from the first call>",
     "NOT optional: the enrichment calls return a HANDLE and every value "
     "arrives only here. Poll until it reads completed. CG-32 is a blocking "
     "gate that exists because this was got wrong once and 20 resolved "
     "contacts were lost between the tool and the producer"),
)

#: The SECOND Clay pass, and the one that used to be skipped. T3's peer card
#: asks whether running this product is normal among the peers the run
#: already froze in `Peer_Benchmarks` — a question about THEIR domains, not
#: the client's. Nothing upstream ever asked it, so the producer was left to
#: research peers inside the synthesis session, which is the work a turn
#: budget drops first.
CLAY_PEER_PLAN = (
    ("mcp__Clay__find-and-enrich-company",
     'companyIdentifier=<the PEER\'s registrable domain>, '
     'companyDataPoints=[{"type": "Tech Stack"}]',
     "one call per peer in Peer_Benchmarks; the peer's estate, not the "
     "client's"),
    ("engine.cli techscan peer-record",
     "--ts <TS-nnn> --peer <name> --deployed|--not-deployed --basis <clause> "
     "[--url <source>]",
     "one row per (product, peer). `peer_coverage` is COMPUTED from these "
     "— deployed over examined — and never typed"),
)

#: Explorium exports carry a category, not our layer. These are matched by
#: substring against the category and the product text together. A row that
#: matches nothing is REPORTED, never guessed into a layer — a mis-layered
#: row silently moves a gap from one pillar to another.
_LAYER_HINTS = (
    ("OPS", ("core banking", "core platform", "loan origin", "los ",
             "servicing", "payment", "ach", "wire", "card process",
             "workflow", "erp", "accounting", "general ledger", "procure",
             "human resources", "hris", "payroll", "back office",
             "treasury", "lending", "mortgage", "teller", "branch")),
    ("CUST", ("crm", "marketing", "digital banking", "online banking",
              "mobile banking", "onboarding", "chat", "contact center",
              "call center", "cms", "content management", "ecommerce",
              "commerce", "customer", "email marketing", "campaign",
              "personaliz", "survey", "experience", "portal", "website",
              "advertis", "social")),
    ("DATA", ("data warehouse", "warehouse", "lakehouse", "data lake",
              "business intelligence", " bi ", "analytics", "etl", "elt",
              "cdp", "customer data platform", "machine learning", "ml ",
              "artificial intelligence", " ai ", "reporting", "dashboard",
              "database", "data quality", "master data", "data catalog")),
    ("INFRA", ("cloud", "hosting", "aws", "azure", "google cloud", "cdn",
               "dns", "identity", "sso", "single sign", "iam", "security",
               "firewall", "endpoint", "network", "monitoring", "observab",
               "devops", "ci/cd", "container", "kubernetes", "backup",
               "disaster recovery", "email security", "vpn")),
)


def layer_for(*texts) -> str | None:
    """The layer a piece of vendor text implies, or None. None is a result."""
    blob = " " + " ".join(str(t or "") for t in texts).lower() + " "
    for layer, hints in _LAYER_HINTS:
        if any(h in blob for h in hints):
            return layer
    return None


#: Sheet names the export has actually used, in priority order, and the
#: header labels that identify a real header row. Both are lifted from the
#: app's own parser (apps/dma-insights/.../parsers/dma_package.py), which
#: has read this shape across five client packages — Alma, WSFS, CalPrivate,
#: Nicola and Odlum — each of which formats it slightly differently.
_EXPORT_HEADER_LABELS = {
    "category", "vendor", "vendor / product", "vendor/product", "technology",
    "product", "company", "name", "layer", "confidence", "tier",
    "evidence id", "evidence_id", "deploy status", "deployment_confirmed",
    "deployment confirmed", "presence", "validation method",
    "validation_method", "source",
}
_EXPORT_VENDOR_KEYS = ("vendor", "company", "technology", "name",
                       "vendor / product", "vendor/product",
                       "vendor / technology", "vendor/technology")


def _is_header_row(row) -> bool:
    cells = [str(c).strip().lower() if c is not None else "" for c in row or ()]
    hits = sum(1 for c in cells if c in _EXPORT_HEADER_LABELS
               or any(c.startswith(k) for k in _EXPORT_HEADER_LABELS))
    return hits >= 2


def read_explorium_export(path) -> dict:
    """Parse an Explorium technographic export into candidate rows.

    Reads, decides nothing. Every row comes back with the layer the text
    implies or `None`, and the caller sees the unmapped ones rather than
    receiving a register that quietly guessed.
    """
    from openpyxl import load_workbook
    path = Path(path)
    wb = load_workbook(path, read_only=True, data_only=True)
    ws = None
    if "Confirmed_Tech_Stack" in wb.sheetnames:
        ws = wb["Confirmed_Tech_Stack"]
    else:
        for sn in wb.sheetnames:
            low = sn.lower()
            if low.endswith("_tech_stack") and not low.endswith(
                    "_recommendations_map"):
                ws = wb[sn]
                break
    ws = ws or wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return {"file": str(path), "sheet": ws.title if ws else None,
                "rows": [], "unmapped": [], "skipped": 0}
    head = 0
    for i in range(min(6, len(rows))):
        if _is_header_row(rows[i]):
            head = i
            break
    headers = [str(c).strip().lower() if c is not None else ""
               for c in rows[head]]

    def pick(cells, keys):
        for k in keys:
            if cells.get(k):
                return str(cells[k]).strip()
        for hk, hv in cells.items():
            if hv and any(k in hk for k in keys):
                return str(hv).strip()
        return None

    out, unmapped, skipped = [], [], 0
    for raw in rows[head + 1:]:
        if not raw or all(c is None for c in raw):
            continue
        cells = {headers[i]: raw[i]
                 for i in range(min(len(headers), len(raw)))}
        vendor = pick(cells, _EXPORT_VENDOR_KEYS)
        if not vendor or vendor.lower().startswith(("source:", "note:", "—")):
            skipped += 1
            continue
        product = pick(cells, ("product", "subcategory")) or vendor
        category = pick(cells, ("category",))
        stated = (pick(cells, ("layer",)) or "").strip().upper()
        layer = stated if stated in C.TECH_LAYERS else layer_for(
            category, product, vendor)
        row = {"product": product, "vendor": vendor, "category": category,
               "layer": layer,
               "confidence": pick(cells, ("confidence", "tier")),
               "as_of": pick(cells, ("as of", "as_of", "date"))}
        (out if layer else unmapped).append(row)
    return {"file": str(path), "sheet": ws.title, "rows": out,
            "unmapped": unmapped, "skipped": skipped}


def import_explorium(wb: RunWorkbook, path, *, status: str = "CLAIMED",
                     as_of: str | None = None) -> dict:
    """Record an Explorium export into the Tech_Register, as CLAIMED.

    CLAIMED is not a demotion, it is the accurate word: the export is a
    third party asserting a deployment with no primary source attached, and
    `record` refuses to let a broker-only row wear CONFIRMED. Promoting one
    is a separate, evidenced act — find the vendor announcement or the job
    posting, then re-record the row with that provider alongside.
    """
    parsed = read_explorium_export(path)
    if not parsed["rows"] and not parsed["unmapped"]:
        raise ScanRefused(
            f"{path}: no technographic rows found. Checked sheet "
            f"{parsed['sheet']!r}. An export that parses to nothing is a "
            f"file problem, not a clean estate — do not record the absence.")
    recorded, failed = [], []
    for r in parsed["rows"]:
        basis = (f"Explorium technographic export row: "
                 f"{r['vendor']} / {r['product']}"
                 + (f", category {r['category']}" if r["category"] else "")
                 + (f", broker confidence {r['confidence']}"
                    if r["confidence"] else "")
                 + f" — read from {Path(parsed['file']).name}")
        try:
            ts = record(wb, product=r["product"], vendor=r["vendor"],
                        layer=r["layer"], status=status,
                        method="technographic_scan", basis=basis,
                        providers=["explorium"],
                        as_of=r["as_of"] or as_of)
            recorded.append(ts)
        except ScanRefused as exc:                              # noqa: PERF203
            failed.append({"product": r["product"], "reason": str(exc)})
    return {"file": parsed["file"], "sheet": parsed["sheet"],
            "recorded": recorded, "refused": failed,
            "skipped_preamble": parsed["skipped"],
            # Reported, never guessed. An unmapped row is work for a person.
            "unmapped_layer": parsed["unmapped"],
            "state": scan_state(wb)}


# ── rendering: docx for people, json for the app ─────────────────────────

def render(wb: RunWorkbook, out_dir, *, force: bool = False) -> dict:
    md = wb.metadata()
    rows = wb.rows("Tech_Register")
    state = scan_state(wb)
    if not rows and not force:
        raise ScanRefused(
            "REFUSED: the Tech_Register is empty. A scan that ran and found "
            "nothing is renderable — record ABSENT rows with the searches "
            "that establish them, or pass --force to render the empty state "
            "as an explicit NOT_RUN document. A blank scan that looks like a "
            "clean scan is the defect.")
    entity = str(md.get("entity_name") or "client")
    date = str(md.get("reference_date"))[:10]
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    # the machine copy first — it is the one the app reads
    doc = {
        "artefact": "technographic_scan",
        "run_id": md.get("run_id"),
        "entity_name": entity,
        "entity_id": md.get("entity_id"),
        "reference_date": md.get("reference_date"),
        "generated_at": _utcnow(),
        "engine_version": md.get("engine_version"),
        "vocabulary": {"layers": list(C.TECH_LAYERS),
                       "statuses": list(C.TECH_STATUS),
                       "providers": list(C.TECH_PROVIDERS),
                       "brokers": list(C.TECH_BROKERS)},
        "counts": state,
        # The app declares this facet's sources as {explorium, clay}. Say
        # per source whether it RAN, so a provider that was never reached is
        # a stated NOT_RUN rather than an absence the reader has to infer
        # from a smaller number.
        "sources": {
            b: {"rows": state["by_provider"].get(b, 0),
                "status": "RAN" if state["by_provider"].get(b) else "NOT_RUN",
                "reason": (None if state["by_provider"].get(b) else
                           ("no Clay enrichment was run for this entity"
                            if b == "clay" else
                            "no Explorium export was present in the client "
                            "folder; Explorium has no MCP server and no key "
                            "in Secret Manager, so it cannot be called"))}
            for b in C.TECH_BROKERS},
        "detections": [
            {"ts_id": r["TS_ID"], "product": r["Product"],
             "vendor": r["Vendor"], "layer": r["Layer"],
             "status": r["Status"], "evidence_level": r["Evidence_Level"],
             "detection_basis": r["Detection_Basis"],
             "detection_method": r["Detection_Method"],
             "providers": _providers_of(r),
             # T3 — what the drilldown renders. Absent is a real state and
             # says so; it is not the same as a row nobody looked at.
             "dma_impact": r.get("DMA_Impact") or None,
             "peer_deployments": peer_rows(wb, str(r["TS_ID"])),
             "peer_coverage": peer_state(
                 wb, str(r["TS_ID"]))["peer_coverage"],
             "subcap_ids": _split_ids(r.get("SubCap_IDs")),
             "evidence_ids": _split_ids(r.get("Evidence_IDs")),
             "source_urls": _split_ids(r.get("Source_URLs")),
             "as_of": r["As_Of"]} for r in rows],
        "not_run": (None if rows else
                    "the scan produced no register rows; this document "
                    "records that it RAN EMPTY, which is different from a "
                    "clean estate"),
    }
    json_path = out_dir / JSON_NAME
    json_path.write_text(json.dumps(doc, indent=2, default=str))

    from docx import Document
    from docx.shared import Pt, RGBColor
    d = Document()
    d.styles["Normal"].font.name = "Calibri"
    d.styles["Normal"].font.size = Pt(10.5)
    d.add_heading("Technographic Scan", level=0)
    d.add_paragraph(entity)
    d.add_paragraph(
        f"Run {md.get('run_id')} · reference date {date} · generated "
        f"{doc['generated_at']} · engine {md.get('engine_version')}")
    d.add_paragraph(
        "Every row below is read from the run's Tech_Register at render "
        "time. Four layers (OPS · CUST · DATA · INFRA), four statuses — and "
        "CLAIMED is a status, not a confirmation: a row says how it was "
        "detected and what that detection rests on, so a reader can weigh "
        "it rather than trust it.")
    d.add_heading("Coverage", level=1)
    t = d.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    for i, c in enumerate(("Layer", "Detections", "Statuses")):
        t.rows[0].cells[i].text = c
    for layer in C.TECH_LAYERS:
        cells = t.add_row().cells
        mine = [r for r in rows if str(r["Layer"]) == layer]
        cells[0].text = layer
        cells[1].text = str(len(mine))
        cells[2].text = ", ".join(
            f"{s}×{sum(1 for r in mine if str(r['Status']) == s)}"
            for s in C.TECH_STATUS
            if any(str(r["Status"]) == s for r in mine)) or "—"
    never = state["layers_never_looked_at"]
    if never:
        p = d.add_paragraph()
        r = p.add_run(
            f"NOT SCANNED: {', '.join(never)} — no detection was attempted "
            f"for these layers. That is a gap in the scan, not a clean "
            f"estate; nothing here may be read as ABSENT.")
        r.bold = True
        r.font.color.rgb = RGBColor(0xB0, 0x00, 0x20)
    d.add_heading("Register", level=1)
    if rows:
        t = d.add_table(rows=1, cols=7)
        t.style = "Light Grid Accent 1"
        for i, c in enumerate(("ID", "Product", "Vendor", "Layer", "Status",
                               "Basis", "As of")):
            t.rows[0].cells[i].text = c
        for r in rows:
            cells = t.add_row().cells
            for i, v in enumerate((r["TS_ID"], r["Product"], r["Vendor"],
                                   r["Layer"], r["Status"],
                                   r["Detection_Basis"], r["As_Of"])):
                cells[i].text = "" if v is None else str(v)
    else:
        p = d.add_paragraph()
        r = p.add_run("NOT RUN — the register is empty and this render was "
                      "forced. No detection, no absence, no estate claim.")
        r.bold = True
        r.font.color.rgb = RGBColor(0xB0, 0x00, 0x20)
    docx_path = out_dir / DOCX_NAME.format(entity=_slug(entity), date=date)
    d.save(docx_path)
    return {"docx": str(docx_path), "json": str(json_path),
            "detections": len(rows), "counts": state, "forced": not rows}


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description=__doc__.split("\n")[0])
    sub = ap.add_subparsers(dest="cmd", required=True)
    for name in ("record", "render", "status", "import-explorium",
                 "clay-plan", "impact", "peer-record", "peers"):
        s = sub.add_parser(name)
        s.add_argument("--run", required=True)
        s.add_argument("--root")
        if name == "import-explorium":
            s.add_argument("--file", required=True,
                           help="the *_Explorium_Tech_Stack.xlsx export")
            s.add_argument("--status", default="CLAIMED",
                           choices=("CLAIMED", "INFERRED"),
                           help="a broker row is CLAIMED; anything stronger "
                                "needs a non-broker provider beside it")
            s.add_argument("--as-of")
        if name == "impact":
            s.add_argument("--ts", required=True, help="TS-nnn")
            s.add_argument("--text", required=True,
                           help=f"{IMPACT_MIN_WORDS}-{IMPACT_MAX_WORDS} words: "
                                f"what running this does to the assessment")
        if name in ("peer-record", "peers"):
            s.add_argument("--ts", required=(name == "peer-record"))
        if name == "peer-record":
            s.add_argument("--peer", required=True)
            g = s.add_mutually_exclusive_group(required=True)
            g.add_argument("--deployed", action="store_true")
            g.add_argument("--not-deployed", dest="not_deployed",
                           action="store_true")
            g.add_argument("--unknown", action="store_true",
                           help="looked and could not establish it — a real "
                                "answer, and the only honest one when the "
                                "search came back ambiguous")
            s.add_argument("--basis", required=True)
            s.add_argument("--url")
            s.add_argument("--as-of")
        if name == "record":
            s.add_argument("--product", required=True)
            s.add_argument("--vendor")
            s.add_argument("--layer", required=True, choices=C.TECH_LAYERS)
            s.add_argument("--status", required=True, choices=C.TECH_STATUS)
            s.add_argument("--method", required=True, choices=C.TECH_METHODS)
            s.add_argument("--basis", required=True)
            s.add_argument("--provider", action="append", default=[],
                           choices=C.TECH_PROVIDERS, required=True,
                           help="who saw it; repeat for more than one")
            s.add_argument("--subcap", action="append", default=[])
            s.add_argument("--evidence-id", action="append", default=[])
            s.add_argument("--url", action="append", default=[])
            s.add_argument("--as-of")
            s.add_argument("--impact",
                           help=f"the T3 drilldown's headline card, "
                                f"{IMPACT_MIN_WORDS}-{IMPACT_MAX_WORDS} words")
        if name == "render":
            s.add_argument("--out")
            s.add_argument("--force", action="store_true")
    a = ap.parse_args(argv)
    run = runstate.locate(a.run, Path(a.root) if a.root else None)
    wb = run.open()
    if a.cmd == "record":
        ts = record(wb, product=a.product, vendor=a.vendor, layer=a.layer,
                    status=a.status, method=a.method, basis=a.basis,
                    providers=a.provider, subcaps=a.subcap,
                    evidence_ids=a.evidence_id,
                    source_urls=a.url, as_of=a.as_of, impact=a.impact)
        print(json.dumps({"ts_id": ts, **scan_state(wb)}, indent=2))
        return 0
    if a.cmd == "render":
        out = render(wb, Path(a.out) if a.out else run.deliverables,
                     force=a.force)
        print(json.dumps(out, indent=2))
        return 0
    if a.cmd == "status":
        print(json.dumps(scan_state(wb), indent=2))
        return 0
    if a.cmd == "import-explorium":
        out = import_explorium(wb, a.file, status=a.status, as_of=a.as_of)
        print(json.dumps(out, indent=2, default=str))
        if out["unmapped_layer"]:
            print(f"\n{len(out['unmapped_layer'])} row(s) matched no layer "
                  f"and were NOT recorded. Layer them by hand with "
                  f"`techscan record --provider explorium`, or leave them "
                  f"out and say so — a guessed layer moves a gap from one "
                  f"pillar to another:", file=sys.stderr)
            for r in out["unmapped_layer"][:20]:
                print(f"    {r['vendor']} / {r['product']} "
                      f"(category {r['category']!r})", file=sys.stderr)
        return 0
    if a.cmd == "impact":
        print(json.dumps(set_impact(wb, a.ts, a.text), indent=2))
        return 0
    if a.cmd == "peer-record":
        print(json.dumps(peer_record(
            wb, ts_id=a.ts, peer=a.peer,
            deployed=(True if a.deployed else None if a.unknown else False),
            basis=a.basis, source_url=a.url, as_of=a.as_of), indent=2))
        return 0
    if a.cmd == "peers":
        out = peer_state(wb, a.ts)
        out["rows"] = peer_rows(wb, a.ts) if a.ts else [
            dict(r) for r in wb.rows("Tech_Peer_Deployments")]
        print(json.dumps(out, indent=2, default=str))
        return 0
    if a.cmd == "clay-plan":
        md = wb.metadata()
        print(f"Clay technographic plan for {md.get('entity_name')}\n")
        for i, (tool, argstr, why) in enumerate(CLAY_PLAN, 1):
            print(f"{i}. {tool}\n     {argstr}\n     why: {why}\n")
        print("Then, once per peer in Peer_Benchmarks — T3's peer card:\n")
        for i, (tool, argstr, why) in enumerate(CLAY_PEER_PLAN, 1):
            print(f"{i}. {tool}\n     {argstr}\n     why: {why}\n")
        print("Then record each returned row:\n"
              "    engine.cli techscan record --provider clay "
              "--method technographic_scan --status CLAIMED ...\n"
              "CLAIMED, because a broker row with no primary source behind "
              "it is a claim — `record` refuses a broker-only CONFIRMED.")
        print(json.dumps(scan_state(wb), indent=2))
        return 0
    return 2


if __name__ == "__main__":
    sys.exit(main())
