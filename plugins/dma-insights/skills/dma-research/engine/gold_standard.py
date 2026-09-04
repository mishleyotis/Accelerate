#!/usr/bin/env python3
"""The gold standard a client package must meet BEFORE anyone believes it.

    python3 -m engine.gold_standard workbook <scoring_workbook.xlsx>
    python3 -m engine.gold_standard report   <report.docx> [--template t.docx] [--kind research|assessment]
    python3 -m engine.gold_standard package   <client_folder> [--json]

CALIBRATED TO THE REFERENCE PACKAGE: Golden 1 Credit Union (DMA-2026-GOLDEN1-001),
named by the engagement owner as "the best gold standard so far" (2026-09-01). Every
threshold below is what that package meets, not an invention.

WHY THIS EXISTS (the goeasy-Ltd findings). The engine already had gates — `validator`
checks the RESEARCH workbook's shape, `completeness` checks tabs, `quality` measures
content. Each is real, and NONE of them is the gold standard: the reference ASSESSMENT
workbook (43 sheets, weighted rollups, an Executive_Summary dashboard, Firmographics,
Focus_Areas, an Issue_Register, a Coverage sheet that discloses evidence gaps as
"Unknown", M-band labels) FAILS the research validator's rule 2/4 outright, because it
is a different and richer artefact produced by the assessment stage, not the research
stage. goeasy shipped the 23-sheet RESEARCH workbook hand-scored in place — which is
why it lacked the dashboard, the firmographics tab, the weighted rollups and the
coverage disclosure the reference carries. THE ROOT CAUSE was skipping the assessment
stage that builds the gold-standard workbook, and then having no single gate that knew
what that workbook should contain. This is that gate.

FINDINGS -> GATES (full register: docs/goeasy-findings-register.md):
  GSY-01 unscored cells render blank        -> GS-WB-SCORES (every subcap numeric 1..5)
  GSY-02 zero / hedge in a value column     -> GS-WB-NOZERO, GS-WB-NOHEDGE
  GSY-03 blank SubCap_Name                   -> GS-WB-NAMES
  GSY-04 peer scores "Not established"       -> GS-WB-PEERS, GS-RPT-NOHEDGE
  GSY-05 report built as a blank docx        -> GS-RPT-BRANDING (branded header kept)
  GSY-06 report missing template sections    -> GS-RPT-SECTIONS
  GSY-07 leftover {{template tokens}}        -> GS-RPT-NOTOKENS
  GSY-08 shallow / thin evidence use         -> GS-RPT-CITATIONS, GS-RPT-LENGTH
  GSY-09 no AI-and-data overlay              -> GS-RPT-AIOVERLAY (assessment, per pillar)
  GSY-10 recommendations with no rebuttal    -> GS-RPT-REBUTTALS
  GSY-11 a fifth (Transformational) BAND     -> GS-RPT-BANDS  (M1..M5 the scale is fine)
  GSY-12 grains missing / duplicated         -> GS-WB-GRAINS
  GSY-13 numbers drift report<->workbook     -> GS-RPT-RECONCILE
  GSY-14 not conducive for app ingestion     -> GS-ING-*
  GSY-15 wrong artefact: RESEARCH workbook    -> GS-WB-STAGE (the 43-sheet assessment set)
         shipped where the ASSESSMENT workbook belonged
  GSY-16 coverage hidden, gaps proxied silently -> GS-WB-COVERAGE (Coverage discloses
         Scored / Unknown_EvidenceGap / Coverage_Pct; Executive_Summary headlines it)
  GSY-17 no Executive_Summary dashboard       -> GS-WB-DASHBOARD
  GSY-18 no 5-year financial trajectory        -> GS-WB-FINANCIALS, GS-RPT-FINANCIALS
         (depth: >=5 fiscal years of real metrics, in a Financial_Trends sheet or
          dispersed as the reference carries it; the report renders it with a trend)
"""
from __future__ import annotations

import json
import math
import re
import sys
import zipfile
from pathlib import Path

#: The pinned templates and the measured reference (references/templates/).
#: Every threshold below is read against gold_reference.json by
#: tests/skills/research_engine/test_gold_reference.py: a floor this gate
#: demands that the Golden 1 package itself would fail is a floor nobody
#: measured, and is refused by the suite.
_TEMPLATES = Path(__file__).resolve().parents[3] / "references" / "templates"


def _pinned_sections(kind: str) -> list[str]:
    """`N. Heading` for every numbered section the pinned Doc carries."""
    try:
        from . import report_spec as RS
    except Exception:            # noqa: BLE001 — the gate must still run
        return []
    key = "assessment" if kind == "assessment" else "client_research"
    return RS.numbered_headings(key)


def gold_reference() -> dict:
    p = _TEMPLATES / "gold_reference.json"
    return json.loads(p.read_text(encoding="utf-8")) if p.is_file() else {}

# ── the verdict shape ────────────────────────────────────────────────────

class Finding(dict):
    def __init__(self, code: str, detail: str, prevents: str = "", **kw):
        super().__init__(code=code, detail=detail, prevents=prevents, **kw)

    def __str__(self):
        p = f" [{self['prevents']}]" if self.get("prevents") else ""
        return f"{self['code']}: {self['detail']}{p}"


# The literals a client-facing deliverable must never ship — every one appeared in a
# shipped goeasy draft. Matched as a standalone value or leading phrase, casefolded.
BANNED_HEDGES = (
    "not established this run", "to be established at the platform",
    "to be established at surface", "surface-production stage",
    "no score yet", "no score exists yet", "queued for enrichment",
    "tbd", "todo", "placeholder", "lorem ipsum", "coming soon",
)
# A FIFTH band is the invariant breach — not the M1..M5 maturity SCALE, which the
# reference package uses throughout ("2.25 (M2)"). Only a reachable 5th band word.
BANNED_BAND_WORDS = ("transformational",)

# The gold-standard ASSESSMENT workbook's sheet set (from the Golden 1 reference).
# A workbook missing these is a RESEARCH workbook shipped where an assessment belongs.
GOLD_SHEETS = (
    "Executive_Summary", "P1_Subcap_Scoring", "P2_Subcap_Scoring",
    "P3_Subcap_Scoring", "P4_Subcap_Scoring", "Pillar_Summary", "Category_Detail",
    "Coverage", "Peer_Benchmarks", "Firmographics", "Focus_Areas",
    "Issue_Register", "Recommendations", "Tech_Register",
)
# The dashboard's own fields — the numbers an executive reads first.
EXEC_FIELDS = ("Institution", "Sub-Vertical", "Evidence Mode", "Overall Maturity",
               "Peer Median", "Gap to Peer", "Subcaps Scored",
               "Evidence Gaps", "Headline")

WORKBOOK_EMPTY_OK = {"Source_URLs"}
#: Columns that are legitimately blank on SOME rows of an engine-built
#: assessment workbook, by sheet — structurally optional fields whose emptiness
#: is a readable state, not an unfinished cell. Measured 2026-09-03: a flat
#: exemption of `Source_URLs` alone flagged every ABSENT firmographic's Value,
#: every peer row without quartiles and every issue without a cap, so the
#: engine's own artefact emitted GS-WB-EMPTY by construction.
#: The scoring sheets' WORKING AREA (columns L onward — the synthesis,
#: ladder and challenge fields) is analysis, not the reader-facing core A–K:
#: `Contradiction_Disposition` is blank when nothing contradicted,
#: `Negative_Ladder` when the cell has evidence, `Ceiling_Band` on a documented
#: absence (null means no score, invariant 9). The reference's own scoring
#: sheets carry the eleven core columns; the gate judges those.
_WORKING_AREA_BLANK_OK = {
    "Dominant_Claim", "Claim_Label", "What_We_Found", "Facet_Coverage",
    "DQ_Works", "DQ_Fails", "DQ_Value", "DQ_Corroborates", "Triangulation",
    "Ceiling_Reasoning", "Why_It_Matters", "DMA_Impact", "DQ_Contradicts",
    "Contradiction_Disposition", "Absence_Claimed", "Proxy_Log",
    "Negative_Ladder", "Discovery_Questions", "Challenge_Verdict",
    "Ceiling_Band", "Uncertainty", "Retrieved_At",
    # core A–K fields that are blank by contract when nothing applies
    "Caps_Applied",
}
OPTIONAL_BLANK: dict[str, set[str]] = {
    "P1_Subcap_Scoring": _WORKING_AREA_BLANK_OK,
    "P2_Subcap_Scoring": _WORKING_AREA_BLANK_OK,
    "P3_Subcap_Scoring": _WORKING_AREA_BLANK_OK,
    "P4_Subcap_Scoring": _WORKING_AREA_BLANK_OK,
    "Firmographics": {"Value", "Unit", "As at", "Evidence", "Conf.", "Reason", "Route"},
    "Peer_Benchmarks": {"Peer_Names", "Peer_Scores", "Peer_P25", "Peer_P75",
                        "Entity_Score", "Peer_Median", "Source_Cell", "As_Of",
                        "Category_Name"},
    "Category_Detail": {"Gap_to_Peer", "Priority_Tier", "Peer_Median", "Coverage",
                        "Priority_Score"},
    "Pillar_Summary": {"Gap_to_Peer", "Peer_Median"},
    "Issue_Register": {"Cap", "As_Of"},
    "Tech_Register": {"Evidence_IDs", "Source_URLs", "SubCap_IDs", "As_Of",
                      "Providers", "DMA_Impact"},
    # Projected from the report's REC cards (`engine.grains recommendations`):
    # a card that names no single category, owner or horizon projects blanks.
    "Recommendations": {"Owner", "Horizon", "Category_ID"},
    "Focus_Areas": {"Currency_Note", "Page"},
    "Coverage": {"Verdict"},
}
#: Per-row rule: on these sheets a row in an ABSENT/QUARANTINED state may
#: leave its value columns blank — the state IS the value.
_STATE_BLANK_OK = {"Firmographics": ("State", ("ABSENT", "QUARANTINED")),
                   "Tech_Register": ("Status", ("ABSENT",))}
#: A grain row whose SCORE is blank is a pillar or category this engagement
#: did not assess (a focused engagement states its scope); its other value
#: columns are blank by construction and are not unfinished cells.
_UNASSESSED_GRAIN_OK = {"Pillar_Summary": "Score", "Category_Detail": "Score"}
SCORE_SHEETS = ("P1_Subcap_Scoring", "P2_Subcap_Scoring",
                "P3_Subcap_Scoring", "P4_Subcap_Scoring")


def _norm(v) -> str:
    return re.sub(r"\s+", " ", str(v if v is not None else "")).strip()


def _is_hedge(text: str) -> bool:
    t = _norm(text).casefold().strip(" .()[]-—:")
    return bool(t) and any(t == b or t.startswith(b) for b in BANNED_HEDGES)


# ── WORKBOOK gate ────────────────────────────────────────────────────────

def workbook_findings(path) -> list[Finding]:
    import openpyxl
    path = Path(path)
    out: list[Finding] = []
    wb = openpyxl.load_workbook(path, data_only=True)
    have = set(wb.sheetnames)

    # GS-WB-STAGE / GS-WB-DASHBOARD — the assessment (not research) artefact.
    missing = [s for s in GOLD_SHEETS if s not in have]
    if missing:
        out.append(Finding("GS-WB-STAGE",
            f"missing gold-standard sheet(s): {missing} — this looks like a "
            f"research workbook, not the assessment package", "GSY-15/17"))

    def rows(sh):
        ws = wb[sh]; rr = list(ws.iter_rows(values_only=True))
        return (rr[0], rr[1:]) if rr else ((), [])

    # GS-WB-DASHBOARD — the Executive_Summary carries the headline numbers.
    if "Executive_Summary" in have:
        _, data = rows("Executive_Summary")
        fields = {_norm(r[0]).casefold(): _norm(r[1]) for r in data if r and r[0]}
        for f in EXEC_FIELDS:
            if not any(f.casefold() in k for k in fields):
                out.append(Finding("GS-WB-DASHBOARD",
                    f"Executive_Summary missing field ~{f!r}", "GSY-17"))

    # GS-WB-COVERAGE — evidence gaps DISCLOSED (Scored / Unknown / Coverage_Pct), not hidden.
    # Two shapes satisfy it: the reference's own `Coverage` sheet (Category,
    # Subcaps, Scored, Unknown_EvidenceGap, Coverage_Pct) and the engine's
    # `Coverage_Map` (category_id, subcaps, evidenced, evidence_gap,
    # coverage_pct), which is the same disclosure under the contract's names.
    # Measured 2026-09-03: the gate read only the sheet named `Coverage`, whose
    # engine columns are the research FLOORS, so an engine-built assessment
    # workbook could never pass its own gold gate.
    def _discloses(sheet):
        hdr, _ = rows(sheet)
        low = {_norm(h).casefold() for h in hdr}
        gap = any(("unknown" in h) or ("evidence_gap" in h) or ("gap" in h) for h in low)
        pct = any("coverage" in h for h in low)
        return gap and pct
    if "Coverage_Map" in have and _discloses("Coverage_Map"):
        pass
    elif "Coverage" in have and _discloses("Coverage"):
        pass
    elif "Coverage" in have or "Coverage_Map" in have:
        out.append(Finding("GS-WB-COVERAGE",
            "neither Coverage_Map nor Coverage discloses the evidence gap "
            "(Unknown_EvidenceGap / evidence_gap) with a coverage percentage", "GSY-16"))

    # GS-WB-SCORES / GS-WB-NOZERO / GS-WB-NAMES — every subcap valued and named.
    for sh in SCORE_SHEETS:
        if sh not in have:
            continue
        hdr, data = rows(sh); ix = {h: i for i, h in enumerate(hdr)}
        for r in data:
            sid = r[ix.get("SubCap_ID", 0)] if hdr else None
            if not sid:
                continue
            sc = r[ix["Score"]] if "Score" in ix else None
            if not isinstance(sc, (int, float)):
                out.append(Finding("GS-WB-SCORES",
                    f"{sid}: Score not numeric ({_norm(sc)!r})", "GSY-01"))
            elif sc == 0 or not (1.0 <= sc <= 5.0):
                out.append(Finding("GS-WB-NOZERO", f"{sid}: Score {sc} outside 1..5", "GSY-02"))
            if "SubCap_Name" in ix:
                nm = r[ix["SubCap_Name"]]
                if not _norm(nm) or _is_hedge(nm) or _norm(nm).casefold() in ("n/a", "na"):
                    out.append(Finding("GS-WB-NAMES", f"{sid}: SubCap_Name blank/placeholder", "GSY-03"))

    # GS-WB-EMPTY / GS-WB-NOHEDGE — no blank or hedge in a reader-facing sheet.
    for sh in list(SCORE_SHEETS) + ["Pillar_Summary", "Category_Detail", "Coverage",
                                    "Peer_Benchmarks", "Firmographics", "Focus_Areas",
                                    "Issue_Register", "Recommendations", "Tech_Register"]:
        if sh not in have:
            continue
        hdr, data = rows(sh); ncol = len([h for h in hdr if h is not None])
        empt = hedged = 0
        optional = WORKBOOK_EMPTY_OK | OPTIONAL_BLANK.get(sh, set())
        state_rule = _STATE_BLANK_OK.get(sh)
        state_ix = (list(hdr).index(state_rule[0])
                    if state_rule and state_rule[0] in hdr else None)
        grain_col = _UNASSESSED_GRAIN_OK.get(sh)
        grain_ix = list(hdr).index(grain_col) if grain_col in hdr else None
        for r in data:
            if not any(_norm(c) for c in r):
                continue
            row_absent = (state_ix is not None and state_rule is not None
                          and _norm(r[state_ix]).upper() in state_rule[1])
            if grain_ix is not None and not _norm(r[grain_ix]):
                row_absent = True            # an unassessed grain, stated as scope
            for j in range(ncol):
                v = r[j]
                if (v is None or not _norm(v)):
                    if hdr[j] in optional or row_absent:
                        continue
                    empt += 1
                elif _is_hedge(v):
                    hedged += 1
        if empt:
            out.append(Finding("GS-WB-EMPTY", f"{sh}: {empt} empty cell(s)", "GSY-01"))
        if hedged:
            out.append(Finding("GS-WB-NOHEDGE", f"{sh}: {hedged} hedge/placeholder cell(s)", "GSY-04"))

    # GS-WB-GRAINS — pillar rollup present (4 pillars, an OVERALL is fine), categories, recs.
    if "Pillar_Summary" in have:
        _, data = rows("Pillar_Summary")
        ids = [_norm(r[0]) for r in data if r and _norm(r[0])]
        pil = {i for i in ids if re.fullmatch(r"P[1-4]", i)}
        if len(ids) != len(set(ids)):
            out.append(Finding("GS-WB-GRAINS", f"duplicate pillar rows: {ids}", "GSY-12"))
        if pil != {"P1", "P2", "P3", "P4"}:
            out.append(Finding("GS-WB-GRAINS", f"pillars not all present: {sorted(pil)}", "GSY-12"))
    # Recommendations may live in the Recommendations tab OR the Solution_Catalogue
    # (the reference carries the recs in the report §8 and the catalogue).
    rec_rows = sol_rows = 0
    if "Recommendations" in have:
        _, data = rows("Recommendations"); rec_rows = len([r for r in data if r and _norm(r[0])])
    if "Solution_Catalogue" in have:
        _, data = rows("Solution_Catalogue"); sol_rows = len([r for r in data if r and _norm(r[0])])
    if rec_rows < 1 and sol_rows < 1:
        out.append(Finding("GS-WB-GRAINS",
            "no recommendations in Recommendations nor Solution_Catalogue", "GSY-12"))

    # GS-WB-PEERS — a peer benchmark is actually established.
    if "Peer_Benchmarks" in have:
        hdr, data = rows("Peer_Benchmarks")
        real = [r for r in data if r and _norm(r[0])]
        if not real:
            out.append(Finding("GS-WB-PEERS", "Peer_Benchmarks is empty", "GSY-04"))
        else:
            hedge = sum(1 for r in real for c in r if _is_hedge(c))
            if hedge:
                out.append(Finding("GS-WB-PEERS", f"{hedge} hedge cell(s) in Peer_Benchmarks", "GSY-04"))

    # GS-WB-FINANCIALS — a real multi-year financial trajectory is present
    # ("depth and all 5-year trends including 5-year financials", GSY-18). The
    # reference (Golden 1) carries it dispersed across its scoring and evidence
    # sheets, so the floor is depth-of-series, NOT a mandated sheet name: at
    # least one sheet must show >=5 distinct fiscal years co-occurring with
    # financial metrics.
    year_re = re.compile(r"(?<!\d)20[0-3]\d(?!\d)")  # matches FY2020, 2020, 2020-24
    fin_re = re.compile(r"revenue|asset|income|deposit|loan|equity|eps|cagr|"
                        r"net charge|roe|roa|dividend|margin|capital", re.I)
    best_years, best_sheet = 0, None
    for sh in wb.sheetnames:
        yrs, kw = set(), False
        for r in wb[sh].iter_rows(values_only=True):
            for c in r:
                t = _norm(c)
                if not t:
                    continue
                yrs.update(year_re.findall(t))
                kw = kw or bool(fin_re.search(t))
        if kw and len(yrs) > best_years:
            best_years, best_sheet = len(yrs), sh
    if best_years < 5:
        out.append(Finding("GS-WB-FINANCIALS",
            f"no 5-year financial trajectory in the workbook (deepest series: "
            f"{best_years} fiscal year(s) in {best_sheet!r})", "GSY-18"))

    # When a dedicated financial-trends sheet exists it must be a real series:
    # >=5 fiscal-year columns, >=5 metric rows, and a growth/CAGR/trend column.
    fin_sheet = next((s for s in wb.sheetnames if s.lower()
                      in ("financial_trends", "financials", "financial_summary")), None)
    if fin_sheet:
        hdr, data = rows(fin_sheet)
        hdr_n = [_norm(h) for h in hdr]
        if "Fiscal_Year" in hdr_n and "Metric" in hdr_n:
            # The ENGINE's long format (contract v7): one row per
            # (metric, fiscal year); the renderer pivots it wide and computes
            # the CAGR. Depth is distinct years × distinct metrics.
            yi, mi = hdr_n.index("Fiscal_Year"), hdr_n.index("Metric")
            years = {_norm(r[yi]) for r in data if r and _norm(r[yi])}
            metrics = {_norm(r[mi]) for r in data if r and _norm(r[mi])}
            if len(years) < 5:
                out.append(Finding("GS-WB-FINANCIALS",
                    f"{fin_sheet}: only {len(years)} fiscal year(s), need >=5", "GSY-18"))
            if len(metrics) < 3:
                out.append(Finding("GS-WB-FINANCIALS",
                    f"{fin_sheet}: only {len(metrics)} metric(s), need >=3", "GSY-18"))
        else:
            # The reference's WIDE shape: metric rows × fiscal-year columns
            # with an explicit CAGR/growth column.
            yr_cols = [h for h in hdr if year_re.search(_norm(h))]
            has_trend = any(re.search(r"cagr|growth|trend|delta|change", _norm(h), re.I)
                            for h in hdr)
            metric_rows = [r for r in data if r and _norm(r[0])]
            if len(yr_cols) < 5:
                out.append(Finding("GS-WB-FINANCIALS",
                    f"{fin_sheet}: only {len(yr_cols)} fiscal-year column(s), need >=5", "GSY-18"))
            if len(metric_rows) < 5:
                out.append(Finding("GS-WB-FINANCIALS",
                    f"{fin_sheet}: only {len(metric_rows)} metric row(s), need >=5", "GSY-18"))
            if not has_trend:
                out.append(Finding("GS-WB-FINANCIALS",
                    f"{fin_sheet}: no CAGR/growth/trend column", "GSY-18"))
    return out


# ── REPORT gate ──────────────────────────────────────────────────────────

def _docx(path):
    import docx
    d = docx.Document(str(path))
    whole = "\n".join(p.text for p in d.paragraphs)
    h1 = [p.text.strip() for p in d.paragraphs
          if p.style and p.style.name in ("Heading 1", "Title") and p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                whole += "\n" + c.text
    with zipfile.ZipFile(str(path)) as z:
        names = z.namelist()
    fonts = [n for n in names if n.startswith("word/fonts/")]
    chrome = [n for n in names if re.match(r"word/(header|footer)\d*\.xml", n)]
    return whole, h1, fonts, chrome


def _template_sections(template_path):
    _, h1, _, _ = _docx(template_path)
    return [h for h in h1 if re.match(r"^\d+\.", h.strip())]


#: Golden 1's own depth, per subcap, as the fallback when gold_reference.json
#: is unreadable: 47 / 115 distinct citations and 4,910 / 11,633 paragraph
#: words over 690 subcaps.
_GOLD_DEPTH_FALLBACK = {"research": (47, 4910), "assessment": (115, 11633)}
_GOLD_SUBCAPS_FALLBACK = 690


def depth_floors(kind: str, subcaps: int | None = None) -> dict:
    """Citation and word floors for a report over `subcaps` cells, at the
    density the Golden 1 reference meets — never above what the reference
    itself would pass (a flat 60 citations failed Golden 1's own research
    report, which carries 47; measured 2026-09-03). `subcaps` defaults to
    the reference's 690, so a bare `gold_standard report <docx>` holds a
    full-size run to the full Golden 1 depth."""
    kind = "assessment" if kind == "assessment" else "research"
    g = gold_reference()
    try:
        ref_sub = int(g["workbook"]["subcaps"])
        ref_c = int(g["reports"][kind]["distinct_e_ids"])
        ref_w = int(g["reports"][kind]["words_paragraphs"])
    except (KeyError, TypeError, ValueError):
        ref_sub = _GOLD_SUBCAPS_FALLBACK
        ref_c, ref_w = _GOLD_DEPTH_FALLBACK[kind]
    n = int(subcaps) if subcaps else ref_sub
    scale = n / ref_sub
    # The WORD floor is the pinned Doc's own contract (the section LENGTH
    # minima summed: 8,400 assessment / 3,050 research at full size), scaled
    # to the run — the Doc is the format the owner asked for, and Golden 1
    # exceeds it (11,633 / 4,910 paragraph words), so the contract is the
    # floor and the reference proves it reachable. Falls back to the
    # reference's own words when the spec cannot be read.
    try:
        from . import report_spec as RS
        spec_min = RS.SPECS["assessment" if kind == "assessment"
                            else "client_research"].min_words
        words = math.ceil(spec_min * scale)
    except Exception:            # noqa: BLE001 — the gate must still run
        words = math.ceil(ref_w * scale)
    words = min(words, math.ceil(ref_w * scale))     # never above the reference
    return {"citations": max(1, math.ceil(ref_c * scale)), "words": max(1, words),
            "subcaps": n, "reference_subcaps": ref_sub}


def report_findings(report_path, template_path=None, scores=None, kind="auto",
                    subcaps: int | None = None) -> list[Finding]:
    report_path = Path(report_path)
    out: list[Finding] = []
    whole, h1, fonts, chrome = _docx(report_path)
    low = whole.casefold()
    if kind == "auto":
        kind = "assessment" if "assessment" in report_path.name.lower() else "research"

    # GS-RPT-SECTIONS — every numbered section of the PINNED template, by
    # number AND heading. A docx template is accepted for a one-off check, but
    # the default is the pin the engine renders to, so the gate and the
    # renderer cannot disagree about what "the required format" is.
    want = _template_sections(template_path) if template_path else _pinned_sections(kind)
    have_h1 = {}
    for h in h1:
        m = re.match(r"^(\d+)\.\s*(.*)$", h.strip())
        if m:
            have_h1[m.group(1)] = m.group(2).strip()
    for s_ in want:
        m = re.match(r"^(\d+)\.\s*(.*)$", s_.strip())
        if not m:
            continue
        n, head = m.group(1), m.group(2).strip()
        if n not in have_h1:
            out.append(Finding("GS-RPT-SECTIONS", f"missing template section {s_!r}", "GSY-06"))
        elif head and have_h1[n].casefold() != head.casefold() \
                and not have_h1[n].casefold().startswith(head.casefold()):
            out.append(Finding("GS-RPT-SECTIONS",
                f"section {n} is titled {have_h1[n]!r}; the template says {head!r}",
                "GSY-06"))

    tok = re.findall(r"\{\{[^}]*\}\}", whole)
    if tok:
        out.append(Finding("GS-RPT-NOTOKENS", f"{len(tok)} leftover token(s): {tok[:3]}", "GSY-07"))

    for bad in BANNED_HEDGES:
        if bad in ("tbd", "todo", "placeholder"):
            if re.search(rf"\b{re.escape(bad)}\b", low):
                out.append(Finding("GS-RPT-NOHEDGE", f"stub marker {bad!r}", "GSY-04"))
        elif bad in low:
            out.append(Finding("GS-RPT-NOHEDGE", f"{low.count(bad)}x hedge {bad!r}", "GSY-04"))

    # GS-RPT-BANDS — a reachable FIFTH band, not the M1..M5 scale, is the breach.
    for bad in BANNED_BAND_WORDS:
        if re.search(rf"\b{re.escape(bad)}\b", low):
            out.append(Finding("GS-RPT-BANDS", f"off-scale band word {bad!r}", "GSY-11"))

    # GS-RPT-BRANDING — the reference brands via a header, not embedded fonts.
    if not chrome and not fonts:
        out.append(Finding("GS-RPT-BRANDING",
            "no header/footer or embedded font — authored as a blank document", "GSY-05"))

    cites = len(set(re.findall(r"\b(?:E|ENR|PB|TS|INT|US)-\d+", whole)))
    floors = depth_floors(kind, subcaps)
    if cites < floors["citations"]:
        out.append(Finding("GS-RPT-CITATIONS",
            f"{cites} distinct citations (< {floors['citations']}, the Golden 1 "
            f"density over {floors['subcaps']} subcaps)", "GSY-08"))
    words = len(re.findall(r"\w+", whole))
    if words < floors["words"]:
        out.append(Finding("GS-RPT-LENGTH",
            f"{words} words (< {floors['words']}, the Golden 1 density over "
            f"{floors['subcaps']} subcaps)", "GSY-08"))

    # GS-RPT-COVERAGE — the report discloses coverage, as the reference does.
    if "coverage" not in low and "unknown" not in low:
        out.append(Finding("GS-RPT-COVERAGE", "no coverage / evidence-gap disclosure", "GSY-16"))

    if kind == "assessment":
        # One overlay per pillar DEEP DIVE the report carries — four on a full
        # engagement, fewer on a focused one that states its scope (the Doc's
        # §5 is one card per pillar in scope). Counted from the card headings
        # the renderer emits; a report with no deep-dive headings at all is
        # held to the full four.
        dives = len(re.findall(r"pillar deep dive \(p[1-4]\)", low))
        need = dives if 1 <= dives <= 4 else 4
        if low.count("ai and data overlay") < need:
            out.append(Finding("GS-RPT-AIOVERLAY",
                f"AI-and-data overlay x{low.count('ai and data overlay')} "
                f"(need {need}, one per pillar deep dive)", "GSY-09"))
        recs = len(set(re.findall(r"rec-r?\d+", low)))
        rebut = max(low.count("strongest counter"), low.count("rebuttal"))
        if recs and rebut < recs:
            out.append(Finding("GS-RPT-REBUTTALS", f"{rebut} rebuttals for {recs} recs", "GSY-10"))

    # GS-RPT-FINANCIALS — the report renders a multi-year financial trajectory
    # ("depth and all 5-year trends including 5-year financials", GSY-18): >=5
    # distinct fiscal years, a financial metric, and an explicit trend word.
    fyears = set(re.findall(r"(?<!\d)20[0-3]\d(?!\d)", whole))
    has_fin = bool(re.search(r"revenue|asset|income|deposit|loan|equity|eps|cagr|"
                             r"net charge|roe|roa|dividend|margin|capital", low))
    has_trend = bool(re.search(r"cagr|growth|grew|year-over-year|yoy|compound|"
                               r"trajectory|five-year|5-year", low))
    if not (len(fyears) >= 5 and has_fin and has_trend):
        out.append(Finding("GS-RPT-FINANCIALS",
            f"no 5-year financial trajectory ({len(fyears)} yrs, "
            f"fin={has_fin}, trend={has_trend})", "GSY-18"))

    if scores and scores.get("overall") is not None:
        ov = scores["overall"]
        if f"{ov:.2f}" not in whole and f"{ov:.1f}" not in whole:
            out.append(Finding("GS-RPT-RECONCILE", f"overall {ov} not in report", "GSY-13"))
    return out


def _subcap_count(workbook_path) -> int | None:
    """Selected cells in the scoring workbook — the size the report floors
    scale by. None when the workbook cannot be read."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(workbook_path, read_only=True, data_only=True)
        n = 0
        for sh in SCORE_SHEETS:
            if sh in wb.sheetnames:
                for r in wb[sh].iter_rows(min_row=2, values_only=True):
                    if r and r[0] is not None and _norm(r[0]):
                        n += 1
        wb.close()
        return n or None
    except Exception:            # noqa: BLE001 — the gate must still run
        return None


# ── PACKAGE / ingestion gate ─────────────────────────────────────────────

def package_findings(folder) -> list[Finding]:
    folder = Path(folder)
    out: list[Finding] = []
    man = folder / "run_manifest.json"
    if not man.exists():
        out.append(Finding("GS-ING-MANIFEST", "run_manifest.json absent", "GSY-14"))
    wb = list(folder.glob("DMA_Scoring_Workbook_*.xlsx")) or list(folder.glob("*Scoring_Workbook*.xlsx"))
    subcaps = None
    if not wb:
        out.append(Finding("GS-ING-DELIVERABLES", "no scoring workbook at root", "GSY-14"))
    else:
        out += workbook_findings(wb[0])
        subcaps = _subcap_count(wb[0])
    for pat, k in ((("Client_Profile_Research_*.docx", "*Research_Report*.docx"), "research"),
                   (("DMA_Assessment_Report_*.docx", "*Assessment_Report*.docx"), "assessment")):
        hit = []
        for p in pat:
            hit += [x for x in folder.glob(p) if not x.name.startswith("DRAFT_")]
        if not hit:
            out.append(Finding("GS-ING-DELIVERABLES", f"no {k} report at root", "GSY-14"))
        else:
            out += report_findings(hit[0], kind=k, subcaps=subcaps)
    if not list(folder.glob("Technographic_Scan_*.docx")) and not list(folder.glob("*Tech*Scan*.docx")):
        out.append(Finding("GS-ING-SCAN", "no technographic scan deliverable", "GSY-14"))
    return out


# ── CLI ──────────────────────────────────────────────────────────────────

def _print(findings, as_json):
    if as_json:
        print(json.dumps([dict(f) for f in findings], indent=1))
    elif not findings:
        print("GOLD STANDARD: PASS — 0 findings")
    else:
        print(f"GOLD STANDARD: {len(findings)} finding(s)")
        for f in findings:
            print("  -", f)
    return 0 if not findings else 1


def main(argv=None):
    import argparse
    ap = argparse.ArgumentParser(description="Gold-standard gate for DMA deliverables.")
    sub = ap.add_subparsers(dest="cmd", required=True)
    w = sub.add_parser("workbook"); w.add_argument("path"); w.add_argument("--json", action="store_true")
    r = sub.add_parser("report"); r.add_argument("path"); r.add_argument("--template")
    r.add_argument("--scores"); r.add_argument("--kind", default="auto"); r.add_argument("--json", action="store_true")
    r.add_argument("--subcaps", type=int, default=None,
                   help="the run's selected cell count, so the depth floors scale "
                        "to this engagement (default: the reference's 690); "
                        "`package` reads it from the workbook")
    r.add_argument("--workbook", help="read --subcaps from this scoring workbook")
    p = sub.add_parser("package"); p.add_argument("folder"); p.add_argument("--json", action="store_true")
    a = ap.parse_args(argv)
    if a.cmd == "workbook":
        return _print(workbook_findings(a.path), a.json)
    if a.cmd == "report":
        scores = json.loads(Path(a.scores).read_text()) if a.scores else None
        subcaps = a.subcaps or (_subcap_count(a.workbook) if a.workbook else None)
        return _print(report_findings(a.path, a.template, scores, a.kind,
                                      subcaps=subcaps), a.json)
    if a.cmd == "package":
        return _print(package_findings(a.folder), a.json)


if __name__ == "__main__":
    sys.exit(main())
