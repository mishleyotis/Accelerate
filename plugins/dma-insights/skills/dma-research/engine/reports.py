#!/usr/bin/env python3
"""The two client-facing reports, curated from the workbook.

    reports.py --run RUN --report client_research|assessment|both

WHY THIS EXISTS.

  AUD-0003  of the three canonical artefacts, the Assessment Report had NO
      PRODUCER AT ALL ("grep for any render/generate/build_report entry point
      against the pinned v8 template across apps/worker, apps/api, apps/mcp
      and both skill trees returns zero"), and the Client Research Report was
      emitted as `client_profile.md`, which the app's classifier returns None
      for. One artefact could not be made and the other could not be ingested.
  AUD-0052  and the renderer that did exist read five JSON files and zero
      sheets, so "a curated report and a recording workbook can disagree about
      the same assessment and both ship".
  AUD-0033  which is how a delivered report shipped with 6 of 21 cited E-ids
      unresolvable, with the check that would catch it documented and never
      implemented.

Three rules, all enforced before a file is written:

  1. Every section is curated from named workbook sheets. Nothing is read
     from anywhere else.
  2. Every `[E-xxx]` in the rendered text is resolved against Evidence_Detail.
     One dead citation refuses the render. There is no "warn and ship".
  3. Every blocking minimum is measured — section word counts, the report
     word count, the insight-card floor.

A refusal names what is missing and what would fix it, because an unattended
run has to be able to act on it.
"""
from __future__ import annotations

# Runnable both ways. `python3 -m engine.<mod>` is the documented invocation,
# but every audit and every operator reaches for `python3 <path> --help`
# first, and a relative import dies there. Binding __package__ makes the two
# equivalent instead of making one of them a trap.
if __package__ in (None, ""):  # noqa: E402  (must precede the relative imports)
    import os as _os
    import sys as _sys
    _sys.path.insert(0, _os.path.dirname(_os.path.dirname(
        _os.path.abspath(__file__))))
    __package__ = "engine"

import argparse
import datetime as _dt
import json
import re
import sys
from pathlib import Path

import math

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from . import contract as C
from . import quality as Q
from . import report_spec as RS
from . import runstate
from . import template as T
from .workbook import RunWorkbook, _split_ids

CITE_RE = re.compile(r"\[(E-\d+)(?::F\d+)?\]")


class ReportRefused(SystemExit):
    pass


def _words(text: str) -> int:
    return len(re.findall(r"[A-Za-z0-9][A-Za-z0-9'’\-]*", text or ""))


def _slug(s: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "_", (s or "")).strip("_") or "client"


# ── curation ─────────────────────────────────────────────────────────────

def curate(wb: RunWorkbook, spec: RS.ReportSpec) -> dict:
    """Assemble the report's blocks from the workbook, without rendering.

    Separate from rendering so the checks below run against exactly what
    would be written — not against a summary of it."""
    md = wb.metadata()
    narrative = [r for r in wb.rows("Report_Narrative")
                 if str(r.get("Report") or "") == spec.key]
    by_section: dict[str, list[dict]] = {}
    for r in narrative:
        by_section.setdefault(str(r.get("Section_ID") or ""), []).append(r)

    blocks = []
    for sec in spec.sections:
        rows = by_section.get(sec.id, [])
        body = "\n\n".join(str(r.get("Body") or "").strip() for r in rows
                           if str(r.get("Body") or "").strip())
        tables = _tables_for(wb, sec)
        if "Evidence_Detail" in sec.inputs:
            cited = _evidence_cited_table(wb, sorted(set(CITE_RE.findall(body))))
            if cited:
                tables.append(cited)
        declared = [s for s in sec.inputs if s in C.SHEETS]
        empty_inputs = [s for s in declared if not wb.rows(s)]
        # A section has NO SOURCE only when EVERY declared input is empty.
        # A focused engagement legitimately leaves three pillar sheets empty,
        # and calling that "no source" would refuse a correct run — the
        # opposite error to AUD-0107's, and just as wrong. Partially empty
        # inputs are STATED in the document as scope, not treated as a defect.
        no_source = bool(declared) and len(empty_inputs) == len(declared)
        blocks.append({
            "section": sec, "rows": rows, "body": body, "tables": tables,
            "words": _words(body) + sum(t["words"] for t in tables),
            "citations": sorted(set(CITE_RE.findall(body))
                                | {c for t in tables for c in t["citations"]}),
            "empty_inputs": empty_inputs, "no_source": no_source,
            "declared_inputs": declared,
        })
    return {"meta": md, "spec": spec, "blocks": blocks}


#: Every declared input renders as ONE table, titled, with the columns a
#: reader can argue with. Measured 2026-09-03: the old if-chain knew eight
#: sheet names of the ~28 the pinned Docs declare as inputs, so
#: `Peer_Benchmarks`, `Firmographics`, `Focus_Areas`, `Issue_Register`,
#: `Tech_Register`, `Solution_Catalogue`, `Recommendations` … produced no
#: table at all — while the full evidence register was rendered SIX times in
#: one document. A sheet not listed here falls back to its contract columns.
_TABLE_TITLES: dict[str, tuple[str, tuple[str, ...] | None]] = {
    "Coverage": ("Coverage by category (research floors)", None),
    "Coverage_Map": ("Coverage disclosure — scored, unknown, coverage %", None),
    "Pillar_Rollup": ("Pillar rollup (stated)", None),
    "Category_Rollup": ("Category rollup (stated)", None),
    "Pillar_Summary": ("Pillar summary", None),
    "Category_Detail": ("Category detail", None),
    "Firmographics": ("Firmographics", None),
    "Issue_Register": ("Issue register", None),
    "Focus_Areas": ("Client priorities, in the client's words", None),
    "Tech_Register": ("Technology register",
                      ("TS_ID", "Product", "Vendor", "Layer", "Status",
                       "Evidence_Level", "Detection_Basis", "As_Of")),
    "Tech_Peer_Deployments": ("Peer deployments", None),
    "Peer_Benchmarks": ("Peer benchmarks", None),
    "Platform_Peer_Adoption": ("Platform adoption among peers", None),
    "Caps_Applied_Log": ("Caps applied", None),
    "Cap_Triggers": ("Cap rules", None),
    "Pillar_Weights": ("Pillar weights", None),
    "Maturity_Rubric": ("Maturity rubric", None),
    "Catalogue_Meta": ("Catalogue binding", None),
    "Solution_Catalogue": ("Solution catalogue", None),
    "Recommendations": ("Recommendations (projected from §8)", None),
    "Enrichment_Needed": ("Enrichment still needed", None),
    "Entity_Timeline": ("Digital evolution timeline",
                        ("Event_Date", "Title", "Kind", "Signal",
                         "Maturity_Effect", "Evidence_IDs")),
    "Gate_Log": ("Gates run on this assessment",
                 ("Gate", "Scope", "Verdict", "Detail")),
    "Run_Metadata": ("Run metadata", None),
    "Handoff_Lock": ("Catalogue lock", None),
    "Subcap_Scores": ("Subcapability scores",
                      ("subcap_id", "subcap_name", "category", "score",
                       "confidence", "evidence_ceiling", "caps_applied",
                       "ai_applicability", "data_readiness")),
}
#: Sheets a section may declare as an INPUT without rendering as a table —
#: they are the section's own prose source or a per-cell working area.
_NO_TABLE = frozenset({"Report_Narrative", "Search_Log"})


def _tables_for(wb: RunWorkbook, sec: RS.Section, card: str | None = None) -> list[dict]:
    """The workbook-derived tables a section carries.

    These are the CURATION: the numbers in the report are the numbers in the
    sheets, read at render time, so the two cannot disagree. Each declared
    input renders ONCE; the evidence register renders as "Evidence cited in
    this section" (the rows this section actually cites) rather than the
    whole register per section; a pillar card's score table is filtered to
    its pillar; the financial trajectory is pivoted wide with its CAGR."""
    out = []
    for name in sec.inputs:
        if name in _NO_TABLE or name == "Evidence_Detail" or name not in C.SHEETS:
            continue
        if name == "Coverage":
            rows = wb.coverage()
        else:
            rows = wb.rows(name)
        if not rows:
            continue
        if name == "Financial_Trends":
            out.append(_financial_table(rows))
            continue
        title, cols = _TABLE_TITLES.get(name, (name.replace("_", " "), None))
        cols = list(cols or C.SHEETS[name])
        if name == "Subcap_Scores" and card and re.fullmatch(r"P[1-4]", card):
            rows = [r for r in rows
                    if str(r.get("subcap_id") or "").startswith(card)]
            title = f"{title} — {card}"
        if name == "Search_Log":
            title = f"Searches run ({len(rows)})"
        out.append(_table(title, cols, [[r.get(c) for c in cols] for r in rows]))
    return out


def _evidence_cited_table(wb: RunWorkbook, citations: list[str]) -> dict | None:
    """The register rows THIS section cites — the drawer a reader opens from
    the section, not the whole register repeated per section."""
    if not citations:
        return None
    register = wb.evidence_index()
    cols = ["E_ID", "Source_Name", "Tier", "Date_Published", "Recency",
            "Claim_Type", "Origin"]
    rows = [[register[c].get(k) for k in cols] for c in citations if c in register]
    return _table("Evidence cited in this section", cols, rows) if rows else None


def _financial_table(rows: list[dict]) -> dict:
    """Financial_Trends (long) pivoted wide: metric × fiscal year, plus the
    compound annual growth rate over the series — the five-year trajectory
    the gold standard carries (GSY-18), computed at render time."""
    years = sorted({str(r.get("Fiscal_Year") or "") for r in rows if r.get("Fiscal_Year")})
    by_metric: dict[str, dict[str, float]] = {}
    units: dict[str, str] = {}
    cites: set[str] = set()
    for r in rows:
        m = str(r.get("Metric") or "").strip()
        if not m:
            continue
        try:
            by_metric.setdefault(m, {})[str(r.get("Fiscal_Year"))] = float(r.get("Value"))
        except (TypeError, ValueError):
            continue
        units[m] = str(r.get("Unit") or "")
        cites.update(_split_ids(r.get("Evidence_IDs")))
    body = []
    for m, series in by_metric.items():
        vals = [series.get(y) for y in years]
        present = [(y, v) for y, v in zip(years, vals) if v is not None]
        cagr = "—"
        if len(present) >= 2 and present[0][1] and present[0][1] > 0 and present[-1][1] > 0:
            n = len(present) - 1
            cagr = f"{((present[-1][1] / present[0][1]) ** (1 / n) - 1) * 100:.1f}%"
        body.append([m, units.get(m, "")] + [("" if v is None else f"{v:,.1f}") for v in vals]
                    + [cagr, " ".join(f"[{c}]" for c in sorted(cites))])
    return _table(f"Financial trajectory — {len(years)} fiscal years "
                  f"({years[0] if years else ''}–{years[-1] if years else ''})",
                  ["Metric", "Unit"] + years + ["CAGR", "Evidence"], body)


#: A section's declared block, as `narrative.write` requires it in the body.
_BLOCK_LINE = re.compile(r"^\s*##\s+(.+?)\s*$")


def _emit_body(doc, body: str) -> None:
    """Write a section body, promoting its `## ` block lines to Heading2.

    The blocks are not decoration. The app parses a report at Heading2
    grain (`report_parser`) and scopes its vectors from tokens inside those
    headings (`embed._PILLAR_TOKEN`), so a section rendered as one
    undivided run of paragraphs arrives as a single row belonging to no
    pillar — which is what every section did before `Section.blocks`
    existed. Promoting them here is what makes the declared anatomy real in
    the artefact rather than only in the workbook.
    """
    buf: list[str] = []

    def flush():
        if buf:
            doc.add_paragraph("\n".join(buf).strip())
            buf.clear()

    for line in (body or "").splitlines():
        m = _BLOCK_LINE.match(line)
        if m:
            flush()
            doc.add_heading(m.group(1), level=2)
        elif not line.strip():
            flush()
        else:
            buf.append(line)
    flush()


def _table(title, cols, rows) -> dict:
    text = " ".join(str(c) for r in rows for c in r if c is not None)
    return {"title": title, "cols": cols, "rows": rows,
            "words": _words(text),
            "citations": sorted(set(CITE_RE.findall(text)))}


# ── the checks that refuse ───────────────────────────────────────────────

def check(wb: RunWorkbook, curated: dict) -> list[str]:
    spec = curated["spec"]
    register = wb.evidence_index()
    problems: list[str] = []
    total = 0

    from . import narrative as _N
    for b in curated["blocks"]:
        sec = b["section"]
        total += b["words"]
        if b["words"] < _N.min_words_for(wb, sec):
            problems.append(
                f"§{sec.id} {sec.heading}: {b['words']} words against a "
                f"blocking minimum of {_N.min_words_for(wb, sec)}. Write it into "
                f"Report_Narrative (Report={spec.key}, Section_ID={sec.id}).")
        dead = [c for c in b["citations"] if c not in register]
        if dead:
            problems.append(
                f"§{sec.id} {sec.heading}: cites {dead}, which do not resolve "
                f"in Evidence_Detail. A citation that opens an empty drawer "
                f"under the client's name is the AUD-0033 defect.")
        if sec.requires_citation and not b["citations"]:
            problems.append(
                f"§{sec.id} {sec.heading}: carries no citation at all.")
        if b["no_source"]:
            problems.append(
                f"§{sec.id} {sec.heading}: every declared input "
                f"{b['declared_inputs']} is empty, so the section has no "
                f"source at all. Populate them or remove the section from "
                f"the spec.")
        if sec.kind in RS.CARD_KINDS:
            n = len([r for r in b["rows"]
                     if str(r.get("Kind") or "") == sec.kind])
            floor = _N.card_floor_for(wb, sec)
            if n < floor:
                problems.append(
                    f"§{sec.id} {sec.heading}: {n} {sec.kind.replace('_', ' ')}"
                    f"(s) against the template's blocking minimum of "
                    f"{floor}.")
            elif sec.cards_max and n > int(sec.cards_max):
                problems.append(
                    f"§{sec.id} {sec.heading}: {n} {sec.kind}s, the template "
                    f"allows at most {sec.cards_max}.")
        # The countable MINIMUM DATA / MUST NOT rules of the Doc's control
        # block, re-checked on what will actually be rendered.
        whole = "\n".join(str(r.get("Body") or "") for r in b["rows"])
        for why in _N._check_counts(sec, whole, per_card=False):
            problems.append(f"§{sec.id} {sec.heading}: {why}")
        if sec.kind in RS.CARD_KINDS:
            for r in b["rows"]:
                for why in _N._check_counts(sec, str(r.get("Body") or ""),
                                            per_card=True):
                    problems.append(
                        f"§{sec.id} {sec.heading} card "
                        f"{r.get('Card_ID')}: {why}")
        # Functional language on everything a client reads: report bodies
        # are impact prose by definition, so both tiers apply — verdict
        # words and blame constructions alike
        # (references/functional_language.md).
        for r in b["rows"]:
            why = Q.accusatory(str(r.get("Body") or ""), impact_field=True)
            if why:
                problems.append(
                    f"§{sec.id} {sec.heading}: {why}")
                break

    if total < _N.report_min_words_for(wb, spec):
        problems.append(
            f"whole report: {total} words against a blocking minimum of "
            f"{_N.report_min_words_for(wb, spec)}")
    # DEPTH, at the Golden 1 density (owner issue 5). The reference cites 47
    # distinct sources in its research report and 115 in its assessment over
    # 690 subcaps; the floor scales with THIS run's cell count so a focused
    # engagement is held to the same density, not the same absolute number.
    distinct = {c for b in curated["blocks"] for c in b["citations"]}
    floor = citation_floor(wb, spec)
    if len(distinct) < floor:
        problems.append(
            f"whole report: {len(distinct)} distinct citations against a floor "
            f"of {floor} (Golden 1 density × {len(wb.selected_subcaps())} "
            f"subcaps). Cite the evidence base, do not summarise it.")
    return problems


def citation_floor(wb: RunWorkbook, spec: RS.ReportSpec) -> int:
    """Distinct citations owed by a report of this run's size, at the density
    the Golden 1 reference meets — never above what the reference itself
    would pass (tests/skills/research_engine/test_gold_reference.py)."""
    from . import gold_standard as GS
    g = GS.gold_reference()
    kind = "assessment" if spec.key == "assessment" else "research"
    try:
        per_subcap = (g["reports"][kind]["distinct_e_ids"]
                      / g["workbook"]["subcaps"])
    except (KeyError, ZeroDivisionError, TypeError):
        per_subcap = (115 if kind == "assessment" else 47) / 690
    return max(1, math.ceil(per_subcap * len(wb.selected_subcaps())))


# ── rendering ────────────────────────────────────────────────────────────

def render(wb: RunWorkbook, spec: RS.ReportSpec, out_dir: Path,
           *, force: bool = False, qa_dir: Path | None = None) -> dict:
    from . import narrative as N
    # THE STAGE PRECONDITIONS, AT RENDER TIME TOO. Measured 2026-09-03: this
    # function never asked whether the run was scored — `narrative.write`
    # did, but only when handed the run, and `--force` swallowed the one
    # readiness check that remained. So a report could be RENDERED on an
    # unscored workbook. The qa_dir is derived from the workbook's own
    # layout when the caller passes none, so no calling shape skips it.
    qa_dir = qa_dir if qa_dir is not None else Path(wb.path).resolve().parent / "07_qa"
    pre = N.stage_preconditions(wb, spec.key, qa_dir)
    curated = curate(wb, spec)
    problems = check(wb, curated)
    # Every section carries an INDEPENDENT verdict, or the report does not
    # render. AUD-0153: the renderer refused a MISSING section and accepted
    # an unreviewed one, so a report could ship on prose nobody had
    # adversarially read.
    if not force:
        try:
            N.require_ready(wb, spec.key)
        except N.NarrativeRefusal as e:
            problems = list(problems) + [str(e)]
    if (pre or problems) and not force:
        # Every reason at once — the run's readiness AND the sections' own
        # problems — so one pass closes them all (an unattended session acts
        # on a list and stalls on a sentence).
        lines = ([f"the run is not ready for it: {p}" for p in pre]
                 + list(problems))
        raise ReportRefused(
            f"REFUSED: {spec.title} cannot be rendered —\n  - "
            + "\n  - ".join(lines)
            + (f"\nRun `engine.cli narrative preconditions --run <R> --root <ROOT> "
               f"--report {spec.key}` until it prints ready. --force does not "
               f"waive this; it writes a DRAFT_ file that no package accepts."
               if pre else ""))
    problems = list(problems) + list(pre)
    draft = bool(force and problems)
    md = curated["meta"]
    entity = str(md.get("entity_name") or "client")
    name = spec.filename.format(entity=_slug(entity),
                                date=str(md.get("reference_date"))[:10])
    if draft:
        # A forced render is a DRAFT: named so no deliverable glob
        # (assemble.DELIVERABLES, gold_standard.package_findings) can pick it
        # up, banner-marked, and recorded as a FAIL on the Gate_Log.
        name = "DRAFT_" + name
    out = Path(out_dir) / name
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = _shell()
    _style(doc)
    _brand(doc, spec, entity, md)
    if draft:
        _warn(doc, f"DRAFT — NOT A DELIVERABLE: rendered with --force while "
                   f"{len(problems)} precondition(s)/problem(s) were open. "
                   f"Every gap below is marked NO SOURCE.")
        from . import ledger as L
        L.append_gate(wb, gate="REPORT_RENDER", scope=spec.key, verdict="FAIL",
                      detail=("forced draft: " + "; ".join(problems))[:900],
                      blocking=False)
    doc.add_heading(spec.title, level=0)
    p = doc.add_paragraph(entity)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(
        f"Run {md.get('run_id')} · catalogue {md.get('catalogue_version')} "
        f"({str(md.get('catalogue_hash'))[:12]}) · reference date "
        f"{md.get('reference_date')} · engine {md.get('engine_version')}")
    doc.add_paragraph(
        "Every figure and every citation in this document is read from the "
        "assessment workbook at render time. The report and the workbook "
        "cannot disagree, because there is only one of them.")
    _front_matter(doc, wb, spec, md)

    for b in curated["blocks"]:
        sec = b["section"]
        doc.add_heading(f"{sec.id}. {sec.heading}", level=1)
        if sec.kind in RS.CARD_KINDS and b["rows"]:
            _emit_cards(doc, wb, sec, b["rows"])
        elif b["body"]:
            _emit_body(doc, b["body"])
        elif force:
            _warn(doc, f"NO SOURCE — §{sec.id} has no narrative in "
                       f"Report_Narrative and its inputs are "
                       f"{list(sec.inputs)}.")
        for t in b["tables"]:
            doc.add_heading(t["title"], level=2)
            _write_table(doc, t)
        if b["no_source"] and force:
            _warn(doc, f"NO SOURCE — every declared input "
                       f"{b['declared_inputs']} is empty for this section.")
        elif b["empty_inputs"]:
            doc.add_paragraph(
                "Scope: this section draws on "
                + ", ".join(i for i in b["declared_inputs"]
                            if i not in b["empty_inputs"])
                + ". Not in this engagement's scope, and therefore not "
                  "reported on: " + ", ".join(b["empty_inputs"]) + ".")

    doc.add_heading("Sources cited", level=1)
    register = wb.evidence_index()
    used = sorted({c for b in curated["blocks"] for c in b["citations"]})
    if not used:
        _warn(doc, "This report cites nothing.")
    bullet = _style_named(doc, "List Bullet", "List Paragraph")
    for e in used:
        r = register.get(e, {})
        line = (f"[{e}] {r.get('Source_Name')} — {r.get('Source_URL')} "
                f"(tier {r.get('Tier')}, {r.get('Recency')}, published "
                f"{r.get('Date_Published') or 'undated'})")
        if bullet is not None:
            doc.add_paragraph(line, style=bullet)
        else:
            doc.add_paragraph("• " + line)

    doc.save(out)
    return {
        "report": spec.key, "path": str(out), "sections": len(curated["blocks"]),
        "words": sum(b["words"] for b in curated["blocks"]),
        "citations": len(used), "unresolved": [c for c in used
                                               if c not in register],
        "forced": draft, "draft": draft, "problems": problems,
        "shell": str(T.REPORT_SHELL),
    }


def _shell():
    """The branded .docx shell, emptied of its own (older) body.

    goeasy GSY-05: the reports were built as a blank `Document()`, throwing
    the client's fonts, header and footer away; until 2026-09-03 the fix was
    a header STRING synthesised onto a blank document, which passed the
    branding gate by the letter and matched the reference by nothing else.
    The shell carries header1.xml, footer1.xml and the embedded DM Sans faces
    (python-docx cannot embed a font into a blank document); its body — an
    older twelve-section outline — is cleared, and the pinned Docs' sections
    are emitted into its styles. Format from the Docs; chrome from the shell.
    """
    if not T.REPORT_SHELL.is_file():
        raise ReportRefused(
            f"REFUSED: the branded report shell is missing at {T.REPORT_SHELL}; "
            f"the plugin ships it, so this is a partial install — reinstall.")
    doc = Document(str(T.REPORT_SHELL))
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)
    return doc


def _style(doc) -> None:
    """The shell's own styles carry the brand; nothing is overridden here.
    The size is pinned so a shell edit cannot silently shrink the body."""
    st = _style_named(doc, "Normal")
    if st is not None and st.font.size is None:
        st.font.size = Pt(10.5)


def _brand(doc, spec, entity: str, md: dict) -> None:
    """Fill the shell's header and footer with THIS report's identity. The
    chrome itself (header1.xml / footer1.xml, the DM Sans faces) is the
    shell's — goeasy GSY-05, gold gate GS-RPT-BRANDING."""
    sec = doc.sections[0]
    h = sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
    for r in list(h.runs)[1:]:
        r._r.getparent().remove(r._r)
    if h.runs:
        h.runs[0].text = f"{spec.title} | ZENNIFY · {entity}"
    else:
        h.text = f"{spec.title} | ZENNIFY · {entity}"
    f = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    tail = (f" · Run {md.get('run_id')} · catalogue {md.get('catalogue_version')} "
            f"({str(md.get('catalogue_hash'))[:8]}) · Confidential")
    if f.runs:
        f.add_run(tail).font.size = Pt(8)
    else:
        f.text = "Zennify" + tail


def _front_matter(doc, wb, spec, md: dict) -> None:
    """The Doc's two unnumbered front sections, filled from the run.

    'Document Control and Catalogue Binding' resolves every value from the
    workbook — nothing typed — and 'Surface Alignment' is the Doc's own
    section -> app-surface table, rendered from the pinned spec so it cannot
    drift from what the sections declare they feed. Unnumbered Heading1s are
    front matter to the app's parser and are never stored as sections."""
    doc.add_heading("Document Control and Catalogue Binding", level=1)
    doc.add_paragraph(
        "Every value below is resolved at render time from the run's own "
        "workbook and the active catalogue. Nothing here is typed by hand.")
    tax = C.taxonomy()
    lock = wb.handoff_lock()
    rows = [
        ["Catalogue version", str(md.get("catalogue_version"))],
        ["Catalogue content hash", str(md.get("catalogue_hash"))],
        ["Structure counts", f"{tax.n_pillars} pillars / {tax.n_categories} "
                             f"categories / {tax.n_capabilities} capabilities / "
                             f"{tax.n_cells} subcapabilities"],
        ["Sub-vertical", str(md.get("sub_vertical") or "")],
        ["Evidence mode", str(md.get("evidence_mode") or "")],
        ["Scope", f"{md.get('scope_mode')} — {md.get('subcaps_selected')} "
                  f"subcapabilities selected"],
        ["Reference date", str(md.get("reference_date"))],
        ["Workbook contract", f"{md.get('workbook_contract')} · engine "
                              f"{md.get('engine_version')}"],
        ["Scoring workbook", wb.path.name],
        ["Template binding", str(md.get("template_binding") or "UNBOUND")],
        ["Peer set", str(lock.get("locked_peer_set") or "not locked")],
        ["Stage", str(md.get("stage") or "research")],
    ]
    in_scope = sorted({c[:2] for c in wb.selected_subcaps()})
    out_of_scope = [p for p in ("P1", "P2", "P3", "P4") if p not in in_scope]
    rows.append(["Pillars in scope", ", ".join(in_scope) or "none"])
    _write_table(doc, {"cols": ["Field", "Value"], "rows": rows})
    if out_of_scope:
        # A focused engagement STATES its scope rather than refusing: the
        # sheets it leaves empty are named here, once, so a reader does not
        # take an unassessed pillar for a silent one.
        doc.add_paragraph(
            "Pillars assessed in this engagement: " + ", ".join(in_scope)
            + ". Not in this engagement's scope, and therefore not reported on: "
            + ", ".join(f"{p}_Subcap_Scoring" for p in out_of_scope) + ".")
    doc.add_heading("Surface Alignment", level=1)
    doc.add_paragraph(
        "This report is one input to the DMA Insights app. Each section "
        "feeds the named app surfaces; producing a section without knowing "
        "what it feeds is how a report and a dashboard end up disagreeing "
        "under the same client name.")
    _write_table(doc, {
        "cols": ["Section", "Feeds app surface"],
        "rows": [[f"{s.id}. {s.heading}",
                  ", ".join(s.surfaces) or "No served surface; internal"]
                 for s in spec.sections]})


def _card_heading(wb, sec, row) -> str:
    """The Doc's own heading for one card — '5.N Pillar deep dive (P1): …'
    with the served score and median, or 'REC-NN: Title'."""
    card = str(row.get("Card_ID") or "").strip()
    title = str(row.get("Heading") or "").strip()
    if sec.kind == "pillar":
        n = card[1:] if card.startswith("P") else "?"
        score = median = gap = "—"
        for r in wb.rows("Pillar_Rollup"):
            if str(r.get("pillar_id") or "").strip() == card:
                score = str(r.get("score") or "—")
                median = str(r.get("peer_median") or "—")
                gap = str(r.get("gap") or "—")
        name = C.PILLAR_NAMES.get(card, title or card)
        # '(P1)' is the token the app's parser (DEEP_DIVE_PILLAR) and the
        # embedder (_PILLAR_TOKEN) scope on; it stays in the heading.
        return (f"{sec.id}.{n} Pillar deep dive ({card}): {name} — score "
                f"{score} against median {median} ({gap})")
    if title and title != sec.heading and not title.startswith(card):
        return f"{card}: {title}"
    return title or card


def _emit_cards(doc, wb, sec, rows) -> None:
    """A list section: one Heading2 per card, its blocks as Heading3s.

    That is the Doc's shape ('## 5.N …' / '### Capability scorecard';
    '## REC-NN: …' / '#### Root cause') and the grain the app parses: rows
    land per Heading2, so a pillar's four blocks arrive under one heading
    carrying its pillar token, and a recommendation's under its REC id."""
    def key(r):
        c = str(r.get("Card_ID") or "")
        m = re.search(r"(\d+)$", c)
        return (0, int(m.group(1))) if m else (1, c)
    for r in sorted(rows, key=key):
        doc.add_heading(_card_heading(wb, sec, r), level=2)
        buf: list[str] = []

        def flush():
            if buf:
                doc.add_paragraph("\n".join(buf).strip())
                buf.clear()
        for line in str(r.get("Body") or "").splitlines():
            m = _BLOCK_LINE.match(line)
            if m:
                flush()
                doc.add_heading(m.group(1), level=3)
            elif not line.strip():
                flush()
            else:
                buf.append(line)
        flush()


def _warn(doc, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = RGBColor(0xB0, 0x00, 0x20)


def _style_named(doc, *names):
    """The first of `names` the document defines, by NAME (the shell's
    `Normal` is registered as `normal`; python-docx's id lookup is
    deprecated and case-sensitive). None when the shell has none of them."""
    want = {n.casefold() for n in names}
    for s in doc.styles:
        if s.name and s.name.casefold() in want:
            return s
    return None


def _write_table(doc, t) -> None:
    table = doc.add_table(rows=1, cols=len(t["cols"]))
    st = _style_named(doc, "Light Grid Accent 1", "Table Grid", "TableNormal",
                      "Normal Table")
    if st is not None:
        table.style = st
    for i, c in enumerate(t["cols"]):
        cell = table.rows[0].cells[i]
        cell.text = str(c)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in t["rows"]:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "" if v is None else str(v)


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description=__doc__.split("\n")[0])
    ap.add_argument("--run", required=True)
    ap.add_argument("--root")
    ap.add_argument("--report", default="both",
                    choices=["client_research", "assessment", "both"])
    ap.add_argument("--out")
    ap.add_argument("--spec", help="JSON export of a template's control blocks")
    ap.add_argument("--force", action="store_true",
                    help="render a DRAFT_ file with every gap marked NO SOURCE "
                         "and a REPORT_RENDER FAIL on the Gate_Log; never a "
                         "deliverable — no package glob picks it up")
    a = ap.parse_args(argv)
    r = runstate.locate(a.run, Path(a.root) if a.root else None)
    wb = r.open()
    out_dir = Path(a.out) if a.out else r.deliverables
    keys = list(RS.SPECS) if a.report == "both" else [a.report]
    override = RS.from_json(json.loads(Path(a.spec).read_text())) if a.spec else None
    results = []
    for k in keys:
        spec = override if override and override.key == k else RS.SPECS[k]
        results.append(render(wb, spec, out_dir, force=a.force))
    print(json.dumps(results, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
